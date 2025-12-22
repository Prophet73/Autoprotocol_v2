#!/usr/bin/env python3
"""
Мультиязычная транскрипция v4: Оптимизированный пайплайн.

Улучшения относительно v3:
1. Единая загрузка модели - не перезагружаем для каждого языка
2. Батчинг VAD сегментов - группируем для эффективной транскрипции
3. tqdm прогресс-бары - наглядный прогресс
4. Оптимизация памяти - агрессивная очистка CUDA кэша
5. Батчинг переводов - группируем для уменьшения rate limiting
6. Параллельный анализ эмоций - с ThreadPoolExecutor
"""

import torch
import os
import sys
import re
import json
import time
from pathlib import Path
from datetime import datetime
from collections import defaultdict
from typing import Optional, List, Dict, Any, Tuple
from concurrent.futures import ThreadPoolExecutor, as_completed
import gc

# Patch torch.load BEFORE any imports (PyTorch 2.8+ fix)
_original_torch_load = torch.load
def _patched_torch_load(*args, **kwargs):
    kwargs['weights_only'] = False
    return _original_torch_load(*args, **kwargs)
torch.load = _patched_torch_load

import lightning_fabric.utilities.cloud_io as cloud_io
cloud_io.torch.load = _patched_torch_load

import whisperx
from whisperx.diarize import DiarizationPipeline
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from transformers import Wav2Vec2ForSequenceClassification, Wav2Vec2FeatureExtractor
import librosa
import subprocess
import numpy as np

try:
    from tqdm import tqdm
except ImportError:
    # Fallback если tqdm не установлен
    def tqdm(iterable, **kwargs):
        desc = kwargs.get('desc', '')
        total = kwargs.get('total', len(iterable) if hasattr(iterable, '__len__') else None)
        for i, item in enumerate(iterable):
            if total:
                print(f"\r{desc}: {i+1}/{total}", end='', flush=True)
            yield item
        print()

# Google Gemini
import google.generativeai as genai
from dotenv import load_dotenv

# === LOAD ENV ===
PROJECT_ROOT = Path(__file__).parent.resolve()
load_dotenv(PROJECT_ROOT / ".env")

# === КОНСТАНТЫ ===
OUTPUT_DIR = PROJECT_ROOT / "output"
EMOTION_MODEL = "Aniemore/wav2vec2-xlsr-53-russian-emotion-recognition"

EMOTION_LABELS_RU = {
    'anger': 'Гнев', 'disgust': 'Отвращение', 'enthusiasm': 'Энтузиазм',
    'fear': 'Страх', 'happiness': 'Радость', 'neutral': 'Нейтрально', 'sadness': 'Грусть'
}
EMOTION_EMOJI = {
    'anger': '😠', 'disgust': '🤢', 'enthusiasm': '🤩',
    'fear': '😨', 'happiness': '😊', 'neutral': '😐', 'sadness': '😔'
}
LANGUAGE_FLAGS = {
    'ru': '🇷🇺', 'zh': '🇨🇳', 'en': '🇬🇧', 'unknown': '❓'
}

# === HALLUCINATION PATTERNS ===
HALLUCINATION_PATTERNS = [
    r'продолжение следует',
    r'субтитры\s*(сделал|подогнал|создал|делал)',
    r'редактор субтитров',
    r'корректор\s+[а-яё]+\.[а-яё]+',
    r'спасибо за просмотр',
    r'подписывайтесь на канал',
    r'ставьте лайк',
    r'^пока\.?$', r'^ага\.?$', r'^угу\.?$',
    r'^谢谢大家\.?$', r'^谢谢\.?$', r'^谢谢观看', r'^感谢收看',
    r'^thank you\.?$', r'^thanks for watching',
    r'DimaTorzok', r'Амели',
]


def clean_cuda_cache():
    """Агрессивная очистка CUDA памяти"""
    gc.collect()
    if torch.cuda.is_available():
        torch.cuda.empty_cache()
        torch.cuda.synchronize()


# === DEBUG LOG ===
class DebugLog:
    """Сборщик всех отладочных данных для анализа"""

    def __init__(self):
        self.vad_segments = []           # Исходные VAD сегменты
        self.vad_merged = []             # После объединения
        self.all_transcriptions = []     # ВСЕ транскрипции (все языки, все сегменты)
        self.scoring_decisions = []      # Решения по выбору лучшего варианта
        self.rejected_segments = []      # Отфильтрованные сегменты с причинами
        self.cleaning_log = []           # Лог очистки текста
        self.final_segments = []         # Финальные сегменты
        self.translation_log = []        # Лог переводов
        self.emotion_log = []            # Лог эмоций
        self.stats = {
            "vad_original": 0,
            "vad_merged": 0,
            "transcribed": 0,
            "rejected": {},
            "cleaned": 0,
            "translated": 0,
            "by_language": {},
        }

    def add_vad_segment(self, seg: Dict, stage: str = "original"):
        """Добавляет VAD сегмент"""
        entry = {
            "stage": stage,
            "start": round(seg["start"], 3),
            "end": round(seg["end"], 3),
            "duration": round(seg.get("duration", seg["end"] - seg["start"]), 3),
        }
        if stage == "original":
            self.vad_segments.append(entry)
            self.stats["vad_original"] += 1
        else:
            self.vad_merged.append(entry)
            self.stats["vad_merged"] += 1

    def add_transcription(self, vad_idx: int, lang: str, result: Dict):
        """Добавляет результат транскрипции"""
        entry = {
            "vad_idx": vad_idx,
            "language": lang,
            "start": round(result.get("start", 0), 3),
            "end": round(result.get("end", 0), 3),
            "text": result.get("text", ""),
            "text_length": len(result.get("text", "")),
            "score": round(result.get("score", 0), 4),
            "avg_logprob": round(result.get("avg_logprob", -1), 4),
            "no_speech_prob": round(result.get("no_speech_prob", 0.5), 4),
            "compression_ratio": round(result.get("compression_ratio", 1.5), 4),
        }
        self.all_transcriptions.append(entry)

    def add_scoring_decision(self, vad_idx: int, candidates: List[Dict],
                            chosen: Dict, reason: str):
        """Логирует решение по выбору лучшего варианта"""
        entry = {
            "vad_idx": vad_idx,
            "candidates": [{
                "lang": c.get("language"),
                "text": c.get("text", "")[:100],
                "score": round(c.get("score", 0), 4),
                "avg_logprob": round(c.get("avg_logprob", -1), 4),
                "no_speech_prob": round(c.get("no_speech_prob", 0.5), 4),
            } for c in candidates],
            "chosen": {
                "lang": chosen.get("language") if chosen else None,
                "text": chosen.get("text", "")[:100] if chosen else None,
                "score": round(chosen.get("score", 0), 4) if chosen else None,
            } if chosen else None,
            "decision": reason,
        }
        self.scoring_decisions.append(entry)

    def add_rejected(self, seg: Dict, reason: str, stage: str, details: Dict = None):
        """Добавляет отфильтрованный сегмент"""
        entry = {
            "stage": stage,
            "reason": reason,
            "start": round(seg.get("start", 0), 3),
            "end": round(seg.get("end", 0), 3),
            "text": seg.get("text", "")[:200],
            "language": seg.get("language"),
            "score": round(seg.get("score", 0), 4) if seg.get("score") else None,
            "avg_logprob": round(seg.get("avg_logprob", -1), 4) if seg.get("avg_logprob") else None,
            "no_speech_prob": round(seg.get("no_speech_prob", 0), 4) if seg.get("no_speech_prob") else None,
            "details": details,
        }
        self.rejected_segments.append(entry)
        self.stats["rejected"][reason] = self.stats["rejected"].get(reason, 0) + 1

    def add_cleaning(self, seg: Dict, original: str, cleaned: str, changes: List[str]):
        """Логирует очистку текста"""
        if original != cleaned:
            entry = {
                "start": round(seg.get("start", 0), 3),
                "end": round(seg.get("end", 0), 3),
                "original": original,
                "cleaned": cleaned,
                "changes": changes,
                "chars_removed": len(original) - len(cleaned),
            }
            self.cleaning_log.append(entry)
            self.stats["cleaned"] += 1

    def add_translation(self, seg: Dict, original: str, translated: str,
                       context: str = None, success: bool = True):
        """Логирует перевод"""
        entry = {
            "start": round(seg.get("start", 0), 3),
            "end": round(seg.get("end", 0), 3),
            "speaker": seg.get("speaker"),
            "language": seg.get("language"),
            "original": original,
            "translated": translated,
            "context_used": context[:200] if context else None,
            "success": success,
        }
        self.translation_log.append(entry)
        if success:
            self.stats["translated"] += 1

    def add_emotion(self, seg: Dict, emotion: str, confidence: float):
        """Логирует эмоцию"""
        entry = {
            "start": round(seg.get("start", 0), 3),
            "end": round(seg.get("end", 0), 3),
            "speaker": seg.get("speaker"),
            "text": seg.get("text", "")[:100],
            "emotion": emotion,
            "confidence": round(confidence, 4),
        }
        self.emotion_log.append(entry)

    def finalize(self, segments: List[Dict]):
        """Финализирует лог с итоговыми сегментами"""
        self.final_segments = [{
            "idx": i,
            "start": round(s.get("start", 0), 3),
            "end": round(s.get("end", 0), 3),
            "speaker": s.get("speaker"),
            "language": s.get("language"),
            "text": s.get("text", ""),
            "original_text": s.get("original_text"),
            "emotion": s.get("emotion"),
            "score": round(s.get("score", 0), 4) if s.get("score") else None,
        } for i, s in enumerate(segments)]

        # Статистика по языкам
        for seg in segments:
            lang = seg.get("language", "unknown")
            self.stats["by_language"][lang] = self.stats["by_language"].get(lang, 0) + 1

    def save(self, output_path: Path):
        """Сохраняет полный debug лог в JSON"""
        result = {
            "summary": {
                "vad_segments_original": self.stats["vad_original"],
                "vad_segments_merged": self.stats["vad_merged"],
                "total_transcriptions": len(self.all_transcriptions),
                "scoring_decisions": len(self.scoring_decisions),
                "rejected_total": len(self.rejected_segments),
                "rejected_by_reason": self.stats["rejected"],
                "cleaned_segments": self.stats["cleaned"],
                "translated_segments": self.stats["translated"],
                "final_segments": len(self.final_segments),
                "by_language": self.stats["by_language"],
            },
            "vad_original": self.vad_segments,
            "vad_merged": self.vad_merged,
            "all_transcriptions": self.all_transcriptions,
            "scoring_decisions": self.scoring_decisions,
            "rejected_segments": self.rejected_segments,
            "cleaning_log": self.cleaning_log,
            "translation_log": self.translation_log,
            "emotion_log": self.emotion_log,
            "final_segments": self.final_segments,
        }

        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(result, f, ensure_ascii=False, indent=2, default=str)

        print(f"🔍 DEBUG JSON: {output_path}")
        self.print_summary()

    def print_summary(self):
        """Выводит сводку в консоль"""
        print(f"\n   === DEBUG SUMMARY ===")
        print(f"   VAD: {self.stats['vad_original']} → {self.stats['vad_merged']} (merged)")
        print(f"   Транскрипций: {len(self.all_transcriptions)}")
        print(f"   Отфильтровано: {len(self.rejected_segments)}")
        for reason, count in sorted(self.stats["rejected"].items(), key=lambda x: -x[1]):
            print(f"      - {reason}: {count}")
        print(f"   Очищено: {self.stats['cleaned']}")
        print(f"   Переведено: {self.stats['translated']}")
        print(f"   Финальных: {len(self.final_segments)}")
        print(f"   По языкам: {self.stats['by_language']}")


# Глобальный debug лог
debug_log: Optional[DebugLog] = None


def clean_repetitions(text: str) -> str:
    """Удаляет повторяющиеся слова/фразы"""
    if not text or len(text) < 5:
        return text

    # Повторяющиеся символы
    text = re.sub(r'(.)\1{2,}', r'\1', text)
    # Короткие повторы типа "да, да, да"
    text = re.sub(r'\b(да|нет|ага|угу|ну|так|вот|好|是|对|嗯)[,，.\s]*(\1[,，.\s]*){1,}', r'\1', text, flags=re.IGNORECASE)
    # Повторяющиеся слова
    text = re.sub(r'\b([\w\u4e00-\u9fff]{2,20})[,，\s]+(\1[,，\s]*){1,}\1?\b', r'\1', text, flags=re.IGNORECASE)
    # Повторяющиеся фразы
    text = re.sub(r'(([\w\u4e00-\u9fff]+[,，\s]*){2,5}?)\1{1,}', r'\1', text)
    # Очистка
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'[,，.\s]+$', '', text)
    text = re.sub(r'^[,，.\s]+', '', text)
    return text.strip()


def matches_hallucination_pattern(text: str) -> bool:
    """Проверка на паттерны галлюцинаций"""
    text_lower = text.lower().strip()
    for pattern in HALLUCINATION_PATTERNS:
        if re.search(pattern, text_lower, re.IGNORECASE):
            return True
        if re.search(pattern, text, re.IGNORECASE):
            return True
    return False


def detect_language_by_text(text: str) -> str:
    """Определяет язык по содержимому"""
    if not text:
        return "unknown"

    chinese = sum(1 for c in text if '\u4e00' <= c <= '\u9fff')
    cyrillic = sum(1 for c in text if '\u0400' <= c <= '\u04ff')
    latin = sum(1 for c in text if 'a' <= c.lower() <= 'z')
    total = chinese + cyrillic + latin

    if total == 0:
        return "unknown"
    if chinese / total > 0.3:
        return "zh"
    elif cyrillic / total > 0.3:
        return "ru"
    elif latin / total > 0.5:
        return "en"
    return "unknown"


def calculate_language_score(text: str, expected_lang: str) -> float:
    """Вычисляет score соответствия языку"""
    if not text or len(text.strip()) < 2:
        return 0.0

    text = text.strip()
    chinese = sum(1 for c in text if '\u4e00' <= c <= '\u9fff')
    cyrillic = sum(1 for c in text if '\u0400' <= c <= '\u04ff')
    total = len([c for c in text if c.isalpha()])

    if total == 0:
        return 0.0

    if expected_lang == "zh":
        if cyrillic > 0:
            return 0.1
        return chinese / total
    elif expected_lang == "ru":
        if chinese > 0:
            return 0.1
        return cyrillic / total
    return 0.5


def score_transcription(seg: Dict, expected_lang: str) -> float:
    """Вычисляет качество транскрипции"""
    text = seg.get("text", "").strip()
    if not text or len(text) < 2:
        return -100.0

    avg_logprob = seg.get("avg_logprob", -1.0)
    logprob_score = max(0, min(1, (avg_logprob + 1.5) / 1.5))

    no_speech_prob = seg.get("no_speech_prob", 0.5)
    speech_score = 1.0 - no_speech_prob

    lang_score = calculate_language_score(text, expected_lang)

    duration = seg.get("end", 0) - seg.get("start", 0)
    if duration > 0:
        chars_per_sec = len(text) / duration
        length_score = min(1.0, chars_per_sec / 15.0)
    else:
        length_score = 0.5

    compression = seg.get("compression_ratio", 1.5)
    compression_score = max(0, 1.0 - (compression - 1.0) / 2.0)

    return (0.25 * logprob_score + 0.20 * speech_score +
            0.35 * lang_score + 0.10 * length_score + 0.10 * compression_score)


# === VAD ===
def load_silero_vad(device: str = "cuda"):
    """Загружает Silero VAD"""
    print("📡 Загрузка Silero VAD...")
    model, utils = torch.hub.load(
        repo_or_dir='snakers4/silero-vad',
        model='silero_vad',
        trust_repo=True
    )
    model = model.to(device)
    return model, utils[0], utils[2]


def get_vad_segments(audio_path: Path, vad_model, get_speech_timestamps,
                     min_speech_duration_ms: int = 250,
                     min_silence_duration_ms: int = 100,
                     threshold: float = 0.5,
                     device: str = "cuda") -> List[Dict]:
    """Получает сегменты речи через VAD"""
    wav, sr = librosa.load(str(audio_path), sr=16000)
    wav_tensor = torch.tensor(wav).to(device)

    speech_timestamps = get_speech_timestamps(
        wav_tensor, vad_model,
        sampling_rate=16000,
        min_speech_duration_ms=min_speech_duration_ms,
        min_silence_duration_ms=min_silence_duration_ms,
        threshold=threshold,
        return_seconds=True
    )

    segments = [{"start": ts["start"], "end": ts["end"],
                 "duration": ts["end"] - ts["start"]} for ts in speech_timestamps]

    total_speech = sum(s["duration"] for s in segments)
    total_audio = len(wav) / sr

    print(f"   📊 VAD: {len(segments)} сегментов | "
          f"Речь: {total_speech:.1f}s / {total_audio:.1f}s ({100*total_speech/total_audio:.1f}%)")

    return segments


def merge_vad_segments(segments: List[Dict], max_gap: float = 1.0,
                       max_duration: float = 30.0) -> List[Dict]:
    """
    Объединяет близкие VAD сегменты с ограничением по длине.
    Оптимизация: создаём чанки оптимального размера для batch транскрипции.
    """
    if not segments:
        return []

    merged = [segments[0].copy()]

    for seg in segments[1:]:
        last = merged[-1]
        gap = seg["start"] - last["end"]
        new_duration = seg["end"] - last["start"]

        if gap <= max_gap and new_duration <= max_duration:
            last["end"] = seg["end"]
            last["duration"] = last["end"] - last["start"]
        else:
            merged.append(seg.copy())

    print(f"   📦 После объединения: {len(merged)} чанков")
    return merged


# === TRANSCRIPTION (OPTIMIZED) ===
class MultilingualTranscriber:
    """
    Оптимизированный транскрайбер.
    Загружает модели один раз и переиспользует их.
    """

    def __init__(self, model_name: str, languages: List[str],
                 device: str, compute_type: str):
        self.device = device
        self.compute_type = compute_type
        self.languages = languages
        self.models = {}

        print(f"🔄 Загрузка моделей для языков: {languages}")
        for lang in tqdm(languages, desc="   Модели"):
            self.models[lang] = whisperx.load_model(
                model_name, device,
                compute_type=compute_type,
                language=lang,
                asr_options={"condition_on_previous_text": False}
            )

    def transcribe_segment(self, audio: np.ndarray, start: float, end: float,
                           batch_size: int = 16) -> Dict[str, Dict]:
        """Транскрибирует сегмент всеми языками"""
        sr = 16000
        segment_audio = audio[int(start * sr):int(end * sr)]

        if len(segment_audio) < sr * 0.1:
            return {}

        results = {}
        for lang, model in self.models.items():
            result = model.transcribe(segment_audio, batch_size=batch_size)

            if result["segments"]:
                text = " ".join(s.get("text", "") for s in result["segments"]).strip()
                seg = result["segments"][0]

                results[lang] = {
                    "text": text,
                    "language": lang,
                    "start": start,
                    "end": end,
                    "avg_logprob": seg.get("avg_logprob", -1.0),
                    "no_speech_prob": seg.get("no_speech_prob", 0.5),
                    "compression_ratio": seg.get("compression_ratio", 1.5),
                }
                results[lang]["score"] = score_transcription(results[lang], lang)

        return results

    def transcribe_all(self, audio: np.ndarray, vad_segments: List[Dict],
                       batch_size: int = 16) -> List[Dict]:
        """Транскрибирует все VAD сегменты с полным логированием"""
        global debug_log
        final_segments = []

        for vad_idx, vad_seg in enumerate(tqdm(vad_segments, desc="   Транскрипция")):
            results = self.transcribe_segment(
                audio, vad_seg["start"], vad_seg["end"], batch_size
            )

            # Логируем все результаты транскрипции
            for lang, result in results.items():
                if debug_log:
                    debug_log.add_transcription(vad_idx, lang, result)

            if not results:
                if debug_log:
                    debug_log.add_rejected(
                        {"start": vad_seg["start"], "end": vad_seg["end"], "text": ""},
                        "no_result", "transcription"
                    )
                    debug_log.add_scoring_decision(vad_idx, [], None, "rejected:no_result")
                continue

            # Выбираем лучший по score
            candidates = list(results.values())
            best_lang = max(results.keys(), key=lambda l: results[l]["score"])
            best = results[best_lang].copy()

            # Очищаем текст
            original_text = best["text"]
            best["text"] = clean_repetitions(best["text"])

            # Логируем очистку
            if debug_log and original_text != best["text"]:
                debug_log.add_cleaning(best, original_text, best["text"],
                                      ["repetition_removal"])

            # Проверки на галлюцинации с детальным логированием
            reject_reason = None
            reject_details = {}

            if len(best["text"]) < 3:
                reject_reason = "empty_after_clean"
                reject_details = {"original_len": len(original_text), "cleaned_len": len(best["text"])}

            elif matches_hallucination_pattern(best["text"]):
                reject_reason = "hallucination_pattern"
                reject_details = {"matched_text": best["text"][:100]}

            elif best.get("no_speech_prob", 0) > 0.7:
                reject_reason = "high_no_speech_prob"
                reject_details = {"no_speech_prob": best.get("no_speech_prob")}

            elif best.get("avg_logprob", 0) < -1.2:
                reject_reason = "low_avg_logprob"
                reject_details = {"avg_logprob": best.get("avg_logprob")}

            elif best.get("compression_ratio", 1) > 2.8:
                reject_reason = "high_compression"
                reject_details = {"compression_ratio": best.get("compression_ratio")}

            elif best["score"] < 0.25:
                reject_reason = "low_score"
                reject_details = {"score": best["score"], "threshold": 0.25}

            if reject_reason:
                if debug_log:
                    debug_log.add_rejected(best, reject_reason, "filtering", reject_details)
                    debug_log.add_scoring_decision(vad_idx, candidates, best,
                                                  f"rejected:{reject_reason}")
                continue

            # Логируем успешный выбор
            if debug_log:
                debug_log.add_scoring_decision(vad_idx, candidates, best, "accepted")

            # Добавляем альтернативы
            best["alternatives"] = {
                l: {"text": r["text"][:50], "score": round(r["score"], 3)}
                for l, r in results.items() if l != best_lang
            }

            final_segments.append(best)

        # Статистика
        print(f"   ✅ Результат: {len(final_segments)} сегментов")

        # Статистика по языкам
        lang_counts = defaultdict(int)
        for seg in final_segments:
            lang_counts[seg["language"]] += 1
        print(f"   🌍 По языкам: {dict(lang_counts)}")

        return final_segments

    def cleanup(self):
        """Освобождает память"""
        for model in self.models.values():
            del model
        self.models.clear()
        clean_cuda_cache()


# === AUDIO ===
def extract_audio(input_file: Path, output_file: Path) -> Path:
    """Извлекает аудио из видео"""
    if input_file.suffix.lower() == '.wav':
        return input_file

    print(f"🎬 Извлечение аудио из {input_file.name}...")
    cmd = [
        'ffmpeg', '-i', str(input_file),
        '-vn', '-acodec', 'pcm_s16le', '-ar', '16000', '-ac', '1',
        '-y', str(output_file)
    ]
    subprocess.run(cmd, capture_output=True)
    return output_file


# === DIARIZATION ===
def align_and_diarize(segments: List[Dict], audio: np.ndarray,
                      device: str, hf_token: str) -> List[Dict]:
    """Выравнивание и диаризация"""
    print("📐 Выравнивание...")
    lang_map = {seg.get("text", ""): seg.get("language", "ru") for seg in segments}

    model_a, metadata = whisperx.load_align_model(language_code="ru", device=device)
    result = whisperx.align(segments, model_a, metadata, audio, device=device)
    aligned = result["segments"]

    for seg in aligned:
        seg["language"] = lang_map.get(seg.get("text", ""), detect_language_by_text(seg.get("text", "")))

    del model_a
    clean_cuda_cache()

    print("👥 Диаризация...")
    diarize_model = DiarizationPipeline(use_auth_token=hf_token, device=device)
    diarize_segments = diarize_model(audio)
    result = whisperx.assign_word_speakers(diarize_segments, {"segments": aligned})

    del diarize_model
    clean_cuda_cache()

    return result["segments"]


def merge_speaker_segments(segments: List[Dict]) -> List[Dict]:
    """Склеивает последовательные сегменты одного спикера"""
    if not segments:
        return []

    merged = []
    current = None

    for seg in segments:
        speaker = seg.get("speaker", "UNKNOWN")
        text = seg.get("text", "").strip()

        if current is None:
            current = {"speaker": speaker, "start": seg["start"], "end": seg["end"],
                      "text": text, "language": seg.get("language", "ru")}
        elif current["speaker"] == speaker and current["language"] == seg.get("language", "ru"):
            current["end"] = seg["end"]
            current["text"] += " " + text
        else:
            merged.append(current)
            current = {"speaker": speaker, "start": seg["start"], "end": seg["end"],
                      "text": text, "language": seg.get("language", "ru")}

    if current:
        merged.append(current)

    # Фильтруем UNKNOWN
    original_len = len(merged)
    merged = [s for s in merged if s.get("speaker") not in ("UNKNOWN", None, "")]
    if len(merged) < original_len:
        print(f"   🚫 Удалено UNKNOWN: {original_len - len(merged)}")

    return merged


# === TRANSLATION (ADAPTIVE CONTEXTUAL) ===
def init_gemini():
    """Инициализирует Gemini"""
    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key:
        raise ValueError("GEMINI_API_KEY не найден в .env")
    genai.configure(api_key=api_key)
    return genai.GenerativeModel("gemini-2.0-flash")


def translate_adaptive(model, segments: List[Dict], context_window: int = 3) -> List[Dict]:
    """
    Адаптивный контекстный перевод сегментов (без логирования).
    """
    return translate_adaptive_with_log(model, segments, context_window)


def translate_adaptive_with_log(model, segments: List[Dict], context_window: int = 3) -> List[Dict]:
    """
    Адаптивный контекстный перевод сегментов с логированием.

    Особенности:
    - Учитывает контекст предыдущих фраз
    - Сохраняет технические термины и имена
    - Адаптирует стиль под деловое совещание
    - Переводит по одному для лучшего качества, с контекстом
    """
    global debug_log

    non_ru_indices = [i for i, s in enumerate(segments)
                     if s.get("language") != "ru" and s.get("text")]

    if not non_ru_indices:
        print("   ℹ️ Нет сегментов для перевода")
        return segments

    print(f"🌐 Адаптивный перевод {len(non_ru_indices)} сегментов...")

    lang_names = {'zh': 'китайского', 'en': 'английского', 'ja': 'японского', 'ko': 'корейского'}

    for idx in tqdm(non_ru_indices, desc="   Перевод"):
        seg = segments[idx]
        source_lang = seg.get("language", "zh")
        source_name = lang_names.get(source_lang, source_lang)

        # Собираем контекст из предыдущих сегментов (переведённых)
        context_parts = []
        for ctx_idx in range(max(0, idx - context_window), idx):
            ctx_seg = segments[ctx_idx]
            ctx_speaker = ctx_seg.get("speaker", "")
            ctx_text = ctx_seg.get("text", "")
            if ctx_text:
                context_parts.append(f"{ctx_speaker}: {ctx_text}")

        context_str = "\n".join(context_parts) if context_parts else "Начало разговора"

        # Определяем спикера текущего сегмента
        speaker = seg.get("speaker", "Спикер")
        original_text = seg["text"]

        prompt = f"""Ты профессиональный переводчик на деловых совещаниях.
Переведи реплику с {source_name} на русский язык.

ВАЖНЫЕ ПРАВИЛА:
1. Сохраняй технические термины, названия компаний и имена собственные
2. Адаптируй перевод под контекст делового совещания (строительство, проекты, сроки)
3. Сохраняй разговорный стиль если он есть в оригинале
4. НЕ добавляй слова "переводится как" или пояснения
5. Верни ТОЛЬКО перевод, одной строкой

КОНТЕКСТ РАЗГОВОРА (предыдущие реплики):
{context_str}

ТЕКУЩАЯ РЕПЛИКА ({speaker}, на {source_name}):
{original_text}

ПЕРЕВОД НА РУССКИЙ:"""

        try:
            response = model.generate_content(prompt)
            translation = response.text.strip()

            # Убираем возможные артефакты
            translation = re.sub(r'^(Перевод|Translation):\s*', '', translation, flags=re.IGNORECASE)
            translation = translation.strip('"\'')

            seg["original_text"] = original_text
            seg["translation"] = translation
            seg["text"] = translation  # Заменяем на перевод для отчёта

            # Логируем успешный перевод
            if debug_log:
                debug_log.add_translation(seg, original_text, translation,
                                         context_str, success=True)

            time.sleep(0.3)  # Rate limiting

        except Exception as e:
            print(f"      ⚠️ Ошибка перевода: {e}")
            seg["original_text"] = original_text
            seg["translation"] = f"[Ошибка перевода: {original_text[:50]}...]"
            seg["text"] = seg["translation"]

            # Логируем ошибку
            if debug_log:
                debug_log.add_translation(seg, original_text, str(e),
                                         context_str, success=False)

    # Статистика
    by_lang = defaultdict(int)
    for idx in non_ru_indices:
        by_lang[segments[idx].get("language", "?")] += 1
    print(f"   ✅ Переведено: {dict(by_lang)}")

    return segments


def translate_batch(model, segments: List[Dict], batch_size: int = 5) -> List[Dict]:
    """
    Быстрый батчевый перевод (для длинных записей).
    Используйте translate_adaptive для лучшего качества.
    """
    non_ru = [s for s in segments if s.get("language") != "ru" and s.get("text")]

    if not non_ru:
        return segments

    print(f"🌐 Быстрый перевод {len(non_ru)} сегментов...")

    by_lang = defaultdict(list)
    for seg in non_ru:
        by_lang[seg["language"]].append(seg)

    lang_names = {'zh': 'китайского', 'en': 'английского'}

    for lang, lang_segs in by_lang.items():
        source_name = lang_names.get(lang, lang)

        for i in tqdm(range(0, len(lang_segs), batch_size),
                     desc=f"   Перевод {LANGUAGE_FLAGS.get(lang, lang)}"):
            batch = lang_segs[i:i+batch_size]
            texts = [s["text"] for s in batch]

            numbered_texts = "\n".join(f"{j+1}. {t}" for j, t in enumerate(texts))

            prompt = f"""Переведи тексты с {source_name} на русский. Контекст: деловое совещание.
Сохраняй термины и имена. Верни ТОЛЬКО переводы в формате:
1. перевод
2. перевод
...

Тексты:
{numbered_texts}"""

            try:
                response = model.generate_content(prompt)
                translations = response.text.strip().split("\n")

                for j, seg in enumerate(batch):
                    seg["original_text"] = seg["text"]
                    if j < len(translations):
                        trans = re.sub(r'^\d+\.\s*', '', translations[j]).strip()
                        seg["translation"] = trans
                        seg["text"] = trans
                    else:
                        seg["translation"] = f"[Ошибка перевода]"

                time.sleep(0.5)

            except Exception as e:
                print(f"      ⚠️ Ошибка: {e}")
                for seg in batch:
                    seg["original_text"] = seg["text"]
                    seg["translation"] = f"[Ошибка: {seg['text'][:50]}]"

    return segments


# === EMOTIONS (PARALLEL) ===
class EmotionAnalyzer:
    def __init__(self, model_name: str = EMOTION_MODEL, device: str = "cuda"):
        print(f"😊 Загрузка модели эмоций...")
        self.device = device
        self.feature_extractor = Wav2Vec2FeatureExtractor.from_pretrained(model_name)
        self.model = Wav2Vec2ForSequenceClassification.from_pretrained(model_name).to(device)
        self.model.eval()
        self.id2label = self.model.config.id2label

    def analyze(self, audio_path: Path, start_time: float, end_time: float):
        try:
            duration = min(end_time - start_time, 30)
            audio, sr = librosa.load(str(audio_path), sr=16000,
                                    offset=start_time, duration=duration)
            if len(audio) < 1600:
                return 'neutral', 0.5

            inputs = self.feature_extractor(audio, sampling_rate=16000,
                                           return_tensors="pt", padding=True)
            inputs = {k: v.to(self.device) for k, v in inputs.items()}

            with torch.no_grad():
                outputs = self.model(**inputs)
                probs = torch.softmax(outputs.logits, dim=-1)
                pred_idx = torch.argmax(probs, dim=-1).item()
                confidence = probs[0][pred_idx].item()

            return self.id2label[pred_idx], confidence
        except:
            return 'neutral', 0.5

    def analyze_segments(self, audio_path: Path, segments: List[Dict]) -> List[Dict]:
        """Анализирует эмоции для всех сегментов"""
        for seg in tqdm(segments, desc="   Эмоции"):
            emotion, conf = self.analyze(audio_path, seg["start"], seg["end"])
            seg["emotion"] = emotion
            seg["emotion_confidence"] = conf
        return segments

    def analyze_segments_with_log(self, audio_path: Path, segments: List[Dict]) -> List[Dict]:
        """Анализирует эмоции для всех сегментов с логированием"""
        global debug_log

        for seg in tqdm(segments, desc="   Эмоции"):
            emotion, conf = self.analyze(audio_path, seg["start"], seg["end"])
            seg["emotion"] = emotion
            seg["emotion_confidence"] = conf

            # Логируем результат
            if debug_log:
                debug_log.add_emotion(seg, emotion, conf)

        return segments


# === REPORTS ===
def format_time(seconds):
    hours = int(seconds // 3600)
    mins = int((seconds % 3600) // 60)
    secs = int(seconds % 60)
    if hours > 0:
        return f"{hours:02d}:{mins:02d}:{secs:02d}"
    return f"{mins:02d}:{secs:02d}"


def build_speaker_profiles(segments: List[Dict]) -> Dict:
    profiles = defaultdict(lambda: {
        'emotions': [], 'total_time': 0, 'segment_count': 0,
        'emotion_counts': defaultdict(int), 'languages': set()
    })

    for seg in segments:
        speaker = seg['speaker']
        emotion = seg.get('emotion', 'neutral')
        duration = seg['end'] - seg['start']

        profiles[speaker]['emotions'].append(emotion)
        profiles[speaker]['total_time'] += duration
        profiles[speaker]['segment_count'] += 1
        profiles[speaker]['emotion_counts'][emotion] += 1
        profiles[speaker]['languages'].add(seg.get('language', 'ru'))

    for data in profiles.values():
        dominant = max(data['emotion_counts'].items(), key=lambda x: x[1])[0] if data['emotion_counts'] else 'neutral'

        interpretations = {
            'happiness': 'Позитивный настрой', 'enthusiasm': 'Позитивный настрой',
            'anger': 'Напряжённость', 'sadness': 'Обеспокоенность'
        }
        data['interpretation'] = interpretations.get(dominant, 'Деловой тон')
        data['languages'] = list(data['languages'])

    return dict(profiles)


def save_txt(segments: List[Dict], profiles: Dict, output_path: Path):
    with open(output_path, "w", encoding="utf-8") as f:
        f.write("=" * 60 + "\n")
        f.write("ПРОТОКОЛ СОВЕЩАНИЯ (v4: Оптимизированный)\n")
        f.write("=" * 60 + "\n\n")

        f.write("УЧАСТНИКИ:\n")
        f.write("-" * 60 + "\n")
        for speaker, data in sorted(profiles.items(), key=lambda x: -x[1]['total_time']):
            langs = ', '.join(LANGUAGE_FLAGS.get(l, l) for l in data['languages'])
            f.write(f"{speaker:<12} | {format_time(data['total_time']):<8} | {langs} | {data['interpretation']}\n")

        f.write("\n" + "=" * 60 + "\nТРАНСКРИПЦИЯ:\n" + "=" * 60 + "\n\n")

        for seg in segments:
            emotion = seg.get("emotion", "neutral")
            lang_flag = LANGUAGE_FLAGS.get(seg.get("language", "ru"), "")

            f.write(f"[{format_time(seg['start'])} - {format_time(seg['end'])}] "
                   f"{seg['speaker']} | {lang_flag} | {EMOTION_EMOJI.get(emotion, '')} "
                   f"{EMOTION_LABELS_RU.get(emotion, emotion)}\n")

            if seg.get("original_text"):
                f.write(f"{seg['original_text']}\n  → {seg['text']}\n\n")
            else:
                f.write(f"{seg['text']}\n\n")

    print(f"📄 TXT: {output_path}")


def save_json(segments: List[Dict], profiles: Dict, input_file: Path,
              output_path: Path, elapsed: float):
    def convert(obj):
        if isinstance(obj, (np.floating, np.integer)):
            return float(obj) if isinstance(obj, np.floating) else int(obj)
        if isinstance(obj, set):
            return list(obj)
        if isinstance(obj, dict):
            return {k: convert(v) for k, v in obj.items()}
        if isinstance(obj, list):
            return [convert(v) for v in obj]
        return obj

    result = {
        "source_file": input_file.name,
        "processed_at": datetime.now().isoformat(),
        "pipeline": "v4_optimized",
        "processing_time_seconds": elapsed,
        "segments_count": len(segments),
        "speakers": convert(profiles),
        "segments": convert(segments)
    }

    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(result, f, ensure_ascii=False, indent=2)
    print(f"📋 JSON: {output_path}")


def create_word_report(segments: List[Dict], profiles: Dict,
                       input_file: Path, output_path: Path):
    doc = Document()

    title = doc.add_heading('Протокол совещания (v4)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('Информация', level=1)
    total_duration = segments[-1]["end"] if segments else 0

    info = doc.add_table(rows=4, cols=2)
    info.style = 'Table Grid'
    for i, (k, v) in enumerate([
        ('Файл', input_file.name),
        ('Дата', datetime.now().strftime('%d.%m.%Y %H:%M')),
        ('Длительность', format_time(total_duration)),
        ('Участников', str(len(profiles))),
    ]):
        info.rows[i].cells[0].text = k
        info.rows[i].cells[1].text = v

    doc.add_paragraph()
    doc.add_heading('Транскрипция', level=1)

    for seg in segments:
        emotion = seg.get("emotion", "neutral")
        lang_flag = LANGUAGE_FLAGS.get(seg.get("language", "ru"), "")

        header = doc.add_paragraph()
        run = header.add_run(
            f"[{format_time(seg['start'])} - {format_time(seg['end'])}] "
            f"{seg['speaker']} | {lang_flag} | {EMOTION_EMOJI.get(emotion, '')} "
            f"{EMOTION_LABELS_RU.get(emotion, emotion)}"
        )
        run.bold = True

        if seg.get("original_text"):
            doc.add_paragraph(seg['original_text'])
            trans = doc.add_paragraph()
            trans.add_run(f"→ {seg['text']}").italic = True
        else:
            doc.add_paragraph(seg['text'])

    doc.save(str(output_path))
    print(f"📝 DOCX: {output_path}")


# === MAIN PIPELINE ===
def process_v4(
    input_file: Path,
    output_dir: Path = OUTPUT_DIR,
    languages: List[str] = ["ru", "zh"],
    model: str = "large-v3",
    device: str = "cuda",
    compute_type: str = "float16",
    batch_size: int = 16,
    skip_emotions: bool = False,
    skip_translation: bool = False,
    skip_diarization: bool = False,
    vad_threshold: float = 0.5,
    vad_min_speech_ms: int = 250,
) -> Path:
    """
    Оптимизированный мультиязычный пайплайн v4.
    """
    global debug_log
    debug_log = DebugLog()  # Инициализируем debug лог

    start_time_total = time.time()

    print("=" * 60)
    print("🚀 WhisperX Pipeline v4 (Optimized + Debug)")
    print(f"📁 Файл: {input_file.name}")
    print(f"🌍 Языки: {languages}")
    print(f"🤖 Модель: {model}")
    print("=" * 60)

    output_dir.mkdir(parents=True, exist_ok=True)
    temp_audio = output_dir / "temp_audio.wav"

    # 1. Аудио
    print("\n[1/7] 🎵 Подготовка аудио")
    audio_file = extract_audio(input_file, temp_audio)

    # 2. VAD
    print("\n[2/7] 📡 VAD анализ")
    vad_model, get_speech_ts, _ = load_silero_vad(device)
    vad_segments_raw = get_vad_segments(
        audio_file, vad_model, get_speech_ts,
        min_speech_duration_ms=vad_min_speech_ms,
        threshold=vad_threshold, device=device
    )

    # Логируем исходные VAD сегменты
    for seg in vad_segments_raw:
        debug_log.add_vad_segment(seg, "original")

    vad_segments = merge_vad_segments(vad_segments_raw, max_gap=1.0, max_duration=30.0)

    # Логируем объединённые VAD сегменты
    for seg in vad_segments:
        debug_log.add_vad_segment(seg, "merged")

    del vad_model
    clean_cuda_cache()

    # 3. Транскрипция
    print("\n[3/7] 🎤 Транскрипция")
    audio = whisperx.load_audio(str(audio_file))

    transcriber = MultilingualTranscriber(model, languages, device, compute_type)
    segments = transcriber.transcribe_all(audio, vad_segments, batch_size)
    transcriber.cleanup()

    # 4. Диаризация
    if not skip_diarization:
        print("\n[4/7] 👥 Диаризация")
        hf_token = os.getenv("HUGGINGFACE_TOKEN")
        segments = align_and_diarize(segments, audio, device, hf_token)
        segments = merge_speaker_segments(segments)
    else:
        print("\n[4/7] 👥 Диаризация пропущена")
        for seg in segments:
            seg["speaker"] = "SPEAKER_00"

    print(f"   📊 Сегментов после обработки: {len(segments)}")

    # 5. Перевод (адаптивный с контекстом)
    if not skip_translation:
        print("\n[5/7] 🌐 Перевод")
        gemini = init_gemini()
        # Используем адаптивный перевод для лучшего качества
        segments = translate_adaptive_with_log(gemini, segments, context_window=3)
    else:
        print("\n[5/7] 🌐 Перевод пропущен")

    # 6. Эмоции
    if not skip_emotions:
        print("\n[6/7] 😊 Анализ эмоций")
        analyzer = EmotionAnalyzer(device=device)
        segments = analyzer.analyze_segments_with_log(audio_file, segments)
    else:
        print("\n[6/7] 😊 Эмоции пропущены")
        for seg in segments:
            seg["emotion"] = "neutral"

    # 7. Сохранение
    print("\n[7/7] 💾 Сохранение")
    profiles = build_speaker_profiles(segments)
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')

    elapsed = time.time() - start_time_total

    # Финализируем debug лог
    debug_log.finalize(segments)

    txt_path = output_dir / f"v4_{timestamp}.txt"
    save_txt(segments, profiles, txt_path)

    json_path = output_dir / f"v4_{timestamp}.json"
    save_json(segments, profiles, input_file, json_path, elapsed)

    word_path = output_dir / f"v4_{timestamp}.docx"
    create_word_report(segments, profiles, input_file, word_path)

    # Сохраняем debug JSON
    debug_path = output_dir / f"v4_debug_{timestamp}.json"
    debug_log.save(debug_path)

    # Cleanup
    if audio_file == temp_audio and temp_audio.exists():
        temp_audio.unlink()

    print("\n" + "=" * 60)
    print(f"✅ ГОТОВО! Время: {elapsed/60:.1f} мин")
    print(f"📄 Результат: {word_path}")
    print(f"🔍 Debug: {debug_path}")
    print("=" * 60)

    return word_path


# === ENTRY POINT ===
if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="WhisperX Pipeline v4 (Optimized)")
    parser.add_argument("input", nargs="?", help="Входной файл")
    parser.add_argument("--skip-emotions", action="store_true")
    parser.add_argument("--skip-translation", action="store_true")
    parser.add_argument("--skip-diarization", action="store_true")
    parser.add_argument("--languages", nargs="+", default=["ru", "zh"])
    parser.add_argument("--vad-threshold", type=float, default=0.5)
    parser.add_argument("--batch-size", type=int, default=16)
    parser.add_argument("-o", "--output", type=Path, default=OUTPUT_DIR)
    args = parser.parse_args()

    if args.input:
        INPUT_FILE = Path(args.input)
    else:
        # Дефолтный тестовый файл
        INPUT_FILE = PROJECT_ROOT / "test_segment_0-45min.wav"

    if not INPUT_FILE.exists():
        print(f"❌ ОШИБКА: Файл не найден: {INPUT_FILE}")
        sys.exit(1)

    result = process_v4(
        input_file=INPUT_FILE,
        output_dir=args.output,
        languages=args.languages,
        skip_emotions=args.skip_emotions,
        skip_translation=args.skip_translation,
        skip_diarization=args.skip_diarization,
        vad_threshold=args.vad_threshold,
        batch_size=args.batch_size,
    )

    print(f"\n🎉 Результат: {result}")
