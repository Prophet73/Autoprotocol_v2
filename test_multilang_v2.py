#!/usr/bin/env python3
"""
Тест мультиязычной транскрипции (RU + ZH) с фильтрацией галлюцинаций и переводом.
"""

import torch
import os
import sys
import re
import json
import time
from pathlib import Path
from datetime import datetime
from collections import defaultdict
from typing import Optional, List, Dict, Any

# Patch torch.load BEFORE any imports (PyTorch 2.8+ fix)
_original_torch_load = torch.load
def _patched_torch_load(*args, **kwargs):
    kwargs['weights_only'] = False
    return _original_torch_load(*args, **kwargs)
torch.load = _patched_torch_load

import lightning_fabric.utilities.cloud_io as cloud_io
cloud_io.torch.load = _patched_torch_load

import whisperx
from whisperx.diarize import DiarizationPipeline
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from transformers import Wav2Vec2ForSequenceClassification, Wav2Vec2FeatureExtractor
import librosa
import subprocess
import numpy as np

# Google Gemini
import google.generativeai as genai
from dotenv import load_dotenv

# === LOAD ENV ===
PROJECT_ROOT = Path(__file__).parent.resolve()
load_dotenv(PROJECT_ROOT / ".env")

# === КОНСТАНТЫ ===
OUTPUT_DIR = PROJECT_ROOT / "output"
EMOTION_MODEL = "Aniemore/wav2vec2-xlsr-53-russian-emotion-recognition"

EMOTION_LABELS_RU = {
    'anger': 'Гнев', 'disgust': 'Отвращение', 'enthusiasm': 'Энтузиазм',
    'fear': 'Страх', 'happiness': 'Радость', 'neutral': 'Нейтрально', 'sadness': 'Грусть'
}
EMOTION_EMOJI = {
    'anger': '😠', 'disgust': '🤢', 'enthusiasm': '🤩',
    'fear': '😨', 'happiness': '😊', 'neutral': '😐', 'sadness': '😔'
}

LANGUAGE_FLAGS = {
    'ru': '🇷🇺', 'zh': '🇨🇳', 'en': '🇬🇧', 'unknown': '❓'
}


def detect_language_by_text(text: str) -> str:
    """Определяет язык по содержимому текста"""
    if not text:
        return "ru"

    # Считаем символы разных типов
    chinese_chars = sum(1 for c in text if '\u4e00' <= c <= '\u9fff')
    cyrillic_chars = sum(1 for c in text if '\u0400' <= c <= '\u04ff')
    latin_chars = sum(1 for c in text if 'a' <= c.lower() <= 'z')

    total = chinese_chars + cyrillic_chars + latin_chars
    if total == 0:
        return "ru"

    # Если больше 30% китайских символов — это китайский
    if chinese_chars / total > 0.3:
        return "zh"
    # Если больше 30% латиницы — это английский
    elif latin_chars / total > 0.5 and cyrillic_chars / total < 0.2:
        return "en"
    else:
        return "ru"


def is_valid_for_language(text: str, expected_lang: str) -> bool:
    """
    Проверяет что текст действительно на ожидаемом языке.
    Отсекает мусор когда Whisper пытается натянуть не тот язык.
    """
    if not text or len(text.strip()) < 2:
        return False

    text = text.strip()

    # Считаем символы
    chinese_chars = sum(1 for c in text if '\u4e00' <= c <= '\u9fff')
    cyrillic_chars = sum(1 for c in text if '\u0400' <= c <= '\u04ff')

    total_meaningful = chinese_chars + cyrillic_chars
    if total_meaningful == 0:
        return False

    if expected_lang == "zh":
        # Для китайского: минимум 30% иероглифов
        return chinese_chars / total_meaningful > 0.3 if total_meaningful > 0 else False
    elif expected_lang == "ru":
        # Для русского: минимум 50% кириллицы
        return cyrillic_chars / total_meaningful > 0.5 if total_meaningful > 0 else False
    else:
        return True


def is_valid_for_language_strict(text: str, expected_lang: str) -> bool:
    """
    СТРОГАЯ проверка соответствия текста языку.

    Ключевое отличие от is_valid_for_language:
    - Если ожидаем RU, но в тексте ЕСТЬ китайские иероглифы — отклоняем
    - Если ожидаем ZH, но в тексте ЕСТЬ кириллица — отклоняем

    Это предотвращает ситуацию когда Whisper "переводит" китайский на русский
    или генерирует смешанный текст.
    """
    if not text or len(text.strip()) < 2:
        return False

    text = text.strip()

    # Считаем символы разных типов
    chinese_chars = sum(1 for c in text if '\u4e00' <= c <= '\u9fff')
    cyrillic_chars = sum(1 for c in text if '\u0400' <= c <= '\u04ff')
    latin_chars = sum(1 for c in text if 'a' <= c.lower() <= 'z')

    total_chars = len([c for c in text if c.isalpha()])
    if total_chars == 0:
        return False

    if expected_lang == "zh":
        # Для китайского:
        # - Должно быть минимум 60% иероглифов
        # - НЕ должно быть кириллицы (если Whisper "перевёл" — отклоняем)
        if cyrillic_chars > 0:
            return False  # Кириллица в китайском тексте = галлюцинация/перевод
        return chinese_chars / total_chars > 0.6 if total_chars > 0 else False

    elif expected_lang == "ru":
        # Для русского:
        # - Должно быть минимум 60% кириллицы
        # - НЕ должно быть китайских иероглифов
        if chinese_chars > 0:
            return False  # Иероглифы в русском тексте = что-то пошло не так
        return cyrillic_chars / total_chars > 0.6 if total_chars > 0 else False

    else:
        return True


def smart_deduplicate(segments: List[Dict], overlap_threshold: float = 0.3) -> List[Dict]:
    """
    Умная дедупликация сегментов из разных языковых проходов.

    Логика:
    1. Сортируем все сегменты по времени начала
    2. Для каждого временного слота выбираем ЛУЧШИЙ сегмент:
       - Тот, где язык текста совпадает с заявленным языком
       - При равенстве — более длинный текст
    3. Сегменты без пересечения — добавляем все

    Args:
        segments: Список сегментов из всех языковых проходов
        overlap_threshold: Минимальное пересечение для считания дубликатом (0.3 = 30%)

    Returns:
        Дедуплицированный список сегментов
    """
    if not segments:
        return []

    # Сортируем по времени начала
    sorted_segs = sorted(segments, key=lambda x: x.get("start", 0))

    result = []

    for seg in sorted_segs:
        start = seg.get("start", 0)
        end = seg.get("end", 0)
        duration = end - start
        text = seg.get("text", "").strip()
        lang = seg.get("language", "ru")

        if duration <= 0 or not text:
            continue

        # Проверяем соответствие языка текста заявленному языку
        detected_lang = detect_language_by_text(text)
        seg_matches_lang = (detected_lang == lang)

        # Ищем пересечения с уже добавленными сегментами
        best_overlap_idx = -1
        best_overlap_ratio = 0

        for i, existing in enumerate(result):
            ex_start = existing.get("start", 0)
            ex_end = existing.get("end", 0)

            # Вычисляем пересечение
            overlap_start = max(start, ex_start)
            overlap_end = min(end, ex_end)
            overlap = max(0, overlap_end - overlap_start)

            # Отношение пересечения к меньшему сегменту
            min_duration = min(duration, ex_end - ex_start)
            if min_duration > 0:
                overlap_ratio = overlap / min_duration
                if overlap_ratio > overlap_threshold and overlap_ratio > best_overlap_ratio:
                    best_overlap_ratio = overlap_ratio
                    best_overlap_idx = i

        if best_overlap_idx >= 0:
            # Есть пересечение — решаем какой сегмент лучше
            existing = result[best_overlap_idx]
            ex_text = existing.get("text", "").strip()
            ex_lang = existing.get("language", "ru")
            ex_detected = detect_language_by_text(ex_text)
            ex_matches_lang = (ex_detected == ex_lang)

            # ПРИОРИТЕТ: Китайский текст (иероглифы) важнее русского
            # потому что русский на месте китайского — это обычно бессмыслица
            seg_has_chinese = any('\u4e00' <= c <= '\u9fff' for c in text)
            ex_has_chinese = any('\u4e00' <= c <= '\u9fff' for c in ex_text)

            should_replace = False

            # Если новый сегмент — настоящий китайский, а старый — русский
            if seg_has_chinese and lang == "zh" and not ex_has_chinese:
                should_replace = True
            # Если старый — китайский, не заменяем на русский
            elif ex_has_chinese and existing.get("language") == "zh" and not seg_has_chinese:
                should_replace = False
            # Стандартная логика для остальных случаев
            elif seg_matches_lang and not ex_matches_lang:
                should_replace = True
            elif seg_matches_lang == ex_matches_lang:
                if len(text) > len(ex_text) * 1.2:
                    should_replace = True

            if should_replace:
                result[best_overlap_idx] = seg
        else:
            # Нет пересечения — добавляем
            result.append(seg)

    # Финальная сортировка
    result = sorted(result, key=lambda x: x.get("start", 0))

    return result


def merge_and_deduplicate(segments: List[Dict], overlap_threshold: float = 0.5) -> List[Dict]:
    """
    Мержит сегменты из разных языков и убирает дубликаты по времени.
    Если два сегмента пересекаются более чем на overlap_threshold — оставляем один.
    """
    if not segments:
        return []

    # Сортируем по времени начала
    sorted_segs = sorted(segments, key=lambda x: x.get("start", 0))

    result = []
    for seg in sorted_segs:
        start = seg.get("start", 0)
        end = seg.get("end", 0)
        duration = end - start

        # Проверяем пересечение с последним добавленным
        if result:
            last = result[-1]
            last_start = last.get("start", 0)
            last_end = last.get("end", 0)

            # Вычисляем пересечение
            overlap_start = max(start, last_start)
            overlap_end = min(end, last_end)
            overlap = max(0, overlap_end - overlap_start)

            # Если пересечение > 50% от меньшего сегмента — это дубликат
            min_duration = min(duration, last_end - last_start)
            if min_duration > 0 and overlap / min_duration > overlap_threshold:
                # Выбираем тот, у которого язык "правильнее" по тексту
                seg_detected = detect_language_by_text(seg.get("text", ""))
                last_detected = detect_language_by_text(last.get("text", ""))

                seg_matches = seg_detected == seg.get("language")
                last_matches = last_detected == last.get("language")

                if seg_matches and not last_matches:
                    # Новый лучше — заменяем
                    result[-1] = seg
                # Иначе оставляем старый
                continue

        result.append(seg)

    return result

# === ФИЛЬТРАЦИЯ ГАЛЛЮЦИНАЦИЙ ===
HALLUCINATION_PATTERNS = [
    # Русские галлюцинации
    r'продолжение следует',
    r'субтитры (сделал|подогнал|создал)',
    r'спасибо за просмотр',
    r'подписывайтесь на канал',
    r'ставьте лайк',
    r'DimaTorzok',
    r'Симон',
    # Китайские галлюцинации
    r'^谢谢大家\.?$',  # "Спасибо всем" - классика
    r'^谢谢\.?$',      # "Спасибо"
    r'^对不起.*谢谢',  # "Прошу прощения... спасибо"
    r'^請訂閱',        # "Подпишитесь"
    r'^感謝收看',      # "Спасибо за просмотр"
    r'^謝謝觀看',      # "Спасибо за просмотр"
    # Английские галлюцинации
    r'^thank you',
    r'^thanks for watching',
    r'^please subscribe',
    # Повторы
    r'(.)\1{5,}',  # Повторяющиеся символы (ээээээ, ааааа)
]

def has_repeated_phrases(text: str, min_length: int = 5) -> bool:
    """
    Проверяет есть ли повторяющиеся фразы/символы в тексте.
    Типичные галлюцинации: "個個個個", "они, они, они", "A-A-A-A-A"
    """
    text = text.strip()
    if len(text) < 10:
        return False

    # 1. Повторяющиеся символы подряд (個個個個, AAAA)
    for i in range(len(text) - 3):
        char = text[i]
        if char.isalnum() and text[i:i+4] == char * 4:
            return True

    # 2. Повторяющиеся короткие паттерны (A-A-A, 10-10-10)
    import re
    # Паттерн: что-то повторяется 4+ раз
    repeat_pattern = r'(.{2,10}?)\1{3,}'
    if re.search(repeat_pattern, text):
        return True

    # 3. Повторяющиеся слова/фразы
    words = text.split()
    if len(words) >= 6:
        # Ищем повторяющиеся биграммы/триграммы
        for n in [2, 3]:
            ngrams = [' '.join(words[i:i+n]) for i in range(len(words) - n + 1)]
            for ngram in set(ngrams):
                if ngrams.count(ngram) >= 3:  # Фраза повторяется 3+ раз
                    return True

    return False


def is_hallucination_smart(seg: Dict, no_speech_threshold: float = 0.6,
                           logprob_threshold: float = -1.0) -> bool:
    """
    Умная проверка на галлюцинации используя метаданные Whisper.

    Args:
        seg: Сегмент с метаданными от Whisper
        no_speech_threshold: Порог no_speech_prob (выше = скорее галлюцинация)
        logprob_threshold: Порог avg_logprob (ниже = скорее галлюцинация)

    Returns:
        True если сегмент похож на галлюцинацию
    """
    text = seg.get("text", "").strip()
    text_lower = text.lower()
    start = seg.get("start", 0)
    end = seg.get("end", 0)
    duration = end - start

    # 1. Пустой или очень короткий текст
    if len(text_lower) < 3:
        return True

    # 2. Проверка по паттернам галлюцинаций
    for pattern in HALLUCINATION_PATTERNS:
        if re.search(pattern, text_lower, re.IGNORECASE):
            return True
        # Также проверяем оригинальный текст (для китайского)
        if re.search(pattern, text, re.IGNORECASE):
            return True

    # 3. Проверка на повторяющиеся фразы
    if has_repeated_phrases(text):
        return True

    # 4. Whisper metadata: no_speech_prob
    # Если модель считает что там нет речи — это галлюцинация
    no_speech_prob = seg.get("no_speech_prob", 0.0)
    if no_speech_prob > no_speech_threshold:
        return True

    # 5. Whisper metadata: avg_logprob (уверенность модели)
    # Очень низкая уверенность = подозрительно
    avg_logprob = seg.get("avg_logprob", 0.0)
    if avg_logprob < logprob_threshold:
        return True

    # 6. Whisper metadata: compression_ratio
    # Высокий compression ratio = повторяющийся/странный текст
    compression_ratio = seg.get("compression_ratio", 1.0)
    if compression_ratio > 2.8:  # Дефолт Whisper = 2.4, берём чуть выше
        return True

    # 7. Проверка плотности текста (для длинных сегментов)
    # 50 сек сегмент с 6 иероглифами — явно галлюцинация
    if duration > 15:  # Для сегментов > 15 сек
        # Считаем "значимые" символы (буквы + иероглифы)
        meaningful_chars = len([c for c in text if c.isalnum() or '\u4e00' <= c <= '\u9fff'])
        chars_per_sec = meaningful_chars / duration if duration > 0 else 0

        # Минимум 1 значимый символ в секунду
        if chars_per_sec < 1:
            return True

    return False


def is_hallucination(text: str, duration: float = 0) -> bool:
    """
    Простая проверка на галлюцинации (для обратной совместимости).
    Используется когда нет метаданных Whisper.
    """
    text_lower = text.lower().strip()

    # Пустой или очень короткий текст
    if len(text_lower) < 3:
        return True

    # Проверка по паттернам
    for pattern in HALLUCINATION_PATTERNS:
        if re.search(pattern, text_lower, re.IGNORECASE):
            return True

    # Проверка на повторяющиеся фразы
    if has_repeated_phrases(text):
        return True

    return False


def filter_segments(segments: List[Dict], use_smart_filter: bool = True) -> List[Dict]:
    """
    Фильтрует сегменты: убирает UNKNOWN и галлюцинации.

    Args:
        segments: Список сегментов
        use_smart_filter: Использовать умную фильтрацию по метаданным Whisper
    """
    filtered = []
    removed_unknown = 0
    removed_hallucination = 0
    removed_no_speech = 0
    removed_low_confidence = 0

    for seg in segments:
        speaker = seg.get("speaker", "UNKNOWN")
        text = seg.get("text", "").strip()

        # Убираем UNKNOWN спикеров
        if speaker == "UNKNOWN" or not speaker:
            removed_unknown += 1
            continue

        if use_smart_filter:
            # Умная фильтрация по метаданным Whisper
            no_speech_prob = seg.get("no_speech_prob", 0.0)
            avg_logprob = seg.get("avg_logprob", 0.0)

            if is_hallucination_smart(seg):
                if no_speech_prob > 0.6:
                    removed_no_speech += 1
                elif avg_logprob < -1.0:
                    removed_low_confidence += 1
                else:
                    removed_hallucination += 1
                continue
        else:
            # Простая фильтрация
            if is_hallucination(text):
                removed_hallucination += 1
                continue

        filtered.append(seg)

    total_removed = removed_unknown + removed_hallucination + removed_no_speech + removed_low_confidence
    print(f"    Фильтрация: удалено {total_removed} сегментов:")
    print(f"      - UNKNOWN спикеры: {removed_unknown}")
    print(f"      - Галлюцинации (паттерны): {removed_hallucination}")
    print(f"      - no_speech_prob > 0.6: {removed_no_speech}")
    print(f"      - avg_logprob < -1.0: {removed_low_confidence}")
    return filtered


# === GEMINI TRANSLATION ===
def init_gemini():
    """Инициализирует Gemini API"""
    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key:
        raise ValueError("GEMINI_API_KEY не найден в .env")
    genai.configure(api_key=api_key)
    return genai.GenerativeModel("gemini-2.5-flash")


def translate_text(model, text: str, source_lang: str, target_lang: str = "ru") -> str:
    """Переводит текст через Gemini"""
    if source_lang == target_lang:
        return text

    lang_names = {'zh': 'китайского', 'en': 'английского', 'ru': 'русский'}
    source_name = lang_names.get(source_lang, source_lang)

    prompt = f"""Переведи следующий текст с {source_name} на русский язык.
Сохрани смысл и стиль оригинала. Верни ТОЛЬКО перевод, без пояснений.

Текст: {text}"""

    try:
        response = model.generate_content(prompt)
        return response.text.strip()
    except Exception as e:
        print(f"    Ошибка перевода: {e}")
        return f"[Ошибка перевода: {text}]"


def translate_segments(segments: List[Dict], model, target_lang: str = "ru") -> List[Dict]:
    """Переводит все не-русские сегменты"""
    translated = []
    translate_count = 0

    for seg in segments:
        lang = seg.get("language", "ru")

        if lang != target_lang and seg.get("text"):
            # Нужен перевод
            original_text = seg["text"]
            translation = translate_text(model, original_text, lang, target_lang)

            seg["original_text"] = original_text
            seg["translation"] = translation
            seg["text"] = translation  # Основной текст теперь перевод
            translate_count += 1

            # Rate limiting
            time.sleep(1)  # 1 секунда между запросами

        translated.append(seg)

    print(f"    Переведено сегментов: {translate_count}")
    return translated


# === AUDIO & TRANSCRIPTION ===
def extract_audio(input_file: Path, output_file: Path) -> Path:
    """Извлекает аудио из видео или копирует WAV если уже в нужном формате"""

    # Если входной файл уже WAV — используем его напрямую
    if input_file.suffix.lower() == '.wav':
        print(f"Входной файл уже WAV: {input_file.name}")
        return input_file

    print(f"Извлечение аудио из {input_file.name}...")
    cmd = [
        'ffmpeg', '-i', str(input_file),
        '-vn', '-acodec', 'pcm_s16le', '-ar', '16000', '-ac', '1',
        '-y', str(output_file)
    ]
    subprocess.run(cmd, capture_output=True)
    print(f"Аудио сохранено: {output_file}")
    return output_file


def quick_language_scan(audio, device: str = "cuda", compute_type: str = "int8",
                        batch_size: int = 16, languages: List[str] = ["ru", "zh"]) -> List[Dict]:
    """
    Быстрый скан аудио для определения языков по таймкодам.

    Делает отдельный проход tiny моделью для каждого языка,
    затем выбирает лучший результат для каждого интервала.

    Args:
        audio: Загруженное аудио
        device: Устройство (cuda/cpu)
        compute_type: int8 для скорости
        batch_size: Размер батча
        languages: Список языков для сканирования

    Returns:
        Список сегментов с определёнными языками и таймкодами
    """
    print("\n  === БЫСТРЫЙ СКАН ЯЗЫКОВ (tiny model, 2 прохода) ===")

    all_segments = []

    for lang in languages:
        print(f"  Проход tiny ({lang})...")
        model = whisperx.load_model(
            "tiny", device,
            compute_type=compute_type,
            language=lang,
            asr_options={
                "condition_on_previous_text": False,
            }
        )

        result = model.transcribe(audio, batch_size=batch_size)

        for seg in result["segments"]:
            text = seg.get("text", "").strip()
            # Проверяем соответствие языка
            if is_valid_for_language_strict(text, lang):
                all_segments.append({
                    "start": seg.get("start", 0),
                    "end": seg.get("end", 0),
                    "language": lang,
                    "text": text,
                    "no_speech_prob": seg.get("no_speech_prob", 0.0),
                    "avg_logprob": seg.get("avg_logprob", 0.0),
                    "compression_ratio": seg.get("compression_ratio", 1.0),
                })

        print(f"    {lang}: {len([s for s in all_segments if s['language'] == lang])} валидных сегментов")

        del model
        torch.cuda.empty_cache()

    # Дедупликация — выбираем лучший сегмент для каждого интервала
    language_map = smart_deduplicate(all_segments, overlap_threshold=0.3)

    # Статистика
    ru_count = sum(1 for s in language_map if s["language"] == "ru")
    zh_count = sum(1 for s in language_map if s["language"] == "zh")

    print(f"  Итого в карте: {len(language_map)} сегментов")
    print(f"    - Русский (RU): {ru_count}")
    print(f"    - Китайский (ZH): {zh_count}")

    return language_map


def get_language_for_time(language_map: List[Dict], time: float) -> str:
    """
    Определяет язык для заданного момента времени по карте языков.

    Args:
        language_map: Карта языков от quick_language_scan
        time: Момент времени в секундах

    Returns:
        Язык ('ru', 'zh', 'en') или 'ru' по умолчанию
    """
    for seg in language_map:
        if seg["start"] <= time <= seg["end"]:
            return seg["language"]
    return "ru"  # По умолчанию русский


def transcribe_language(audio, model_name: str, language: str, device: str,
                        compute_type: str, batch_size: int,
                        language_map: List[Dict] = None) -> List[Dict]:
    """
    Транскрибирует аудио для одного языка.

    ВАЖНО: condition_on_previous_text=False предотвращает "заражение" контекстом,
    когда модель слышит один язык, но пытается продолжить на другом.

    Args:
        audio: Загруженное аудио
        model_name: Название модели (large-v3)
        language: Язык для транскрипции
        device: Устройство
        compute_type: Тип вычислений
        batch_size: Размер батча
        language_map: Карта языков от quick_language_scan (опционально)
    """
    print(f"\n  Загрузка модели для языка: {language}")

    # ASR options для предотвращения галлюцинаций при смешанных языках
    asr_options = {
        "condition_on_previous_text": False,  # КЛЮЧЕВОЕ: не использовать предыдущий текст как контекст
        "no_speech_threshold": 0.4,           # Чуть ниже дефолта (0.6) для лучшего определения тишины
        "compression_ratio_threshold": 2.4,   # Стандартное значение
        "log_prob_threshold": -1.0,           # Стандартное значение
    }

    model = whisperx.load_model(
        model_name, device,
        compute_type=compute_type,
        language=language,
        asr_options=asr_options
    )

    print(f"  Транскрипция ({language}) с condition_on_previous_text=False...")
    result = model.transcribe(audio, batch_size=batch_size)

    # Помечаем язык и фильтруем сегменты
    valid_segments = []
    rejected_lang = 0
    rejected_map = 0

    for seg in result["segments"]:
        text = seg.get("text", "").strip()
        start = seg.get("start", 0)
        end = seg.get("end", 0)
        mid_time = (start + end) / 2

        # Сохраняем метаданные Whisper для умной фильтрации
        seg["language"] = language
        seg["no_speech_prob"] = seg.get("no_speech_prob", 0.0)
        seg["avg_logprob"] = seg.get("avg_logprob", 0.0)
        seg["compression_ratio"] = seg.get("compression_ratio", 1.0)

        # Если есть карта языков — проверяем соответствие
        if language_map:
            expected_lang = get_language_for_time(language_map, mid_time)
            # ВАЖНО: EN часто = ошибка tiny (путает с ZH)
            # Поэтому EN пропускаем в оба прохода, пусть валидация решает
            if expected_lang not in (language, "en", "unknown"):
                # Этот интервал принадлежит другому языку — пропускаем
                rejected_map += 1
                continue

        # Строгая валидация: текст должен соответствовать заявленному языку
        if is_valid_for_language_strict(text, language):
            valid_segments.append(seg)
        else:
            rejected_lang += 1

    print(f"  Найдено сегментов ({language}): {len(result['segments'])}")
    if language_map:
        print(f"  Отфильтровано по карте языков: {rejected_map}")
    print(f"  После валидации языка: {len(valid_segments)} (отброшено: {rejected_lang})")

    del model
    torch.cuda.empty_cache()

    return valid_segments


def merge_multilang_segments(all_segments: List[Dict]) -> List[Dict]:
    """Сортирует и объединяет сегменты из разных языков по времени"""
    # Сортируем по времени начала
    sorted_segs = sorted(all_segments, key=lambda x: x.get("start", 0))
    return sorted_segs


def align_segments(segments: List[Dict], audio, language: str, device: str) -> List[Dict]:
    """Выравнивание сегментов"""
    print(f"  Выравнивание ({language})...")
    model_a, metadata = whisperx.load_align_model(language_code=language, device=device)
    result = whisperx.align(segments, model_a, metadata, audio, device=device)

    del model_a
    torch.cuda.empty_cache()

    return result["segments"]


def diarize_and_assign(segments: List[Dict], audio, device: str, hf_token: str) -> List[Dict]:
    """Диаризация и присвоение спикеров"""
    print("\nДиаризация спикеров...")
    diarize_model = DiarizationPipeline(use_auth_token=hf_token, device=device)
    diarize_segments = diarize_model(audio)

    # Присваиваем спикеров
    result = whisperx.assign_word_speakers(diarize_segments, {"segments": segments})

    del diarize_model
    torch.cuda.empty_cache()

    return result["segments"]


def merge_speaker_segments(segments: List[Dict]) -> List[Dict]:
    """Склеивает последовательные сегменты одного спикера"""
    if not segments:
        return []

    merged = []
    current = None

    for seg in segments:
        speaker = seg.get("speaker", "UNKNOWN")
        text = seg.get("text", "").strip()
        start = seg.get("start", 0)
        end = seg.get("end", 0)
        language = seg.get("language", "ru")

        if current is None:
            current = {
                "speaker": speaker, "start": start, "end": end,
                "text": text, "language": language
            }
        elif current["speaker"] == speaker and current["language"] == language:
            current["end"] = end
            current["text"] += " " + text
        else:
            merged.append(current)
            current = {
                "speaker": speaker, "start": start, "end": end,
                "text": text, "language": language
            }

    if current:
        merged.append(current)

    return merged


# === EMOTIONS ===
class EmotionAnalyzer:
    def __init__(self, model_name: str = EMOTION_MODEL, device: str = "cuda"):
        print(f"Загрузка модели эмоций: {model_name}")
        self.device = device
        self.feature_extractor = Wav2Vec2FeatureExtractor.from_pretrained(model_name)
        self.model = Wav2Vec2ForSequenceClassification.from_pretrained(model_name).to(device)
        self.model.eval()
        self.id2label = self.model.config.id2label

    def analyze(self, audio_path: Path, start_time: float, end_time: float):
        try:
            duration = end_time - start_time
            if duration > 30:
                samples = []
                for offset in [0, duration/2 - 5, duration - 10]:
                    audio, sr = librosa.load(str(audio_path), sr=16000, offset=start_time + offset, duration=10)
                    if len(audio) >= 1600:
                        samples.append(audio)
                if not samples:
                    return 'neutral', 0.5
                all_probs = []
                for audio in samples:
                    inputs = self.feature_extractor(audio, sampling_rate=16000, return_tensors="pt", padding=True)
                    inputs = {k: v.to(self.device) for k, v in inputs.items()}
                    with torch.no_grad():
                        outputs = self.model(**inputs)
                        probs = torch.softmax(outputs.logits, dim=-1)
                        all_probs.append(probs[0].cpu().numpy())
                avg_probs = np.mean(all_probs, axis=0)
                pred_idx = np.argmax(avg_probs)
                confidence = avg_probs[pred_idx]
            else:
                audio, sr = librosa.load(str(audio_path), sr=16000, offset=start_time, duration=duration)
                if len(audio) < 1600:
                    return 'neutral', 0.5
                inputs = self.feature_extractor(audio, sampling_rate=16000, return_tensors="pt", padding=True)
                inputs = {k: v.to(self.device) for k, v in inputs.items()}
                with torch.no_grad():
                    outputs = self.model(**inputs)
                    probs = torch.softmax(outputs.logits, dim=-1)
                    pred_idx = torch.argmax(probs, dim=-1).item()
                    confidence = probs[0][pred_idx].item()

            return self.id2label[pred_idx], confidence
        except Exception as e:
            print(f"Ошибка анализа эмоций: {e}")
            return 'neutral', 0.5


# === REPORTS ===
def format_time(seconds):
    hours = int(seconds // 3600)
    mins = int((seconds % 3600) // 60)
    secs = int(seconds % 60)
    if hours > 0:
        return f"{hours:02d}:{mins:02d}:{secs:02d}"
    return f"{mins:02d}:{secs:02d}"


def build_speaker_profiles(merged_segments: List[Dict]) -> Dict:
    """Строит профиль эмоций для каждого спикера"""
    profiles = defaultdict(lambda: {
        'emotions': [], 'total_time': 0, 'segment_count': 0,
        'emotion_counts': defaultdict(int), 'languages': set()
    })

    for seg in merged_segments:
        speaker = seg['speaker']
        emotion = seg.get('emotion', 'neutral')
        duration = seg['end'] - seg['start']
        language = seg.get('language', 'ru')

        profiles[speaker]['emotions'].append(emotion)
        profiles[speaker]['total_time'] += duration
        profiles[speaker]['segment_count'] += 1
        profiles[speaker]['emotion_counts'][emotion] += 1
        profiles[speaker]['languages'].add(language)

    for speaker, data in profiles.items():
        dominant_emotions = sorted(data['emotion_counts'].items(), key=lambda x: -x[1])
        dominant = dominant_emotions[0][0] if dominant_emotions else 'neutral'

        if dominant in ('happiness', 'enthusiasm'):
            data['interpretation'] = 'Позитивный настрой, энтузиазм'
        elif dominant == 'anger':
            data['interpretation'] = 'Напряжённость, возможно недовольство'
        elif dominant == 'sadness':
            data['interpretation'] = 'Обеспокоенность, усталость'
        elif dominant == 'fear':
            data['interpretation'] = 'Неуверенность, тревожность'
        elif dominant == 'disgust':
            data['interpretation'] = 'Негативное отношение'
        else:
            data['interpretation'] = 'Деловой, сдержанный тон'

        data['emoji_string'] = ''.join(EMOTION_EMOJI.get(e, '😐') for e in data['emotions'][:10])
        if len(data['emotions']) > 10:
            data['emoji_string'] += f'... (+{len(data["emotions"]) - 10})'

        data['languages'] = list(data['languages'])

    return dict(profiles)


def convert_to_serializable(obj):
    """Конвертирует numpy типы в Python типы для JSON"""
    if isinstance(obj, dict):
        return {k: convert_to_serializable(v) for k, v in obj.items()}
    elif isinstance(obj, list):
        return [convert_to_serializable(v) for v in obj]
    elif isinstance(obj, np.floating):
        return float(obj)
    elif isinstance(obj, np.integer):
        return int(obj)
    elif isinstance(obj, np.ndarray):
        return obj.tolist()
    elif isinstance(obj, set):
        return list(obj)
    else:
        return obj


def save_json(merged_segments: List[Dict], speaker_profiles: Dict,
              input_file: Path, output_path: Path):
    """Сохраняет JSON с результатами"""
    result = {
        "source_file": input_file.name,
        "processed_at": datetime.now().isoformat(),
        "segments_count": len(merged_segments),
        "speakers": convert_to_serializable(speaker_profiles),
        "segments": convert_to_serializable(merged_segments)
    }

    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(result, f, ensure_ascii=False, indent=2)
    print(f"JSON сохранён: {output_path}")


def save_txt(merged_segments: List[Dict], speaker_profiles: Dict, output_path: Path):
    """Сохраняет TXT отчёт"""
    with open(output_path, "w", encoding="utf-8") as f:
        f.write("=" * 60 + "\n")
        f.write("ПРОТОКОЛ СОВЕЩАНИЯ (мультиязычный)\n")
        f.write("=" * 60 + "\n\n")

        f.write("ПРОФИЛЬ УЧАСТНИКОВ:\n")
        f.write("-" * 60 + "\n")

        sorted_speakers = sorted(speaker_profiles.items(), key=lambda x: -x[1]['total_time'])
        for speaker, data in sorted_speakers:
            langs = ', '.join(LANGUAGE_FLAGS.get(l, l) for l in data['languages'])
            f.write(f"{speaker:<12} | {format_time(data['total_time']):<8} | {langs} | {data['interpretation']}\n")

        f.write("\n" + "=" * 60 + "\n")
        f.write("ТРАНСКРИПЦИЯ:\n")
        f.write("=" * 60 + "\n\n")

        for seg in merged_segments:
            emotion = seg.get("emotion", "neutral")
            emotion_str = f"{EMOTION_EMOJI.get(emotion, '')} {EMOTION_LABELS_RU.get(emotion, emotion)}"
            lang = seg.get("language", "ru")
            lang_flag = LANGUAGE_FLAGS.get(lang, lang)

            f.write(f"[{format_time(seg['start'])} - {format_time(seg['end'])}] {seg['speaker']} | {lang_flag} | {emotion_str}\n")

            # Если есть перевод — показываем оригинал и перевод
            if seg.get("original_text"):
                f.write(f"{seg['original_text']}\n")
                f.write(f"  → {seg['text']}\n\n")
            else:
                f.write(f"{seg['text']}\n\n")

    print(f"TXT сохранён: {output_path}")


def create_word_report(merged_segments: List[Dict], speaker_profiles: Dict,
                       input_file: Path, output_path: Path):
    """Создаёт Word отчёт с поддержкой мультиязычности"""
    doc = Document()

    title = doc.add_heading('Протокол совещания', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('Информация о записи', level=1)
    total_duration = merged_segments[-1]["end"] if merged_segments else 0

    info_table = doc.add_table(rows=5, cols=2)
    info_table.style = 'Table Grid'
    info_data = [
        ('Файл', input_file.name),
        ('Дата обработки', datetime.now().strftime('%d.%m.%Y %H:%M')),
        ('Длительность', format_time(total_duration)),
        ('Участников', str(len(speaker_profiles))),
        ('Сегментов', str(len(merged_segments))),
    ]
    for i, (label, value) in enumerate(info_data):
        row = info_table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        row.cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()
    doc.add_heading('Транскрипция', level=1)

    for seg in merged_segments:
        speaker = seg["speaker"]
        start = seg["start"]
        end = seg["end"]
        emotion = seg.get("emotion", "neutral")
        lang = seg.get("language", "ru")

        emotion_str = f"{EMOTION_EMOJI.get(emotion, '')} {EMOTION_LABELS_RU.get(emotion, emotion)}"
        lang_flag = LANGUAGE_FLAGS.get(lang, lang)

        # Заголовок сегмента
        header = doc.add_paragraph()
        header_run = header.add_run(f"[{format_time(start)} - {format_time(end)}] {speaker} | {lang_flag} | {emotion_str}")
        header_run.bold = True
        header_run.font.size = Pt(11)

        # Текст
        if seg.get("original_text"):
            # Оригинал
            orig_para = doc.add_paragraph(seg['original_text'])

            # Перевод курсивом
            trans_para = doc.add_paragraph()
            trans_run = trans_para.add_run(f"→ {seg['text']}")
            trans_run.italic = True
            trans_run.font.color.rgb = RGBColor(80, 80, 80)
        else:
            text_para = doc.add_paragraph(seg['text'])

        doc.add_paragraph()  # Отступ

    doc.save(str(output_path))
    print(f"Word отчёт сохранён: {output_path}")


# === MAIN PIPELINE ===
def process_multilang(
    input_file: Path,
    output_dir: Path = OUTPUT_DIR,
    languages: List[str] = ["ru", "zh"],
    model: str = "large-v3",
    device: str = "cuda",
    compute_type: str = "float16",
    batch_size: int = 16,
    skip_emotions: bool = False,
    skip_translation: bool = False,
    use_language_map: bool = True  # Использовать быстрый скан для карты языков
) -> Path:
    """
    Мультиязычный пайплайн транскрипции.

    Новая схема (если use_language_map=True):
    1. Быстрый скан (tiny) с автоопределением языка → карта языков
    2. Точные проходы (large-v3) для каждого языка с фильтрацией по карте
    3. Мердж на основе карты языков (нет угадывания!)
    """
    start_time_total = datetime.now()

    print("=" * 60)
    print("WhisperX Multilingual Pipeline v2")
    print(f"Файл: {input_file}")
    print(f"Языки: {languages}")
    print(f"Модель: {model}")
    print(f"Карта языков: {'Да (tiny scan)' if use_language_map else 'Нет'}")
    print("=" * 60)

    output_dir.mkdir(parents=True, exist_ok=True)
    temp_audio = output_dir / "temp_audio.wav"

    # 1. Извлечение аудио
    print("\n[1/9] Извлечение аудио...")
    audio_file = extract_audio(input_file, temp_audio)

    # 2. Загрузка аудио
    print("\n[2/9] Загрузка аудио...")
    audio = whisperx.load_audio(str(audio_file))

    # 3. НОВОЕ: Быстрый скан для карты языков
    language_map = None
    if use_language_map and len(languages) > 1:
        print("\n[3/9] Быстрый скан языков (tiny model)...")
        language_map = quick_language_scan(audio, device=device, languages=languages)

        # Фильтруем карту от галлюцинаций сразу
        language_map_clean = []
        removed_hallucinations = 0
        for seg in language_map:
            if is_hallucination_smart(seg, no_speech_threshold=0.5, logprob_threshold=-0.8):
                removed_hallucinations += 1
            else:
                language_map_clean.append(seg)

        print(f"  Карта языков: {len(language_map)} → {len(language_map_clean)} (удалено галлюцинаций: {removed_hallucinations})")

        # Сохраняем карту в JSON для отладки
        lang_map_debug = output_dir / "debug_language_map.json"
        with open(lang_map_debug, 'w', encoding='utf-8') as f:
            json.dump(language_map, f, ensure_ascii=False, indent=2)
        print(f"  Карта сохранена: {lang_map_debug}")

        language_map = language_map_clean
    else:
        print("\n[3/9] Быстрый скан пропущен (один язык или отключено)")

    # 4. Транскрипция (точные проходы для каждого языка)
    print("\n[4/9] Транскрипция (точные проходы)...")

    all_segments = []

    for lang in languages:
        print(f"\n  === Проход: {lang.upper()} ===")
        segments = transcribe_language(
            audio, model, lang, device, compute_type, batch_size,
            language_map=language_map  # Передаём карту для фильтрации
        )

        all_segments.extend(segments)
        print(f"  Сегментов ({lang}): {len(segments)}")

    # 5. Умная дедупликация: выбираем лучший сегмент для каждого временного интервала
    print("\n[5/9] Дедупликация...")
    merged_raw = smart_deduplicate(all_segments)
    print(f"    После дедупликации: {len(merged_raw)}")

    # 6. Выравнивание
    print("\n[6/9] Выравнивание...")

    # Сохраняем языки до выравнивания (по тексту)
    lang_map = {seg.get("text", ""): seg.get("language", "ru") for seg in merged_raw}

    model_a, metadata = whisperx.load_align_model(language_code=languages[0], device=device)
    result = whisperx.align(merged_raw, model_a, metadata, audio, device=device)
    aligned_segments = result["segments"]

    # Восстанавливаем языки по тексту
    for seg in aligned_segments:
        text = seg.get("text", "")
        seg["language"] = lang_map.get(text, detect_language_by_text(text))

    del model_a
    torch.cuda.empty_cache()

    # 7. Диаризация
    print("\n[7/9] Диаризация...")
    hf_token = os.getenv("HUGGINGFACE_TOKEN")
    diarize_model = DiarizationPipeline(use_auth_token=hf_token, device=device)
    diarize_segments = diarize_model(audio)
    result = whisperx.assign_word_speakers(diarize_segments, {"segments": aligned_segments})
    segments_with_speakers = result["segments"]

    del diarize_model
    torch.cuda.empty_cache()

    # Склейка по спикерам
    print("\nСклейка сегментов...")
    merged = merge_speaker_segments(segments_with_speakers)
    print(f"    После склейки: {len(merged)}")

    # Фильтрация
    print("\nФильтрация галлюцинаций и UNKNOWN...")
    merged = filter_segments(merged)
    print(f"    После фильтрации: {len(merged)}")

    # 8. Перевод
    if not skip_translation and len(languages) > 1:
        print("\n[8/9] Перевод через Gemini...")
        gemini_model = init_gemini()
        merged = translate_segments(merged, gemini_model, target_lang="ru")
    else:
        print("\n[8/9] Перевод пропущен")

    # 9. Эмоции
    if not skip_emotions:
        print("\n[9/9] Анализ эмоций...")
        emotion_analyzer = EmotionAnalyzer(EMOTION_MODEL, device=device)

        for i, seg in enumerate(merged):
            emotion, confidence = emotion_analyzer.analyze(audio_file, seg["start"], seg["end"])
            seg["emotion"] = emotion
            seg["emotion_confidence"] = confidence
            if (i + 1) % 10 == 0 or i == len(merged) - 1:
                print(f"    Прогресс: {i+1}/{len(merged)}")
    else:
        print("\n[9/9] Анализ эмоций пропущен")
        for seg in merged:
            seg["emotion"] = "neutral"
            seg["emotion_confidence"] = 0.0

    # Профили спикеров
    print("\nПостроение профилей спикеров...")
    speaker_profiles = build_speaker_profiles(merged)

    # Сохранение
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')

    json_path = output_dir / f"multilang_{timestamp}.json"
    save_json(merged, speaker_profiles, input_file, json_path)

    txt_path = output_dir / f"multilang_{timestamp}.txt"
    save_txt(merged, speaker_profiles, txt_path)

    word_path = output_dir / f"multilang_{timestamp}.docx"
    create_word_report(merged, speaker_profiles, input_file, word_path)

    # Удаляем временный файл только если он был создан (не был уже WAV)
    if audio_file == temp_audio and temp_audio.exists():
        temp_audio.unlink()

    elapsed = datetime.now() - start_time_total
    print("\n" + "=" * 60)
    print(f"ГОТОВО! Время обработки: {elapsed}")
    print(f"Результаты:")
    print(f"  - JSON: {json_path}")
    print(f"  - TXT:  {txt_path}")
    print(f"  - DOCX: {word_path}")
    print("=" * 60)

    return word_path


# === ENTRY POINT ===
if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="Мультиязычная транскрипция (RU + ZH)")
    parser.add_argument("input", nargs="?", default=None, help="Входной файл (mp4/wav)")
    parser.add_argument("--skip-emotions", action="store_true", help="Пропустить анализ эмоций")
    parser.add_argument("--skip-translation", action="store_true", help="Пропустить перевод")
    parser.add_argument("--languages", nargs="+", default=["ru", "zh"], help="Языки для транскрипции")
    args = parser.parse_args()

    # Дефолтный тестовый файл
    if args.input:
        INPUT_FILE = Path(args.input)
    else:
        # Используем наш тестовый сегмент 30-45 минут
        INPUT_FILE = PROJECT_ROOT / "test_segment_30-45min.wav"

    print("=" * 60)
    print("ТЕСТ: Мультиязычная транскрипция (RU + ZH)")
    print(f"Файл: {INPUT_FILE}")
    print(f"Языки: {args.languages}")
    print("=" * 60)

    if not INPUT_FILE.exists():
        print(f"ОШИБКА: Файл не найден: {INPUT_FILE}")
        sys.exit(1)

    result = process_multilang(
        input_file=INPUT_FILE,
        output_dir=OUTPUT_DIR,
        languages=args.languages,
        model="large-v3",
        device="cuda",
        skip_emotions=args.skip_emotions,
        skip_translation=args.skip_translation
    )

    print(f"\nРезультат: {result}")
