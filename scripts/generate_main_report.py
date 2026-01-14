#!/usr/bin/env python3
"""
Generate main report from transcript and attach to analytics.

Creates a DOCX report using main_report_prompt and updates the database.

Usage:
    docker exec whisperx-api python /app/scripts/generate_main_report.py
"""
import asyncio
import sys
import json
import os
from pathlib import Path
from datetime import datetime
import warnings

warnings.filterwarnings('ignore')

# Add project root to path
sys.path.insert(0, str(Path(__file__).parent.parent))

from sqlalchemy import select
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import google.generativeai as genai

from backend.shared.database import async_session_factory, init_db
from backend.domains.construction.models import (
    ConstructionReportDB,
)


# Gemini configuration
GEMINI_API_KEY = os.environ.get('GEMINI_API_KEY')
GEMINI_MODEL = os.environ.get('GEMINI_MODEL_NAME', 'gemini-2.5-flash')

# Output directory
OUTPUT_DIR = Path("/app/data/output")


def extract_transcript(docx_path: str) -> str:
    """Extract text from docx file."""
    doc = Document(docx_path)
    return '\n'.join([p.text.strip() for p in doc.paragraphs if p.text.strip()])


def generate_main_report(transcript_text: str, meeting_date: str) -> dict:
    """Generate main report using Gemini API."""
    if not GEMINI_API_KEY:
        print("WARNING: GEMINI_API_KEY not set, using mock data")
        return {
            "summary": "Совещание по бытовому городку PowerChina. Обсуждались вопросы планировки, водоснабжения и электроснабжения.",
            "meeting_domain": "Строительство и СМР",
            "expert_analysis": "Совещание прошло продуктивно, несмотря на некоторые разногласия по нормативам.",
            "tasks": [
                {
                    "category": "Инженерные системы",
                    "items": [
                        {
                            "id_num": 1,
                            "task_description": "Уточнить требования по пожарному водоснабжению",
                            "responsible_parties": ["Кравт"],
                            "deadline": "15.01.2026",
                            "notes_status": "Требуется экспертиза"
                        }
                    ]
                }
            ]
        }

    genai.configure(api_key=GEMINI_API_KEY)

    # Load prompt from prompts.json
    prompts_path = "/app/scripts/../prompts.json"
    if os.path.exists(prompts_path):
        with open(prompts_path, 'r', encoding='utf-8') as f:
            prompts = json.load(f)
    else:
        # Try alternative path
        prompts_path = "/app/prompts.json"
        if os.path.exists(prompts_path):
            with open(prompts_path, 'r', encoding='utf-8') as f:
                prompts = json.load(f)
        else:
            raise FileNotFoundError("prompts.json not found")

    prompt_config = prompts.get("main_report_prompt", {})

    # Build full prompt
    full_prompt = f"""{prompt_config.get('system_role', '')}

{prompt_config.get('task_description', '')}

{prompt_config.get('rules_and_structure', '').replace('{meeting_date}', meeting_date)}

{prompt_config.get('final_instruction', '').replace('{transcript_text}', transcript_text)}"""

    print(f"Calling Gemini API ({GEMINI_MODEL}) for main report...")
    model = genai.GenerativeModel(GEMINI_MODEL)
    response = model.generate_content(
        full_prompt,
        generation_config=genai.GenerationConfig(
            response_mime_type='application/json',
            temperature=0.3,
        )
    )

    return json.loads(response.text)


def create_report_docx(report_data: dict, output_path: Path, meeting_date: str) -> None:
    """Create DOCX report from JSON data."""
    doc = Document()

    # Title
    title = doc.add_heading('Отчёт по совещанию', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Date
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f'Дата: {meeting_date}')
    date_run.font.size = Pt(12)
    date_run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()

    # Domain
    if report_data.get('meeting_domain'):
        domain_para = doc.add_paragraph()
        domain_run = domain_para.add_run(f"Тема: {report_data['meeting_domain']}")
        domain_run.bold = True
        domain_run.font.size = Pt(11)

    # Summary
    doc.add_heading('Краткое содержание', level=1)
    doc.add_paragraph(report_data.get('summary', 'Нет данных'))

    # Expert Analysis
    if report_data.get('expert_analysis'):
        doc.add_heading('Экспертный анализ', level=1)
        doc.add_paragraph(report_data['expert_analysis'])

    # Tasks by category
    tasks = report_data.get('tasks', [])
    if tasks:
        doc.add_heading('Задачи по категориям', level=1)

        for category_group in tasks:
            category_name = category_group.get('category', 'Другое')
            items = category_group.get('items', [])

            if items:
                doc.add_heading(category_name, level=2)

                # Create table for tasks
                table = doc.add_table(rows=1, cols=5)
                table.style = 'Table Grid'

                # Header row
                header_cells = table.rows[0].cells
                headers = ['№', 'Задача', 'Ответственные', 'Срок', 'Статус']
                for i, header in enumerate(headers):
                    header_cells[i].text = header
                    header_cells[i].paragraphs[0].runs[0].bold = True

                # Data rows
                for item in items:
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(item.get('id_num', ''))
                    row_cells[1].text = item.get('task_description', '')

                    responsible = item.get('responsible_parties', [])
                    if isinstance(responsible, list):
                        row_cells[2].text = ', '.join(responsible)
                    else:
                        row_cells[2].text = str(responsible)

                    row_cells[3].text = item.get('deadline', 'Не указан')
                    row_cells[4].text = item.get('notes_status', '')

                doc.add_paragraph()

    # Save
    doc.save(output_path)
    print(f"Report saved to: {output_path}")


async def main():
    """Main function."""
    await init_db()

    print("\n" + "=" * 60)
    print("GENERATING MAIN REPORT FROM TRANSCRIPT")
    print("=" * 60 + "\n")

    # Extract transcript
    transcript_path = "/app/result_transcript.docx"
    if not os.path.exists(transcript_path):
        print(f"ERROR: Transcript file not found: {transcript_path}")
        return

    transcript_text = extract_transcript(transcript_path)
    print(f"Transcript extracted: {len(transcript_text)} chars")

    # Meeting date
    meeting_date = "02.12.2025"

    # Generate report with Gemini
    report_data = generate_main_report(transcript_text, meeting_date)
    print("Report generated successfully")
    print(json.dumps(report_data, indent=2, ensure_ascii=False)[:500] + "...")

    # Create output directory
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    # Generate DOCX report
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    report_filename = f"main_report_{timestamp}.docx"
    report_path = OUTPUT_DIR / report_filename

    create_report_docx(report_data, report_path, meeting_date)

    # Update database
    async with async_session_factory() as session:
        # Get the report
        result = await session.execute(
            select(ConstructionReportDB).where(
                ConstructionReportDB.title == "Совещание по бытовому городку ПОС"
            )
        )
        report = result.scalar_one_or_none()

        if report:
            report.report_path = str(report_path)
            await session.commit()
            print(f"Updated report record with path: {report_path}")
        else:
            print("WARNING: Report record not found in database")

    print("\n" + "=" * 60)
    print("MAIN REPORT GENERATED SUCCESSFULLY!")
    print("=" * 60)
    print(f"\nOutput: {report_path}")
    print("Refresh dashboard to see download button")


if __name__ == "__main__":
    asyncio.run(main())
