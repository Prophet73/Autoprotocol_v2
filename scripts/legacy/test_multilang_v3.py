#!/usr/bin/env python3
"""
Мультиязычная транскрипция v3: VAD-first + параллельная транскрипция.

Схема:
1. VAD (silero) определяет где есть речь
2. Для КАЖДОГО VAD-сегмента транскрибируем ОБОИМИ языками (RU + ZH)
3. Выбираем лучший результат по scoring (logprob, соответствие языку, длина)
4. Диаризация + эмоции + перевод

Преимущества:
- Не теряем контент из-за неправильного определения языка
- VAD отсекает тишину → нет галлюцинаций из ничего
- Scoring объективно выбирает лучший вариант
"""

import torch
import os
import sys
import re
import json
import time
from pathlib import Path
from datetime import datetime
from collections import defaultdict
from typing import Optional, List, Dict, Any, Tuple

# Patch torch.load BEFORE any imports (PyTorch 2.8+ fix)
_original_torch_load = torch.load
def _patched_torch_load(*args, **kwargs):
    kwargs['weights_only'] = False
    return _original_torch_load(*args, **kwargs)
torch.load = _patched_torch_load

import lightning_fabric.utilities.cloud_io as cloud_io
cloud_io.torch.load = _patched_torch_load

import whisperx
from whisperx.diarize import DiarizationPipeline
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from transformers import Wav2Vec2ForSequenceClassification, Wav2Vec2FeatureExtractor
import librosa
import subprocess
import numpy as np

# Google Gemini
import google.generativeai as genai
from dotenv import load_dotenv

# === LOAD ENV ===
PROJECT_ROOT = Path(__file__).parent.resolve()
load_dotenv(PROJECT_ROOT / ".env")

# === КОНСТАНТЫ ===
OUTPUT_DIR = PROJECT_ROOT / "output"
EMOTION_MODEL = "Aniemore/wav2vec2-xlsr-53-russian-emotion-recognition"

# === DEBUG / LOGGING ===
class DebugLog:
    """Сборщик отладочной информации о фильтрации"""

    def __init__(self):
        self.rejected_segments = []  # Отфильтрованные сегменты
        self.all_candidates = []     # Все кандидаты для каждого VAD слота
        self.vad_segments = []       # Исходные VAD сегменты
        self.stats = defaultdict(int)

    def add_rejected(self, segment: Dict, reason: str, stage: str):
        """Добавляет отфильтрованный сегмент"""
        self.rejected_segments.append({
            "segment": segment,
            "reason": reason,
            "stage": stage,
            "timestamp": f"{segment.get('start', 0):.2f}-{segment.get('end', 0):.2f}"
        })
        self.stats[f"{stage}:{reason}"] += 1

    def add_candidates(self, vad_idx: int, candidates: List[Dict], chosen: Dict, chosen_reason: str):
        """Сохраняет все кандидаты для VAD слота"""
        self.all_candidates.append({
            "vad_idx": vad_idx,
            "candidates": [{
                "lang": c.get("language"),
                "text": c.get("text", "")[:100],
                "score": round(c.get("score", 0), 3),
                "no_speech_prob": round(c.get("no_speech_prob", 0), 3),
                "avg_logprob": round(c.get("avg_logprob", 0), 3),
            } for c in candidates],
            "chosen_lang": chosen.get("language") if chosen else None,
            "chosen_reason": chosen_reason
        })

    def save(self, output_path: Path):
        """Сохраняет лог в JSON"""
        result = {
            "summary": {
                "total_rejected": len(self.rejected_segments),
                "by_reason": dict(self.stats),
                "vad_segments_count": len(self.vad_segments),
            },
            "rejected_segments": self.rejected_segments,
            "candidate_choices": self.all_candidates[:50],  # Первые 50 для примера
            "vad_segments": self.vad_segments,
        }

        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(result, f, ensure_ascii=False, indent=2, default=str)

        print(f"DEBUG LOG: {output_path}")

    def print_summary(self):
        """Выводит сводку"""
        print(f"\n  === DEBUG SUMMARY ===")
        print(f"  Всего отфильтровано: {len(self.rejected_segments)}")
        for key, count in sorted(self.stats.items(), key=lambda x: -x[1]):
            print(f"    - {key}: {count}")


# Глобальный лог (инициализируется в process_multilang_v3)
debug_log: Optional[DebugLog] = None

EMOTION_LABELS_RU = {
    'anger': 'Гнев', 'disgust': 'Отвращение', 'enthusiasm': 'Энтузиазм',
    'fear': 'Страх', 'happiness': 'Радость', 'neutral': 'Нейтрально', 'sadness': 'Грусть'
}
EMOTION_EMOJI = {
    'anger': '😠', 'disgust': '🤢', 'enthusiasm': '🤩',
    'fear': '😨', 'happiness': '😊', 'neutral': '😐', 'sadness': '😔'
}

LANGUAGE_FLAGS = {
    'ru': '🇷🇺', 'zh': '🇨🇳', 'en': '🇬🇧', 'unknown': '❓'
}

# === HALLUCINATION PATTERNS ===
HALLUCINATION_PATTERNS = [
    # Русские галлюцинации Whisper
    r'продолжение следует',
    r'субтитры\s*(сделал|подогнал|создал|делал)',
    r'редактор субтитров',
    r'корректор\s+[а-яё]+\.[а-яё]+',  # Корректор А.Егорова
    r'спасибо за просмотр',
    r'подписывайтесь на канал',
    r'ставьте лайк',
    r'не забудьте подписаться',
    r'до новых встреч',
    r'всем пока',
    r'^пока\.?$',
    r'^ага\.?$',
    r'^угу\.?$',
    # Китайские галлюцинации
    r'^谢谢大家\.?$',
    r'^谢谢\.?$',
    r'^谢谢观看',
    r'^感谢收看',
    r'^请订阅',
    r'^再见\.?$',
    # Английские галлюцинации
    r'^thank you\.?$',
    r'^thanks for watching',
    r'^please subscribe',
    r'^bye\.?$',
    r'^goodbye\.?$',
    # Имена/никнеймы (часто галлюцинации)
    r'DimaTorzok',
    r'Амели',
]


def clean_repetitions(text: str) -> str:
    """
    Удаляет повторяющиеся слова/фразы из текста.

    Примеры:
    - "текст Да, да, да, да, да." → "текст Да."
    - "банкомат, банкомат, банкомат, банкомат" → "банкомат"
    - "текст 個個個個 продолжение" → "текст 個 продолжение"
    """
    if not text or len(text) < 5:
        return text

    original_text = text

    # 1. Удаляем повторяющиеся символы подряд (個個個個, аааа) - оставляем 1
    text = re.sub(r'(.)\1{2,}', r'\1', text)

    # 2. Удаляем "Да, да, да, да" и подобные короткие повторы (в начале, чтобы не мешали другим паттернам)
    # Оставляем только первое вхождение
    text = re.sub(r'\b(да|нет|ага|угу|ну|так|вот|好|是|对|嗯)[,，.\s]*(\1[,，.\s]*){1,}', r'\1', text, flags=re.IGNORECASE)

    # 3. Удаляем повторяющиеся слова (слово, слово, слово → слово)
    # Паттерн: слово + разделитель, повторяющиеся 2+ раз
    text = re.sub(r'\b([\w\u4e00-\u9fff]{2,20})[,，\s]+(\1[,，\s]*){1,}\1?\b', r'\1', text, flags=re.IGNORECASE)

    # 4. Удаляем повторяющиеся фразы (2-5 слов)
    # "это хорошо, это хорошо, это хорошо" → "это хорошо"
    text = re.sub(r'(([\w\u4e00-\u9fff]+[,，\s]*){2,5}?)\1{1,}', r'\1', text)

    # 5. Ещё раз проверяем простые повторы типа "word word word"
    text = re.sub(r'\b(\w{2,})\s+(\1\s*){2,}', r'\1', text, flags=re.IGNORECASE)

    # 6. Чистим артефакты: множественные пробелы, запятые/точки в конце и начале
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'[,，.\s]+$', '', text)
    text = re.sub(r'^[,，.\s]+', '', text)
    text = text.strip()

    return text


def has_repeated_words(text: str, min_repeats: int = 3) -> bool:
    """
    Проверяет есть ли повторяющиеся слова/фразы (ПОСЛЕ очистки).
    Используется только для проверки, не для фильтрации.

    Примеры галлюцинаций:
    - "Да, да, да, да, да"
    - "банкомат, банкомат, банкомат"
    - "個個個個"
    """
    if not text or len(text) < 5:
        return False

    text_clean = text.strip().lower()

    # 1. Повторяющиеся символы подряд (個個個個, аааа)
    for i in range(len(text_clean) - 3):
        char = text_clean[i]
        if char.isalnum() and text_clean[i:i+4] == char * 4:
            return True

    # 2. Повторяющиеся слова
    # Разбиваем на слова, убираем пунктуацию
    words = re.findall(r'[\w\u4e00-\u9fff]+', text_clean)
    if len(words) >= min_repeats:
        # Считаем повторения каждого слова
        from collections import Counter
        word_counts = Counter(words)

        for word, count in word_counts.items():
            # Слово повторяется min_repeats+ раз И составляет >40% всех слов
            if count >= min_repeats and count / len(words) > 0.4:
                return True

    # 3. Повторяющиеся короткие паттерны (A-A-A, 10-10-10)
    repeat_pattern = r'(.{2,15}?)\1{2,}'
    if re.search(repeat_pattern, text_clean):
        return True

    return False


def matches_hallucination_pattern(text: str) -> bool:
    """Проверяет совпадение с паттернами галлюцинаций"""
    text_lower = text.lower().strip()

    for pattern in HALLUCINATION_PATTERNS:
        if re.search(pattern, text_lower, re.IGNORECASE):
            return True
        # Также проверяем оригинал (для китайского)
        if re.search(pattern, text, re.IGNORECASE):
            return True

    return False


def is_text_hallucination(text: str, duration: float = 0) -> Tuple[bool, str]:
    """
    Комплексная проверка текста на галлюцинацию.
    НЕ проверяет repetition - это делается через clean_repetitions.

    Returns:
        (is_hallucination, reason)
    """
    if not text or len(text.strip()) < 2:
        return True, "empty"

    text = text.strip()

    # 1. Паттерны галлюцинаций (весь текст = галлюцинация)
    if matches_hallucination_pattern(text):
        return True, "pattern"

    # 2. Слишком короткий текст для длинного сегмента
    if duration > 10:
        meaningful_chars = len([c for c in text if c.isalnum() or '\u4e00' <= c <= '\u9fff'])
        chars_per_sec = meaningful_chars / duration
        if chars_per_sec < 0.5:  # Меньше 0.5 символа в секунду
            return True, "low_density"

    return False, ""


# === VAD ===
def load_silero_vad(device: str = "cuda"):
    """Загружает Silero VAD модель"""
    print("Загрузка Silero VAD...")
    model, utils = torch.hub.load(
        repo_or_dir='snakers4/silero-vad',
        model='silero_vad',
        trust_repo=True
    )
    model = model.to(device)

    get_speech_timestamps = utils[0]
    read_audio = utils[2]

    return model, get_speech_timestamps, read_audio


def get_vad_segments(audio_path: Path, vad_model, get_speech_timestamps,
                     min_speech_duration_ms: int = 250,
                     min_silence_duration_ms: int = 100,
                     threshold: float = 0.5,
                     device: str = "cuda") -> List[Dict]:
    """
    Получает сегменты речи через VAD.

    Args:
        audio_path: Путь к аудио файлу
        vad_model: Модель Silero VAD
        get_speech_timestamps: Функция из utils
        min_speech_duration_ms: Минимальная длина речи в мс
        min_silence_duration_ms: Минимальная пауза между сегментами
        threshold: Порог VAD (0.5 = стандарт)
        device: Устройство (cuda/cpu)

    Returns:
        Список сегментов: [{"start": float, "end": float}, ...]
    """
    print(f"VAD анализ: {audio_path.name}")

    # Загружаем аудио для VAD (16kHz)
    wav, sr = librosa.load(str(audio_path), sr=16000)
    wav_tensor = torch.tensor(wav).to(device)

    # Получаем timestamps
    speech_timestamps = get_speech_timestamps(
        wav_tensor,
        vad_model,
        sampling_rate=16000,
        min_speech_duration_ms=min_speech_duration_ms,
        min_silence_duration_ms=min_silence_duration_ms,
        threshold=threshold,
        return_seconds=True  # Возвращаем в секундах
    )

    # Конвертируем в наш формат
    segments = []
    for ts in speech_timestamps:
        segments.append({
            "start": ts["start"],
            "end": ts["end"],
            "duration": ts["end"] - ts["start"]
        })

    total_speech = sum(s["duration"] for s in segments)
    total_audio = len(wav) / sr

    print(f"  Найдено {len(segments)} сегментов речи")
    print(f"  Речь: {total_speech:.1f}s / {total_audio:.1f}s ({100*total_speech/total_audio:.1f}%)")

    return segments


def merge_close_vad_segments(segments: List[Dict], max_gap: float = 1.0) -> List[Dict]:
    """
    Объединяет близкие VAD сегменты чтобы не резать речь.

    Args:
        segments: Список VAD сегментов
        max_gap: Максимальный промежуток для объединения (сек)
    """
    if not segments:
        return []

    merged = [segments[0].copy()]

    for seg in segments[1:]:
        last = merged[-1]
        gap = seg["start"] - last["end"]

        if gap <= max_gap:
            # Объединяем
            last["end"] = seg["end"]
            last["duration"] = last["end"] - last["start"]
        else:
            merged.append(seg.copy())

    print(f"  После объединения близких: {len(merged)} сегментов")
    return merged


# === LANGUAGE DETECTION ===
def detect_language_by_text(text: str) -> str:
    """Определяет язык по содержимому текста"""
    if not text:
        return "unknown"

    chinese_chars = sum(1 for c in text if '\u4e00' <= c <= '\u9fff')
    cyrillic_chars = sum(1 for c in text if '\u0400' <= c <= '\u04ff')
    latin_chars = sum(1 for c in text if 'a' <= c.lower() <= 'z')

    total = chinese_chars + cyrillic_chars + latin_chars
    if total == 0:
        return "unknown"

    if chinese_chars / total > 0.3:
        return "zh"
    elif cyrillic_chars / total > 0.3:
        return "ru"
    elif latin_chars / total > 0.5:
        return "en"
    else:
        return "unknown"


def calculate_language_score(text: str, expected_lang: str) -> float:
    """
    Вычисляет score соответствия текста ожидаемому языку.

    Returns:
        float: 0.0 - 1.0, где 1.0 = идеальное соответствие
    """
    if not text or len(text.strip()) < 2:
        return 0.0

    text = text.strip()

    chinese_chars = sum(1 for c in text if '\u4e00' <= c <= '\u9fff')
    cyrillic_chars = sum(1 for c in text if '\u0400' <= c <= '\u04ff')
    total_chars = len([c for c in text if c.isalpha()])

    if total_chars == 0:
        return 0.0

    if expected_lang == "zh":
        # Для китайского: доля иероглифов
        # Штраф за кириллицу (Whisper "перевёл")
        if cyrillic_chars > 0:
            return 0.1  # Почти отбраковка
        return chinese_chars / total_chars

    elif expected_lang == "ru":
        # Для русского: доля кириллицы
        # Штраф за иероглифы
        if chinese_chars > 0:
            return 0.1
        return cyrillic_chars / total_chars

    return 0.5


# === SCORING ===
def score_transcription(seg: Dict, expected_lang: str) -> float:
    """
    Вычисляет общий score качества транскрипции.

    Факторы:
    1. avg_logprob - уверенность модели (выше = лучше)
    2. no_speech_prob - вероятность отсутствия речи (ниже = лучше)
    3. language_score - соответствие текста языку
    4. text_length - длина текста (больше обычно лучше)
    5. compression_ratio - повторяемость (ниже = лучше)

    Returns:
        float: общий score (выше = лучше)
    """
    text = seg.get("text", "").strip()

    # Пустой текст = плохо
    if not text or len(text) < 2:
        return -100.0

    # 1. Уверенность модели (обычно от -1.5 до 0, где 0 = идеально)
    avg_logprob = seg.get("avg_logprob", -1.0)
    logprob_score = (avg_logprob + 1.5) / 1.5  # Нормализуем к 0-1
    logprob_score = max(0, min(1, logprob_score))

    # 2. Вероятность отсутствия речи (0-1, где 0 = точно речь)
    no_speech_prob = seg.get("no_speech_prob", 0.5)
    speech_score = 1.0 - no_speech_prob

    # 3. Соответствие языку
    lang_score = calculate_language_score(text, expected_lang)

    # 4. Длина текста (нормализованная)
    # Китайский: ~2-4 символа/сек, Русский: ~10-15 символов/сек
    duration = seg.get("end", 0) - seg.get("start", 0)
    if duration > 0:
        chars_per_sec = len(text) / duration
        # Нормализуем: 0-20 chars/sec -> 0-1
        length_score = min(1.0, chars_per_sec / 15.0)
    else:
        length_score = 0.5

    # 5. Compression ratio (1.0 = идеально, >2.5 = подозрительно)
    compression = seg.get("compression_ratio", 1.5)
    compression_score = max(0, 1.0 - (compression - 1.0) / 2.0)

    # Веса для итогового score
    weights = {
        "logprob": 0.25,
        "speech": 0.20,
        "language": 0.35,  # Самый важный!
        "length": 0.10,
        "compression": 0.10
    }

    total_score = (
        weights["logprob"] * logprob_score +
        weights["speech"] * speech_score +
        weights["language"] * lang_score +
        weights["length"] * length_score +
        weights["compression"] * compression_score
    )

    return total_score


def clean_segment_text(seg: Dict) -> Tuple[Dict, str]:
    """
    Очищает текст сегмента от повторений и галлюцинаций.

    Returns:
        (cleaned_segment, cleaning_info) - сегмент с очищенным текстом и инфо об очистке
    """
    text = seg.get("text", "").strip()
    original_text = text

    # 1. Убираем повторения
    text = clean_repetitions(text)

    # 2. Убираем паттерны галлюцинаций из середины текста
    for pattern in HALLUCINATION_PATTERNS:
        text = re.sub(pattern, '', text, flags=re.IGNORECASE)

    # 3. Чистим артефакты
    text = re.sub(r'\s+', ' ', text).strip()
    text = re.sub(r'^[,，.。\s]+', '', text)
    text = re.sub(r'[,，\s]+$', '', text)

    # Обновляем сегмент
    cleaned_seg = seg.copy()
    cleaned_seg["text"] = text

    # Инфо об очистке
    if text != original_text:
        cleaned_seg["original_text_raw"] = original_text
        return cleaned_seg, f"cleaned:{len(original_text)}→{len(text)}"

    return cleaned_seg, ""


def is_hallucination_by_score(seg: Dict, threshold: float = 0.25) -> Tuple[bool, str]:
    """
    Комплексная проверка сегмента на галлюцинацию.
    Вызывается ПОСЛЕ clean_segment_text.

    Returns:
        (is_hallucination, reason)
    """
    text = seg.get("text", "").strip()
    start = seg.get("start", 0)
    end = seg.get("end", 0)
    duration = end - start

    # 1. Пустой или очень короткий после очистки
    if len(text) < 3:
        return True, "empty_after_clean"

    # 2. Проверка текста на галлюцинацию (паттерны + плотность, БЕЗ repetition)
    is_hall, reason = is_text_hallucination(text, duration)
    if is_hall:
        return True, reason

    # 3. Whisper метаданные: высокая вероятность отсутствия речи
    if seg.get("no_speech_prob", 0) > 0.7:
        return True, "no_speech"

    # 4. Whisper метаданные: очень низкая уверенность
    if seg.get("avg_logprob", 0) < -1.2:
        return True, "low_confidence"

    # 5. Whisper метаданные: высокий compression ratio
    if seg.get("compression_ratio", 1.0) > 2.8:
        return True, "high_compression"

    # 6. Score ниже порога (пересчитываем для очищенного текста)
    lang = seg.get("language", "ru")
    score = score_transcription(seg, lang)
    if score < threshold:
        return True, f"low_score({score:.2f})"

    return False, ""


# === TRANSCRIPTION ===
def transcribe_segment_both_languages(
    audio: np.ndarray,
    start_time: float,
    end_time: float,
    languages: List[str],
    model_name: str,
    device: str,
    compute_type: str
) -> Dict:
    """
    Транскрибирует один сегмент аудио обоими языками и выбирает лучший.

    Returns:
        Dict с лучшей транскрипцией и метаданными
    """
    # Вырезаем сегмент (whisperx audio в samples, 16kHz)
    sr = 16000
    start_sample = int(start_time * sr)
    end_sample = int(end_time * sr)
    segment_audio = audio[start_sample:end_sample]

    results = {}

    for lang in languages:
        # Загружаем модель для языка
        model = whisperx.load_model(
            model_name, device,
            compute_type=compute_type,
            language=lang,
            asr_options={"condition_on_previous_text": False}
        )

        # Транскрибируем
        result = model.transcribe(segment_audio, batch_size=1)

        # Собираем текст и метаданные
        if result["segments"]:
            # Берём первый/главный сегмент
            seg = result["segments"][0]
            text = " ".join(s.get("text", "") for s in result["segments"]).strip()

            results[lang] = {
                "text": text,
                "language": lang,
                "avg_logprob": seg.get("avg_logprob", -1.0),
                "no_speech_prob": seg.get("no_speech_prob", 0.5),
                "compression_ratio": seg.get("compression_ratio", 1.5),
                "start": start_time,
                "end": end_time,
            }

            # Вычисляем score
            results[lang]["score"] = score_transcription(results[lang], lang)
        else:
            results[lang] = {
                "text": "",
                "language": lang,
                "score": -100,
                "start": start_time,
                "end": end_time,
            }

        del model
        torch.cuda.empty_cache()

    # Выбираем лучший по score
    best_lang = max(results.keys(), key=lambda l: results[l]["score"])
    best = results[best_lang]

    # Добавляем альтернативу для отладки
    best["alternatives"] = {l: {"text": r["text"], "score": r["score"]}
                           for l, r in results.items() if l != best_lang}

    return best


def transcribe_vad_segments_parallel(
    audio: np.ndarray,
    vad_segments: List[Dict],
    languages: List[str],
    model_name: str,
    device: str,
    compute_type: str,
    batch_size: int = 16
) -> List[Dict]:
    """
    Транскрибирует все VAD сегменты параллельно обоими языками.

    Оптимизация: загружаем модель один раз для каждого языка,
    транскрибируем все сегменты, потом переключаемся на другой язык.
    """
    print(f"\n=== ПАРАЛЛЕЛЬНАЯ ТРАНСКРИПЦИЯ ({len(vad_segments)} сегментов) ===")

    all_results = {lang: [] for lang in languages}

    # Для каждого языка
    for lang in languages:
        print(f"\n  Проход: {lang.upper()}")

        model = whisperx.load_model(
            model_name, device,
            compute_type=compute_type,
            language=lang,
            asr_options={"condition_on_previous_text": False}
        )

        for i, vad_seg in enumerate(vad_segments):
            start = vad_seg["start"]
            end = vad_seg["end"]

            # Вырезаем сегмент
            sr = 16000
            start_sample = int(start * sr)
            end_sample = int(end * sr)
            segment_audio = audio[start_sample:end_sample]

            if len(segment_audio) < sr * 0.1:  # < 100ms
                continue

            # Транскрибируем
            result = model.transcribe(segment_audio, batch_size=batch_size)

            if result["segments"]:
                # Объединяем все сегменты результата
                text = " ".join(s.get("text", "") for s in result["segments"]).strip()

                # Берём метаданные из первого сегмента
                seg = result["segments"][0]

                seg_result = {
                    "text": text,
                    "language": lang,
                    "start": start,
                    "end": end,
                    "vad_idx": i,
                    "avg_logprob": seg.get("avg_logprob", -1.0),
                    "no_speech_prob": seg.get("no_speech_prob", 0.5),
                    "compression_ratio": seg.get("compression_ratio", 1.5),
                }
                seg_result["score"] = score_transcription(seg_result, lang)

                all_results[lang].append(seg_result)

            if (i + 1) % 20 == 0:
                print(f"    Прогресс: {i+1}/{len(vad_segments)}")

        print(f"    {lang}: {len(all_results[lang])} результатов")

        del model
        torch.cuda.empty_cache()

    # Теперь выбираем лучший результат для каждого VAD сегмента
    print("\n  Выбор лучших результатов...")

    global debug_log
    final_segments = []
    hallucination_reasons = {}  # Статистика по причинам удаления
    cleaned_count = 0  # Сколько сегментов было очищено

    for i in range(len(vad_segments)):
        candidates = []
        for lang in languages:
            for seg in all_results[lang]:
                if seg.get("vad_idx") == i:
                    candidates.append(seg)

        if not candidates:
            # Логируем пустой слот
            if debug_log:
                debug_log.add_rejected(
                    {"start": vad_segments[i]["start"], "end": vad_segments[i]["end"], "text": ""},
                    "no_candidates", "selection"
                )
            continue

        # Выбираем лучший по score
        best = max(candidates, key=lambda x: x["score"])

        # === НОВОЕ: Сначала очищаем текст от повторений ===
        best, clean_info = clean_segment_text(best)
        if clean_info:
            cleaned_count += 1

        # Проверяем на галлюцинацию (после очистки)
        is_hall, reason = is_hallucination_by_score(best)
        if is_hall:
            hallucination_reasons[reason] = hallucination_reasons.get(reason, 0) + 1
            # Логируем отфильтрованный сегмент
            if debug_log:
                debug_log.add_rejected(best, reason, "hallucination_filter")
                debug_log.add_candidates(i, candidates, None, f"rejected:{reason}")
            continue

        # Логируем выбор
        if debug_log:
            debug_log.add_candidates(i, candidates, best, "accepted" + (f" ({clean_info})" if clean_info else ""))

        # Добавляем альтернативы для отладки
        best["alternatives"] = {
            c["language"]: {"text": c["text"][:50], "score": round(c["score"], 3)}
            for c in candidates if c != best
        }

        final_segments.append(best)

    # Статистика по очистке и удалению
    if cleaned_count > 0:
        print(f"  Очищено от повторений: {cleaned_count} сегментов")

    total_removed = sum(hallucination_reasons.values())
    if total_removed > 0:
        print(f"  Удалено галлюцинаций: {total_removed}")
        for reason, count in sorted(hallucination_reasons.items(), key=lambda x: -x[1]):
            print(f"    - {reason}: {count}")

    print(f"  Итого после выбора: {len(final_segments)} сегментов")

    # Статистика по языкам
    lang_counts = defaultdict(int)
    for seg in final_segments:
        lang_counts[seg["language"]] += 1
    print(f"  По языкам: {dict(lang_counts)}")

    return final_segments


# === AUDIO ===
def extract_audio(input_file: Path, output_file: Path) -> Path:
    """Извлекает аудио из видео"""
    if input_file.suffix.lower() == '.wav':
        print(f"Входной файл уже WAV: {input_file.name}")
        return input_file

    print(f"Извлечение аудио из {input_file.name}...")
    cmd = [
        'ffmpeg', '-i', str(input_file),
        '-vn', '-acodec', 'pcm_s16le', '-ar', '16000', '-ac', '1',
        '-y', str(output_file)
    ]
    subprocess.run(cmd, capture_output=True)
    return output_file


# === DIARIZATION & ALIGNMENT ===
def align_and_diarize(segments: List[Dict], audio: np.ndarray,
                      device: str, hf_token: str) -> List[Dict]:
    """Выравнивание и диаризация"""

    print("\nВыравнивание...")
    # Сохраняем языки
    lang_map = {seg.get("text", ""): seg.get("language", "ru") for seg in segments}

    # Выравниваем (используем ru как базовый)
    model_a, metadata = whisperx.load_align_model(language_code="ru", device=device)
    result = whisperx.align(segments, model_a, metadata, audio, device=device)
    aligned = result["segments"]

    # Восстанавливаем языки
    for seg in aligned:
        text = seg.get("text", "")
        seg["language"] = lang_map.get(text, detect_language_by_text(text))

    del model_a
    torch.cuda.empty_cache()

    print("Диаризация...")
    diarize_model = DiarizationPipeline(use_auth_token=hf_token, device=device)
    diarize_segments = diarize_model(audio)
    result = whisperx.assign_word_speakers(diarize_segments, {"segments": aligned})

    del diarize_model
    torch.cuda.empty_cache()

    return result["segments"]


def merge_speaker_segments(segments: List[Dict]) -> List[Dict]:
    """Склеивает последовательные сегменты одного спикера"""
    if not segments:
        return []

    merged = []
    current = None

    for seg in segments:
        speaker = seg.get("speaker", "UNKNOWN")
        text = seg.get("text", "").strip()
        start = seg.get("start", 0)
        end = seg.get("end", 0)
        language = seg.get("language", "ru")

        if current is None:
            current = {
                "speaker": speaker, "start": start, "end": end,
                "text": text, "language": language
            }
        elif current["speaker"] == speaker and current["language"] == language:
            current["end"] = end
            current["text"] += " " + text
        else:
            merged.append(current)
            current = {
                "speaker": speaker, "start": start, "end": end,
                "text": text, "language": language
            }

    if current:
        merged.append(current)

    return merged


def filter_unknown_speakers(segments: List[Dict]) -> List[Dict]:
    """Фильтрует UNKNOWN спикеров"""
    global debug_log
    filtered = []
    removed = 0

    for seg in segments:
        if seg.get("speaker") in ("UNKNOWN", None, ""):
            removed += 1
            if debug_log:
                debug_log.add_rejected(seg, "unknown_speaker", "diarization")
        else:
            filtered.append(seg)

    if removed:
        print(f"  Удалено UNKNOWN: {removed}")
    return filtered


def clean_hallucination_text(segments: List[Dict]) -> List[Dict]:
    """
    Очищает текст от галлюцинаций внутри сегментов.

    Удаляет фразы типа "Продолжение следует..." из середины текста.
    """
    global debug_log

    # Фразы для удаления из текста (не регулярки, а литералы)
    phrases_to_remove = [
        "Продолжение следует...",
        "Продолжение следует",
        "продолжение следует...",
        "продолжение следует",
        "Субтитры сделал DimaTorzok",
        "Редактор субтитров",
        "Корректор",
    ]

    cleaned = []
    removed_count = 0

    for seg in segments:
        text = seg.get("text", "")
        original_text = text

        # Удаляем фразы
        for phrase in phrases_to_remove:
            text = text.replace(phrase, "")

        # Очищаем лишние пробелы
        text = re.sub(r'\s+', ' ', text).strip()

        # Если текст стал слишком коротким — пропускаем сегмент
        if len(text) < 5:
            removed_count += 1
            if debug_log:
                seg_copy = seg.copy()
                seg_copy["original_text_before_clean"] = original_text
                debug_log.add_rejected(seg_copy, "text_cleaned_empty", "post_process")
            continue

        seg["text"] = text
        cleaned.append(seg)

    if removed_count:
        print(f"  Удалено после очистки текста: {removed_count}")

    return cleaned


# === TRANSLATION ===
def init_gemini():
    """Инициализирует Gemini API"""
    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key:
        raise ValueError("GEMINI_API_KEY не найден в .env")
    genai.configure(api_key=api_key)
    return genai.GenerativeModel("gemini-2.5-flash")


def translate_text(model, text: str, source_lang: str) -> str:
    """Переводит текст через Gemini"""
    lang_names = {'zh': 'китайского', 'en': 'английского'}
    source_name = lang_names.get(source_lang, source_lang)

    prompt = f"""Переведи следующий текст с {source_name} на русский язык.
Сохрани смысл и стиль оригинала. Верни ТОЛЬКО перевод, без пояснений.

Текст: {text}"""

    try:
        response = model.generate_content(prompt)
        return response.text.strip()
    except Exception as e:
        print(f"    Ошибка перевода: {e}")
        return f"[Ошибка перевода: {text}]"


def translate_segments(segments: List[Dict], model) -> List[Dict]:
    """Переводит все не-русские сегменты"""
    translated_count = 0

    for seg in segments:
        lang = seg.get("language", "ru")

        if lang != "ru" and seg.get("text"):
            original = seg["text"]
            translation = translate_text(model, original, lang)

            seg["original_text"] = original
            seg["translation"] = translation
            seg["text"] = translation
            translated_count += 1

            time.sleep(1)  # Rate limiting

    print(f"  Переведено: {translated_count} сегментов")
    return segments


# === EMOTIONS ===
class EmotionAnalyzer:
    def __init__(self, model_name: str = EMOTION_MODEL, device: str = "cuda"):
        print(f"Загрузка модели эмоций: {model_name}")
        self.device = device
        self.feature_extractor = Wav2Vec2FeatureExtractor.from_pretrained(model_name)
        self.model = Wav2Vec2ForSequenceClassification.from_pretrained(model_name).to(device)
        self.model.eval()
        self.id2label = self.model.config.id2label

    def analyze(self, audio_path: Path, start_time: float, end_time: float):
        try:
            duration = end_time - start_time
            audio, sr = librosa.load(str(audio_path), sr=16000,
                                    offset=start_time, duration=min(duration, 30))
            if len(audio) < 1600:
                return 'neutral', 0.5

            inputs = self.feature_extractor(audio, sampling_rate=16000,
                                           return_tensors="pt", padding=True)
            inputs = {k: v.to(self.device) for k, v in inputs.items()}

            with torch.no_grad():
                outputs = self.model(**inputs)
                probs = torch.softmax(outputs.logits, dim=-1)
                pred_idx = torch.argmax(probs, dim=-1).item()
                confidence = probs[0][pred_idx].item()

            return self.id2label[pred_idx], confidence
        except Exception as e:
            return 'neutral', 0.5


# === REPORTS ===
def format_time(seconds):
    hours = int(seconds // 3600)
    mins = int((seconds % 3600) // 60)
    secs = int(seconds % 60)
    if hours > 0:
        return f"{hours:02d}:{mins:02d}:{secs:02d}"
    return f"{mins:02d}:{secs:02d}"


def build_speaker_profiles(segments: List[Dict]) -> Dict:
    """Строит профиль для каждого спикера"""
    profiles = defaultdict(lambda: {
        'emotions': [], 'total_time': 0, 'segment_count': 0,
        'emotion_counts': defaultdict(int), 'languages': set()
    })

    for seg in segments:
        speaker = seg['speaker']
        emotion = seg.get('emotion', 'neutral')
        duration = seg['end'] - seg['start']
        language = seg.get('language', 'ru')

        profiles[speaker]['emotions'].append(emotion)
        profiles[speaker]['total_time'] += duration
        profiles[speaker]['segment_count'] += 1
        profiles[speaker]['emotion_counts'][emotion] += 1
        profiles[speaker]['languages'].add(language)

    for speaker, data in profiles.items():
        dominant_emotions = sorted(data['emotion_counts'].items(), key=lambda x: -x[1])
        dominant = dominant_emotions[0][0] if dominant_emotions else 'neutral'

        if dominant in ('happiness', 'enthusiasm'):
            data['interpretation'] = 'Позитивный настрой'
        elif dominant == 'anger':
            data['interpretation'] = 'Напряжённость'
        elif dominant == 'sadness':
            data['interpretation'] = 'Обеспокоенность'
        else:
            data['interpretation'] = 'Деловой, сдержанный тон'

        data['languages'] = list(data['languages'])

    return dict(profiles)


def save_txt(segments: List[Dict], profiles: Dict, output_path: Path):
    """Сохраняет TXT отчёт"""
    with open(output_path, "w", encoding="utf-8") as f:
        f.write("=" * 60 + "\n")
        f.write("ПРОТОКОЛ СОВЕЩАНИЯ (v3: VAD + параллельная транскрипция)\n")
        f.write("=" * 60 + "\n\n")

        f.write("ПРОФИЛЬ УЧАСТНИКОВ:\n")
        f.write("-" * 60 + "\n")

        for speaker, data in sorted(profiles.items(), key=lambda x: -x[1]['total_time']):
            langs = ', '.join(LANGUAGE_FLAGS.get(l, l) for l in data['languages'])
            f.write(f"{speaker:<12} | {format_time(data['total_time']):<8} | {langs} | {data['interpretation']}\n")

        f.write("\n" + "=" * 60 + "\n")
        f.write("ТРАНСКРИПЦИЯ:\n")
        f.write("=" * 60 + "\n\n")

        for seg in segments:
            emotion = seg.get("emotion", "neutral")
            emotion_str = f"{EMOTION_EMOJI.get(emotion, '')} {EMOTION_LABELS_RU.get(emotion, emotion)}"
            lang = seg.get("language", "ru")
            lang_flag = LANGUAGE_FLAGS.get(lang, lang)

            f.write(f"[{format_time(seg['start'])} - {format_time(seg['end'])}] {seg['speaker']} | {lang_flag} | {emotion_str}\n")

            if seg.get("original_text"):
                f.write(f"{seg['original_text']}\n")
                f.write(f"  → {seg['text']}\n\n")
            else:
                f.write(f"{seg['text']}\n\n")

    print(f"TXT: {output_path}")


def save_json(segments: List[Dict], profiles: Dict, input_file: Path, output_path: Path):
    """Сохраняет JSON"""
    def convert(obj):
        if isinstance(obj, (np.floating, np.integer)):
            return float(obj) if isinstance(obj, np.floating) else int(obj)
        if isinstance(obj, set):
            return list(obj)
        if isinstance(obj, dict):
            return {k: convert(v) for k, v in obj.items()}
        if isinstance(obj, list):
            return [convert(v) for v in obj]
        return obj

    result = {
        "source_file": input_file.name,
        "processed_at": datetime.now().isoformat(),
        "pipeline": "v3_vad_parallel",
        "segments_count": len(segments),
        "speakers": convert(profiles),
        "segments": convert(segments)
    }

    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(result, f, ensure_ascii=False, indent=2)
    print(f"JSON: {output_path}")


def create_word_report(segments: List[Dict], profiles: Dict,
                       input_file: Path, output_path: Path):
    """Word отчёт"""
    doc = Document()

    title = doc.add_heading('Протокол совещания (v3)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('Информация', level=1)
    total_duration = segments[-1]["end"] if segments else 0

    info = doc.add_table(rows=4, cols=2)
    info.style = 'Table Grid'
    info_data = [
        ('Файл', input_file.name),
        ('Дата', datetime.now().strftime('%d.%m.%Y %H:%M')),
        ('Длительность', format_time(total_duration)),
        ('Участников', str(len(profiles))),
    ]
    for i, (k, v) in enumerate(info_data):
        info.rows[i].cells[0].text = k
        info.rows[i].cells[1].text = v

    doc.add_paragraph()
    doc.add_heading('Транскрипция', level=1)

    for seg in segments:
        emotion = seg.get("emotion", "neutral")
        emotion_str = f"{EMOTION_EMOJI.get(emotion, '')} {EMOTION_LABELS_RU.get(emotion, emotion)}"
        lang_flag = LANGUAGE_FLAGS.get(seg.get("language", "ru"), "")

        header = doc.add_paragraph()
        run = header.add_run(
            f"[{format_time(seg['start'])} - {format_time(seg['end'])}] "
            f"{seg['speaker']} | {lang_flag} | {emotion_str}"
        )
        run.bold = True

        if seg.get("original_text"):
            doc.add_paragraph(seg['original_text'])
            trans = doc.add_paragraph()
            trans_run = trans.add_run(f"→ {seg['text']}")
            trans_run.italic = True
        else:
            doc.add_paragraph(seg['text'])

        doc.add_paragraph()

    doc.save(str(output_path))
    print(f"DOCX: {output_path}")


# === MAIN PIPELINE ===
def process_multilang_v3(
    input_file: Path,
    output_dir: Path = OUTPUT_DIR,
    languages: List[str] = ["ru", "zh"],
    model: str = "large-v3",
    device: str = "cuda",
    compute_type: str = "float16",
    batch_size: int = 16,
    skip_emotions: bool = False,
    skip_translation: bool = False,
    skip_diarization: bool = False,
    vad_threshold: float = 0.5,
    vad_min_speech_ms: int = 250,
) -> Path:
    """
    VAD-first мультиязычный пайплайн.

    1. VAD → где есть речь
    2. Параллельная транскрипция обоими языками
    3. Scoring → выбор лучшего
    4. Диаризация + эмоции + перевод
    """
    global debug_log
    debug_log = DebugLog()  # Инициализируем лог

    start_time_total = datetime.now()

    print("=" * 60)
    print("WhisperX Multilingual Pipeline v3 (VAD + Parallel)")
    print(f"Файл: {input_file}")
    print(f"Языки: {languages}")
    print(f"Модель: {model}")
    print("=" * 60)

    output_dir.mkdir(parents=True, exist_ok=True)
    temp_audio = output_dir / "temp_audio.wav"

    # 1. Аудио
    print("\n[1/8] Извлечение аудио...")
    audio_file = extract_audio(input_file, temp_audio)

    # 2. VAD
    print("\n[2/8] VAD анализ...")
    vad_model, get_speech_timestamps, _ = load_silero_vad(device)
    vad_segments = get_vad_segments(
        audio_file, vad_model, get_speech_timestamps,
        min_speech_duration_ms=vad_min_speech_ms,
        threshold=vad_threshold,
        device=device
    )
    vad_segments = merge_close_vad_segments(vad_segments, max_gap=1.0)

    # Сохраняем VAD сегменты в лог
    debug_log.vad_segments = [
        {"idx": i, "start": s["start"], "end": s["end"], "duration": s["duration"]}
        for i, s in enumerate(vad_segments)
    ]

    del vad_model
    torch.cuda.empty_cache()

    # 3. Загрузка аудио для whisperx
    print("\n[3/8] Загрузка аудио...")
    audio = whisperx.load_audio(str(audio_file))

    # 4. Параллельная транскрипция
    print("\n[4/8] Транскрипция (параллельно оба языка)...")
    segments = transcribe_vad_segments_parallel(
        audio, vad_segments, languages,
        model, device, compute_type, batch_size
    )

    # 5. Выравнивание и диаризация
    if not skip_diarization:
        print("\n[5/8] Выравнивание и диаризация...")
        hf_token = os.getenv("HUGGINGFACE_TOKEN")
        segments = align_and_diarize(segments, audio, device, hf_token)

        # Склейка по спикерам
        print("\nСклейка сегментов...")
        segments = merge_speaker_segments(segments)
        segments = filter_unknown_speakers(segments)
    else:
        print("\n[5/8] Диаризация пропущена (--skip-diarization)")
        # Присваиваем всем сегментам SPEAKER_00
        for seg in segments:
            seg["speaker"] = "SPEAKER_00"

    segments = clean_hallucination_text(segments)
    print(f"  После обработки: {len(segments)}")

    # 6. Перевод
    if not skip_translation:
        print("\n[6/8] Перевод...")
        gemini_model = init_gemini()
        segments = translate_segments(segments, gemini_model)
    else:
        print("\n[6/8] Перевод пропущен")

    # 7. Эмоции
    if not skip_emotions:
        print("\n[7/8] Анализ эмоций...")
        analyzer = EmotionAnalyzer(device=device)
        for i, seg in enumerate(segments):
            emotion, conf = analyzer.analyze(audio_file, seg["start"], seg["end"])
            seg["emotion"] = emotion
            seg["emotion_confidence"] = conf
    else:
        print("\n[7/8] Эмоции пропущены")
        for seg in segments:
            seg["emotion"] = "neutral"

    # 8. Сохранение
    print("\n[8/8] Сохранение...")
    profiles = build_speaker_profiles(segments)

    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')

    txt_path = output_dir / f"multilang_v3_{timestamp}.txt"
    save_txt(segments, profiles, txt_path)

    json_path = output_dir / f"multilang_v3_{timestamp}.json"
    save_json(segments, profiles, input_file, json_path)

    word_path = output_dir / f"multilang_v3_{timestamp}.docx"
    create_word_report(segments, profiles, input_file, word_path)

    # Сохраняем debug лог
    debug_path = output_dir / f"debug_rejected_{timestamp}.json"
    debug_log.print_summary()
    debug_log.save(debug_path)

    # Cleanup
    if audio_file == temp_audio and temp_audio.exists():
        temp_audio.unlink()

    elapsed = datetime.now() - start_time_total
    print("\n" + "=" * 60)
    print(f"ГОТОВО! Время: {elapsed}")
    print(f"Результаты: {word_path}")
    print(f"Debug log: {debug_path}")
    print("=" * 60)

    return word_path


# === ENTRY POINT ===
if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="Мультиязычная транскрипция v3 (VAD + Parallel)")
    parser.add_argument("input", nargs="?", help="Входной файл")
    parser.add_argument("--skip-emotions", action="store_true")
    parser.add_argument("--skip-translation", action="store_true")
    parser.add_argument("--skip-diarization", action="store_true", help="Пропустить диаризацию (быстрее, но без спикеров)")
    parser.add_argument("--languages", nargs="+", default=["ru", "zh"])
    parser.add_argument("--vad-threshold", type=float, default=0.5)
    args = parser.parse_args()

    if args.input:
        INPUT_FILE = Path(args.input)
    else:
        INPUT_FILE = PROJECT_ROOT / "test_segment_0-45min.wav"

    if not INPUT_FILE.exists():
        print(f"ОШИБКА: Файл не найден: {INPUT_FILE}")
        sys.exit(1)

    result = process_multilang_v3(
        input_file=INPUT_FILE,
        output_dir=OUTPUT_DIR,
        languages=args.languages,
        skip_emotions=args.skip_emotions,
        skip_translation=args.skip_translation,
        skip_diarization=args.skip_diarization,
        vad_threshold=args.vad_threshold
    )

    print(f"\nРезультат: {result}")
