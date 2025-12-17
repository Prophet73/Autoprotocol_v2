#!/usr/bin/env python3
"""
WhisperX Full Pipeline
Транскрипция + Диаризация + Эмоции + Word отчёт
Кроссплатформенный (Windows/Linux)
"""

import torch
import os
import sys
import argparse
import numpy as np
from pathlib import Path
from datetime import datetime
from collections import defaultdict
from typing import Optional

# Patch torch.load BEFORE any imports (PyTorch 2.8+ fix)
_original_torch_load = torch.load
def _patched_torch_load(*args, **kwargs):
    kwargs['weights_only'] = False
    return _original_torch_load(*args, **kwargs)
torch.load = _patched_torch_load

import lightning_fabric.utilities.cloud_io as cloud_io
cloud_io.torch.load = _patched_torch_load

import whisperx
from whisperx.diarize import DiarizationPipeline
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from transformers import Wav2Vec2ForSequenceClassification, Wav2Vec2FeatureExtractor
import librosa
import subprocess

# === КОНСТАНТЫ ===
PROJECT_ROOT = Path(__file__).parent.resolve()
DEFAULT_OUTPUT_DIR = PROJECT_ROOT / "output"

# Модель эмоций - РУССКАЯ (Aniemore)
# Labels: anger, disgust, enthusiasm, fear, happiness, neutral, sadness
EMOTION_MODEL = "Aniemore/wav2vec2-xlsr-53-russian-emotion-recognition"
EMOTION_LABELS_RU = {
    'anger': 'Гнев',
    'disgust': 'Отвращение',
    'enthusiasm': 'Энтузиазм',
    'fear': 'Страх',
    'happiness': 'Радость',
    'neutral': 'Нейтрально',
    'sadness': 'Грусть'
}
EMOTION_EMOJI = {
    'anger': '😠',
    'disgust': '🤢',
    'enthusiasm': '🤩',
    'fear': '😨',
    'happiness': '😊',
    'neutral': '😐',
    'sadness': '😔'
}


def get_hf_token():
    """Получает HuggingFace токен из .env или переменной окружения"""
    # Сначала пробуем из переменной окружения
    token = os.getenv("HUGGINGFACE_TOKEN") or os.getenv("HF_TOKEN")
    if token:
        return token

    # Пробуем загрузить из .env
    env_file = PROJECT_ROOT / ".env"
    if env_file.exists():
        try:
            from dotenv import load_dotenv
            load_dotenv(env_file)
            token = os.getenv("HUGGINGFACE_TOKEN") or os.getenv("HF_TOKEN")
        except ImportError:
            # Читаем вручную если dotenv не установлен
            with open(env_file, 'r') as f:
                for line in f:
                    if line.startswith("HUGGINGFACE_TOKEN=") or line.startswith("HF_TOKEN="):
                        token = line.split('=', 1)[1].strip().strip('"\'')
                        break
    return token


def format_time(seconds):
    """Форматирует секунды в HH:MM:SS или MM:SS"""
    hours = int(seconds // 3600)
    mins = int((seconds % 3600) // 60)
    secs = int(seconds % 60)
    if hours > 0:
        return f"{hours:02d}:{mins:02d}:{secs:02d}"
    return f"{mins:02d}:{secs:02d}"


def extract_audio(input_file: Path, output_file: Path):
    """Извлекает аудио из видео"""
    print(f"Извлечение аудио из {input_file.name}...")
    cmd = [
        'ffmpeg', '-i', str(input_file),
        '-vn', '-acodec', 'pcm_s16le', '-ar', '16000', '-ac', '1',
        '-y', str(output_file)
    ]
    subprocess.run(cmd, capture_output=True)
    print(f"Аудио сохранено: {output_file}")


class EmotionAnalyzer:
    def __init__(self, model_name: str = EMOTION_MODEL, device: str = "cuda"):
        print(f"Загрузка модели эмоций: {model_name}")
        self.device = device
        self.feature_extractor = Wav2Vec2FeatureExtractor.from_pretrained(model_name)
        self.model = Wav2Vec2ForSequenceClassification.from_pretrained(model_name).to(device)
        self.model.eval()
        self.id2label = self.model.config.id2label

    def analyze(self, audio_path: Path, start_time: float, end_time: float):
        """Анализирует эмоцию в сегменте"""
        try:
            duration = end_time - start_time
            if duration > 30:
                samples = []
                for offset in [0, duration/2 - 5, duration - 10]:
                    audio, sr = librosa.load(str(audio_path), sr=16000, offset=start_time + offset, duration=10)
                    if len(audio) >= 1600:
                        samples.append(audio)
                if not samples:
                    return 'neutral', 0.5
                all_probs = []
                for audio in samples:
                    inputs = self.feature_extractor(audio, sampling_rate=16000, return_tensors="pt", padding=True)
                    inputs = {k: v.to(self.device) for k, v in inputs.items()}
                    with torch.no_grad():
                        outputs = self.model(**inputs)
                        probs = torch.softmax(outputs.logits, dim=-1)
                        all_probs.append(probs[0].cpu().numpy())
                avg_probs = np.mean(all_probs, axis=0)
                pred_idx = np.argmax(avg_probs)
                confidence = avg_probs[pred_idx]
            else:
                audio, sr = librosa.load(str(audio_path), sr=16000, offset=start_time, duration=duration)
                if len(audio) < 1600:
                    return 'neutral', 0.5
                inputs = self.feature_extractor(audio, sampling_rate=16000, return_tensors="pt", padding=True)
                inputs = {k: v.to(self.device) for k, v in inputs.items()}
                with torch.no_grad():
                    outputs = self.model(**inputs)
                    probs = torch.softmax(outputs.logits, dim=-1)
                    pred_idx = torch.argmax(probs, dim=-1).item()
                    confidence = probs[0][pred_idx].item()

            return self.id2label[pred_idx], confidence
        except Exception as e:
            print(f"Ошибка анализа эмоций: {e}")
            return 'neutral', 0.5


def merge_speaker_segments(segments: list) -> list:
    """Склеивает последовательные сегменты одного спикера"""
    if not segments:
        return []

    merged = []
    current = None

    for seg in segments:
        speaker = seg.get("speaker", "UNKNOWN")
        text = seg.get("text", "").strip()
        start = seg.get("start", 0)
        end = seg.get("end", 0)

        if current is None:
            current = {"speaker": speaker, "start": start, "end": end, "text": text}
        elif current["speaker"] == speaker:
            current["end"] = end
            current["text"] += " " + text
        else:
            merged.append(current)
            current = {"speaker": speaker, "start": start, "end": end, "text": text}

    if current:
        merged.append(current)

    return merged


def build_speaker_profiles(merged_segments: list) -> dict:
    """Строит профиль эмоций для каждого спикера"""
    profiles = defaultdict(lambda: {
        'emotions': [],
        'total_time': 0,
        'segment_count': 0,
        'emotion_counts': defaultdict(int)
    })

    for seg in merged_segments:
        speaker = seg['speaker']
        emotion = seg.get('emotion', 'neutral')
        duration = seg['end'] - seg['start']

        profiles[speaker]['emotions'].append(emotion)
        profiles[speaker]['total_time'] += duration
        profiles[speaker]['segment_count'] += 1
        profiles[speaker]['emotion_counts'][emotion] += 1

    for speaker, data in profiles.items():
        dominant_emotions = sorted(data['emotion_counts'].items(), key=lambda x: -x[1])
        dominant = dominant_emotions[0][0] if dominant_emotions else 'neutral'

        if dominant in ('happiness', 'enthusiasm'):
            data['interpretation'] = 'Позитивный настрой, энтузиазм'
        elif dominant == 'anger':
            data['interpretation'] = 'Напряжённость, возможно недовольство'
        elif dominant == 'sadness':
            data['interpretation'] = 'Обеспокоенность, усталость'
        elif dominant == 'fear':
            data['interpretation'] = 'Неуверенность, тревожность'
        elif dominant == 'disgust':
            data['interpretation'] = 'Негативное отношение'
        else:
            data['interpretation'] = 'Деловой, сдержанный тон'

        data['emoji_string'] = ''.join(EMOTION_EMOJI.get(e, '😐') for e in data['emotions'][:10])
        if len(data['emotions']) > 10:
            data['emoji_string'] += f'... (+{len(data["emotions"]) - 10})'

    return dict(profiles)


def create_word_report(merged_segments: list, speaker_profiles: dict, audio_file: Path, output_path: Path):
    """Создаёт полный Word отчёт"""
    doc = Document()

    title = doc.add_heading('Протокол совещания', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('Информация о записи', level=1)
    total_duration = merged_segments[-1]["end"] if merged_segments else 0

    info_table = doc.add_table(rows=5, cols=2)
    info_table.style = 'Table Grid'
    info_data = [
        ('Файл', audio_file.name),
        ('Дата обработки', datetime.now().strftime('%d.%m.%Y %H:%M')),
        ('Длительность', format_time(total_duration)),
        ('Участников', str(len(speaker_profiles))),
        ('Сегментов', str(len(merged_segments))),
    ]
    for i, (label, value) in enumerate(info_data):
        row = info_table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        row.cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()
    doc.add_heading('Профиль участников', level=1)

    profile_table = doc.add_table(rows=len(speaker_profiles) + 1, cols=4)
    profile_table.style = 'Table Grid'

    headers = ['Спикер', 'Время', 'Эмоции', 'Интерпретация']
    for i, h in enumerate(headers):
        profile_table.rows[0].cells[i].text = h
        profile_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    sorted_speakers = sorted(speaker_profiles.items(), key=lambda x: -x[1]['total_time'])
    for i, (speaker, data) in enumerate(sorted_speakers, 1):
        row = profile_table.rows[i]
        row.cells[0].text = speaker
        row.cells[1].text = format_time(data['total_time'])
        row.cells[2].text = data['emoji_string']
        row.cells[3].text = data['interpretation']

    doc.add_paragraph()
    doc.add_heading('Общая статистика эмоций', level=1)

    total_emotions = defaultdict(int)
    for seg in merged_segments:
        total_emotions[seg.get('emotion', 'neutral')] += 1

    emotion_table = doc.add_table(rows=len(total_emotions) + 1, cols=3)
    emotion_table.style = 'Table Grid'

    emotion_table.rows[0].cells[0].text = 'Эмоция'
    emotion_table.rows[0].cells[1].text = 'Количество'
    emotion_table.rows[0].cells[2].text = 'Процент'
    for cell in emotion_table.rows[0].cells:
        cell.paragraphs[0].runs[0].bold = True

    total = len(merged_segments)
    for i, (emotion, count) in enumerate(sorted(total_emotions.items(), key=lambda x: -x[1]), 1):
        row = emotion_table.rows[i]
        row.cells[0].text = f"{EMOTION_EMOJI.get(emotion, '')} {EMOTION_LABELS_RU.get(emotion, emotion)}"
        row.cells[1].text = str(count)
        row.cells[2].text = f"{count/total*100:.0f}%"

    doc.add_page_break()
    doc.add_heading('Транскрипция', level=1)

    for seg in merged_segments:
        speaker = seg["speaker"]
        start = seg["start"]
        end = seg["end"]
        text = seg["text"]
        emotion = seg.get("emotion", "neutral")

        emotion_str = f"{EMOTION_EMOJI.get(emotion, '')} {EMOTION_LABELS_RU.get(emotion, emotion)}"

        header = doc.add_paragraph()
        header_run = header.add_run(f"[{format_time(start)} - {format_time(end)}] {speaker} | {emotion_str}")
        header_run.bold = True
        header_run.font.size = Pt(11)

        text_para = doc.add_paragraph(text)
        text_para.paragraph_format.space_after = Pt(12)

    doc.save(str(output_path))
    print(f"Word отчёт сохранён: {output_path}")


def save_txt(merged_segments: list, speaker_profiles: dict, output_path: Path):
    """Сохраняет TXT отчёт"""
    with open(output_path, "w", encoding="utf-8") as f:
        f.write("=" * 60 + "\n")
        f.write("ПРОТОКОЛ СОВЕЩАНИЯ\n")
        f.write("=" * 60 + "\n\n")

        f.write("ПРОФИЛЬ УЧАСТНИКОВ:\n")
        f.write("-" * 60 + "\n")
        f.write(f"{'Спикер':<12} | {'Время':<8} | {'Эмоции':<15} | Интерпретация\n")
        f.write("-" * 60 + "\n")

        sorted_speakers = sorted(speaker_profiles.items(), key=lambda x: -x[1]['total_time'])
        for speaker, data in sorted_speakers:
            emoji_short = ''.join(EMOTION_EMOJI.get(e, '😐') for e in data['emotions'][:5])
            f.write(f"{speaker:<12} | {format_time(data['total_time']):<8} | {emoji_short:<15} | {data['interpretation']}\n")

        f.write("\n" + "=" * 60 + "\n")
        f.write("ТРАНСКРИПЦИЯ:\n")
        f.write("=" * 60 + "\n\n")

        for seg in merged_segments:
            emotion = seg.get("emotion", "neutral")
            emotion_str = f"{EMOTION_EMOJI.get(emotion, '')} {EMOTION_LABELS_RU.get(emotion, emotion)}"
            f.write(f"[{format_time(seg['start'])} - {format_time(seg['end'])}] {seg['speaker']} | {emotion_str}\n")
            f.write(f"{seg['text']}\n\n")

    print(f"TXT сохранён: {output_path}")


def process_file(
    input_file: Path,
    output_dir: Path = DEFAULT_OUTPUT_DIR,
    model: str = "large-v3",
    language: str = "ru",
    device: str = "cuda",
    compute_type: str = "float16",
    batch_size: int = 16,
    skip_emotions: bool = False,
    hf_token: str = None
) -> Path:
    """
    Основная функция обработки файла.
    Возвращает путь к Word отчёту.
    """
    start_time_total = datetime.now()

    print("=" * 60)
    print("WhisperX Full Pipeline")
    print(f"Файл: {input_file}")
    print(f"Модель: {model}")
    print(f"Устройство: {device}")
    print("=" * 60)

    output_dir.mkdir(parents=True, exist_ok=True)
    temp_audio = output_dir / "temp_audio.wav"

    # 1. Извлечение аудио
    print("\n[1/7] Извлечение аудио из видео...")
    extract_audio(input_file, temp_audio)

    # 2. Загрузка модели
    print("\n[2/7] Загрузка модели транскрипции...")
    whisper_model = whisperx.load_model(model, device, compute_type=compute_type, language=language)

    # 3. Загрузка аудио
    print("[3/7] Загрузка аудио...")
    audio = whisperx.load_audio(str(temp_audio))

    # 4. Транскрипция
    print("[4/7] Транскрипция...")
    result = whisper_model.transcribe(audio, batch_size=batch_size)
    print(f"    Найдено сегментов: {len(result['segments'])}")

    # 5. Выравнивание
    print("[5/7] Выравнивание...")
    model_a, metadata = whisperx.load_align_model(language_code=language, device=device)
    result = whisperx.align(result["segments"], model_a, metadata, audio, device=device)

    del whisper_model
    del model_a
    torch.cuda.empty_cache()

    # 6. Диаризация
    print("[6/7] Диаризация спикеров...")
    token = hf_token or get_hf_token()
    if not token:
        print("ВНИМАНИЕ: HF_TOKEN не найден, диаризация может не работать")

    diarize_model = DiarizationPipeline(use_auth_token=token, device=device)
    diarize_segments = diarize_model(audio)
    result = whisperx.assign_word_speakers(diarize_segments, result)

    del diarize_model
    torch.cuda.empty_cache()

    # Склейка
    print("\nСклейка сегментов...")
    merged = merge_speaker_segments(result["segments"])
    print(f"Исходных: {len(result['segments'])} -> После склейки: {len(merged)}")

    # 7. Анализ эмоций
    if not skip_emotions:
        print("\n[7/7] Анализ эмоций...")
        emotion_analyzer = EmotionAnalyzer(EMOTION_MODEL, device=device)

        for i, seg in enumerate(merged):
            emotion, confidence = emotion_analyzer.analyze(temp_audio, seg["start"], seg["end"])
            seg["emotion"] = emotion
            seg["emotion_confidence"] = confidence
            if (i + 1) % 10 == 0 or i == len(merged) - 1:
                print(f"    Прогресс: {i+1}/{len(merged)}")
    else:
        print("\n[7/7] Анализ эмоций пропущен")
        for seg in merged:
            seg["emotion"] = "neutral"
            seg["emotion_confidence"] = 0.0

    # Профили спикеров
    print("\nПостроение профилей спикеров...")
    speaker_profiles = build_speaker_profiles(merged)

    # Сохранение
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')

    txt_path = output_dir / f"protocol_{timestamp}.txt"
    save_txt(merged, speaker_profiles, txt_path)

    word_path = output_dir / f"protocol_{timestamp}.docx"
    create_word_report(merged, speaker_profiles, input_file, word_path)

    # Удаляем временный файл
    if temp_audio.exists():
        temp_audio.unlink()

    elapsed = datetime.now() - start_time_total
    print("\n" + "=" * 60)
    print(f"ГОТОВО! Время обработки: {elapsed}")
    print(f"Результат: {word_path}")
    print("=" * 60)

    return word_path


def build_transcription_result(
    merged_segments: list,
    speaker_profiles: dict,
    input_file: Path,
    processing_time: float,
    model: str,
    language: str
) -> "TranscriptionResult":
    """
    Конвертирует внутренние структуры в Pydantic TranscriptionResult.
    Используется для передачи данных в домены.
    """
    # Ленивый импорт чтобы не ломать CLI если pydantic не установлен
    from schemas.transcription import (
        TranscriptionResult, Segment, SpeakerProfile,
        ProcessingMetadata, Emotion
    )

    # Конвертируем сегменты
    segments = []
    for seg in merged_segments:
        emotion_str = seg.get("emotion", "neutral")
        try:
            emotion = Emotion(emotion_str)
        except ValueError:
            emotion = Emotion.NEUTRAL

        segments.append(Segment(
            start=seg["start"],
            end=seg["end"],
            text=seg["text"],
            speaker=seg["speaker"],
            emotion=emotion,
            emotion_confidence=seg.get("emotion_confidence", 0.0)
        ))

    # Конвертируем профили спикеров
    speakers = []
    for speaker_id, data in speaker_profiles.items():
        dominant_str = max(data["emotion_counts"].items(), key=lambda x: x[1])[0] if data["emotion_counts"] else "neutral"
        try:
            dominant = Emotion(dominant_str)
        except ValueError:
            dominant = Emotion.NEUTRAL

        speakers.append(SpeakerProfile(
            speaker_id=speaker_id,
            total_time=data["total_time"],
            segment_count=data["segment_count"],
            emotion_distribution=dict(data["emotion_counts"]),
            dominant_emotion=dominant,
            interpretation=data["interpretation"]
        ))

    # Сортируем по времени речи
    speakers.sort(key=lambda x: -x.total_time)

    # Общая длительность
    duration = merged_segments[-1]["end"] if merged_segments else 0.0

    # Метаданные
    metadata = ProcessingMetadata(
        source_file=input_file.name,
        duration_seconds=duration,
        processing_time_seconds=processing_time,
        model_name=model,
        language=language
    )

    return TranscriptionResult(
        segments=segments,
        speakers=speakers,
        metadata=metadata
    )


def save_json(result: "TranscriptionResult", output_path: Path):
    """Сохраняет результат в JSON"""
    import json
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(result.model_dump_json(indent=2, ensure_ascii=False))
    print(f"JSON сохранён: {output_path}")


def main():
    parser = argparse.ArgumentParser(
        description="WhisperX Full Pipeline - Транскрипция + Диаризация + Эмоции"
    )
    parser.add_argument("input", type=Path, help="Входной аудио/видео файл")
    parser.add_argument("-o", "--output", type=Path, default=DEFAULT_OUTPUT_DIR, help="Папка для результатов")
    parser.add_argument("-m", "--model", default="large-v3", help="Модель Whisper (default: large-v3)")
    parser.add_argument("-l", "--language", default="ru", help="Язык (default: ru)")
    parser.add_argument("-d", "--device", default="cuda", help="Устройство: cuda/cpu (default: cuda)")
    parser.add_argument("--compute-type", default="float16", help="Тип вычислений (default: float16)")
    parser.add_argument("--batch-size", type=int, default=16, help="Размер батча (default: 16)")
    parser.add_argument("--skip-emotions", action="store_true", help="Пропустить анализ эмоций")
    parser.add_argument("--hf-token", help="HuggingFace токен для диаризации")
    parser.add_argument("--json", action="store_true", help="Дополнительно сохранить результат в JSON")
    parser.add_argument("--open", action="store_true", help="Открыть результат после обработки (только Windows)")

    args = parser.parse_args()

    if not args.input.exists():
        print(f"Ошибка: файл не найден: {args.input}")
        sys.exit(1)

    word_path = process_file(
        input_file=args.input,
        output_dir=args.output,
        model=args.model,
        language=args.language,
        device=args.device,
        compute_type=args.compute_type,
        batch_size=args.batch_size,
        skip_emotions=args.skip_emotions,
        hf_token=args.hf_token
    )

    if args.open and sys.platform == "win32":
        os.startfile(word_path)


if __name__ == "__main__":
    main()
