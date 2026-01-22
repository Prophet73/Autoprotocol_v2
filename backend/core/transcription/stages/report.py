"""
Stage 7: Report Generation

Generates reports in multiple formats:
- DOCX (Word document)
- TXT (plain text)
- JSON (structured data)
"""
import json
import logging
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Any, Optional
from collections import defaultdict

import numpy as np
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from ..config import config

logger = logging.getLogger(__name__)


def format_time(seconds: float) -> str:
    """Format seconds to HH:MM:SS or MM:SS."""
    hours = int(seconds // 3600)
    mins = int((seconds % 3600) // 60)
    secs = int(seconds % 60)
    if hours > 0:
        return f"{hours:02d}:{mins:02d}:{secs:02d}"
    return f"{mins:02d}:{secs:02d}"


def build_speaker_profiles(segments: List[Dict]) -> Dict[str, Dict]:
    """
    Build speaker profiles from segments.

    Args:
        segments: List of transcribed segments

    Returns:
        Dict of speaker profiles
    """
    profiles = defaultdict(lambda: {
        'emotions': [],
        'total_time': 0,
        'segment_count': 0,
        'emotion_counts': defaultdict(int),
        'languages': set()
    })

    for seg in segments:
        speaker = seg.get('speaker', 'UNKNOWN')
        emotion = seg.get('emotion', 'neutral')
        duration = seg['end'] - seg['start']

        profiles[speaker]['emotions'].append(emotion)
        profiles[speaker]['total_time'] += duration
        profiles[speaker]['segment_count'] += 1
        profiles[speaker]['emotion_counts'][emotion] += 1
        profiles[speaker]['languages'].add(seg.get('language', 'ru'))

    # Add interpretation
    for data in profiles.values():
        if data['emotion_counts']:
            dominant = max(data['emotion_counts'].items(), key=lambda x: x[1])[0]
        else:
            dominant = 'neutral'

        interpretations = {
            'happiness': 'Позитивный настрой',
            'enthusiasm': 'Позитивный настрой',
            'anger': 'Напряжённость',
            'sadness': 'Обеспокоенность'
        }
        data['interpretation'] = interpretations.get(dominant, 'Деловой тон')
        data['languages'] = list(data['languages'])

    return dict(profiles)


class ReportGenerator:
    """Generates reports in multiple formats."""

    def __init__(
        self,
        output_dir: Path,
        language_flags: Optional[Dict] = None,
        emotion_labels: Optional[Dict] = None,
        emotion_emoji: Optional[Dict] = None,
    ):
        """
        Initialize report generator.

        Args:
            output_dir: Directory for output files
            language_flags: Language code to flag emoji mapping
            emotion_labels: Emotion to Russian label mapping
            emotion_emoji: Emotion to emoji mapping
        """
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

        self.language_flags = language_flags or config.languages.flags
        self.emotion_labels = emotion_labels or config.emotions.labels_ru
        self.emotion_emoji = emotion_emoji or config.emotions.emoji

    def generate_all(
        self,
        segments: List[Dict],
        input_file: Path,
        elapsed_time: float,
        prefix: str = "v4",
    ) -> Dict[str, Path]:
        """
        Generate all report formats.

        Args:
            segments: List of transcribed segments
            input_file: Original input file
            elapsed_time: Processing time in seconds
            prefix: File name prefix

        Returns:
            Dict of format to file path
        """
        profiles = build_speaker_profiles(segments)
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')

        paths = {}

        # TXT
        txt_path = self.output_dir / f"{prefix}_{timestamp}.txt"
        self.save_txt(segments, profiles, txt_path)
        paths['txt'] = txt_path

        # JSON
        json_path = self.output_dir / f"{prefix}_{timestamp}.json"
        self.save_json(segments, profiles, input_file, json_path, elapsed_time)
        paths['json'] = json_path

        # DOCX
        docx_path = self.output_dir / f"{prefix}_{timestamp}.docx"
        self.save_docx(segments, profiles, input_file, docx_path)
        paths['docx'] = docx_path

        return paths

    def save_txt(
        self,
        segments: List[Dict],
        profiles: Dict,
        output_path: Path,
    ) -> None:
        """Save plain text report."""
        with open(output_path, "w", encoding="utf-8") as f:
            f.write("=" * 60 + "\n")
            f.write("ПРОТОКОЛ СОВЕЩАНИЯ (v4: Оптимизированный)\n")
            f.write("=" * 60 + "\n\n")

            f.write("УЧАСТНИКИ:\n")
            f.write("-" * 60 + "\n")
            for speaker, data in sorted(profiles.items(), key=lambda x: -x[1]['total_time']):
                langs = ', '.join(
                    self.language_flags.get(l, l)
                    for l in data['languages']
                )
                f.write(
                    f"{speaker:<12} | {format_time(data['total_time']):<8} | "
                    f"{langs} | {data['interpretation']}\n"
                )

            f.write("\n" + "=" * 60 + "\nТРАНСКРИПЦИЯ:\n" + "=" * 60 + "\n\n")

            for seg in segments:
                emotion = seg.get("emotion", "neutral")
                lang_flag = self.language_flags.get(seg.get("language", "ru"), "")
                emotion_emoji = self.emotion_emoji.get(emotion, "")
                emotion_label = self.emotion_labels.get(emotion, emotion)

                f.write(
                    f"[{format_time(seg['start'])} - {format_time(seg['end'])}] "
                    f"{seg['speaker']} | {lang_flag} | {emotion_emoji} {emotion_label}\n"
                )

                if seg.get("original_text"):
                    f.write(f"{seg['original_text']}\n  -> {seg['text']}\n\n")
                else:
                    f.write(f"{seg['text']}\n\n")

        logger.info(f"TXT saved: {output_path}")

    def save_json(
        self,
        segments: List[Dict],
        profiles: Dict,
        input_file: Path,
        output_path: Path,
        elapsed_time: float,
    ) -> None:
        """Save JSON report."""
        def convert(obj: Any) -> Any:
            """Convert numpy types and sets for JSON serialization."""
            if isinstance(obj, (np.floating, np.integer)):
                return float(obj) if isinstance(obj, np.floating) else int(obj)
            if isinstance(obj, set):
                return list(obj)
            if isinstance(obj, dict):
                return {k: convert(v) for k, v in obj.items()}
            if isinstance(obj, list):
                return [convert(v) for v in obj]
            return obj

        result = {
            "source_file": input_file.name,
            "processed_at": datetime.now().isoformat(),
            "pipeline": "v4_optimized",
            "processing_time_seconds": elapsed_time,
            "segments_count": len(segments),
            "speakers": convert(profiles),
            "segments": convert(segments)
        }

        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(result, f, ensure_ascii=False, indent=2)

        logger.info(f"JSON saved: {output_path}")

    def save_docx(
        self,
        segments: List[Dict],
        profiles: Dict,
        input_file: Path,
        output_path: Path,
    ) -> None:
        """Save Word document report."""
        doc = Document()

        # Title
        title = doc.add_heading('Протокол совещания (v4)', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Info section
        doc.add_heading('Информация', level=1)
        total_duration = segments[-1]["end"] if segments else 0

        info = doc.add_table(rows=4, cols=2)
        info.style = 'Table Grid'
        info_data = [
            ('Файл', input_file.name),
            ('Дата', datetime.now().strftime('%d.%m.%Y %H:%M')),
            ('Длительность', format_time(total_duration)),
            ('Участников', str(len(profiles))),
        ]
        for i, (k, v) in enumerate(info_data):
            info.rows[i].cells[0].text = k
            info.rows[i].cells[1].text = v

        doc.add_paragraph()

        # Transcription
        doc.add_heading('Транскрипция', level=1)

        for seg in segments:
            emotion = seg.get("emotion", "neutral")
            lang_flag = self.language_flags.get(seg.get("language", "ru"), "")
            emotion_emoji = self.emotion_emoji.get(emotion, "")
            emotion_label = self.emotion_labels.get(emotion, emotion)

            # Header
            header = doc.add_paragraph()
            run = header.add_run(
                f"[{format_time(seg['start'])} - {format_time(seg['end'])}] "
                f"{seg['speaker']} | {lang_flag} | {emotion_emoji} {emotion_label}"
            )
            run.bold = True

            # Content
            if seg.get("original_text"):
                doc.add_paragraph(seg['original_text'])
                trans = doc.add_paragraph()
                trans.add_run(f"-> {seg['text']}").italic = True
            else:
                doc.add_paragraph(seg['text'])

        doc.save(str(output_path))
        logger.info(f"DOCX saved: {output_path}")
