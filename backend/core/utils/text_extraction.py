"""
Text extraction utilities for .docx and .txt files.
Used for direct report generation without transcription.
"""
import re
import logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional, List

logger = logging.getLogger(__name__)

# Pattern: SPEAKER_XX (our transcript speaker label, standalone paragraph)
_SPEAKER_RE = re.compile(r"^SPEAKER_\d+$")

# Pattern: [HH:MM:SS] or [HH:MM] followed by optional lang emoji + emotion emoji + text
# Examples:
#   [00:00] 🇷🇺 😐 текст...
#   [02:03] ? 😐 текст...
#   [01:20:49] 🇷🇺 😐 текст...
_SEGMENT_RE = re.compile(
    r"^\[(\d{1,2}:\d{2}(?::\d{2})?)\]\s*"  # [HH:MM] or [HH:MM:SS]
    r"(?:[^\s]*\s*)?"                         # optional lang emoji (🇷🇺, ?, etc.)
    r"(?:[😐😠😊😢😡😨😮🤔😤😑]+\s*)?"       # optional emotion emoji
    r"(.+)$"                                  # actual text
)

# Translation error lines to skip
_TRANSLATION_ERROR_RE = re.compile(r"^\s*→\s*\[Ошибка перевода:")

# Header paragraphs to skip (our docx transcript boilerplate)
_HEADERS_TO_SKIP = {
    "транскрибация совещания",
    "информация о записи",
    "участники",
    "транскрипция",
}


@dataclass
class ParsedSegment:
    """A parsed transcript segment with speaker, timestamp, and text."""
    speaker: str
    time: str  # "HH:MM" or "HH:MM:SS"
    text: str
    start_seconds: float = 0.0


@dataclass
class ParsedTranscript:
    """Result of parsing a transcript file."""
    segments: List[ParsedSegment] = field(default_factory=list)
    is_transcript: bool = False  # True if file was recognized as our transcript format


def _time_to_seconds(time_str: str) -> float:
    """Convert HH:MM or HH:MM:SS to seconds."""
    parts = time_str.split(":")
    if len(parts) == 3:
        return int(parts[0]) * 3600 + int(parts[1]) * 60 + int(parts[2])
    elif len(parts) == 2:
        return int(parts[0]) * 60 + int(parts[1])
    return 0.0


def parse_transcript(text: str) -> ParsedTranscript:
    """
    Parse text to detect if it's our transcript format and extract structured segments.

    Our transcript docx has this structure:
        SPEAKER_10                      ← speaker label (standalone paragraph)
        [00:00] 🇷🇺 😐 текст реплики   ← segment with timestamp, lang, emotion
        [02:15] 🇷🇺 😐 ещё реплика     ← same speaker, different time
        SPEAKER_11                      ← next speaker
        [02:03] 🇷🇺 😐 текст...        ← segment

    Returns:
        ParsedTranscript with segments if recognized, empty otherwise.
    """
    lines = text.split("\n")
    # Quick check: does this look like our transcript?
    speaker_count = sum(1 for line in lines if _SPEAKER_RE.match(line.strip()))
    if speaker_count < 2:
        return ParsedTranscript(is_transcript=False)

    segments = []
    current_speaker = None
    skip_headers = True

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # Skip header paragraphs at the start
        if skip_headers and line.lower() in _HEADERS_TO_SKIP:
            continue

        # Skip translation error lines
        if _TRANSLATION_ERROR_RE.match(line):
            continue

        # Check if this is a speaker label
        if _SPEAKER_RE.match(line):
            skip_headers = False  # Past headers once we see a speaker
            current_speaker = line
            continue

        # Check if this is a segment with timestamp
        seg_match = _SEGMENT_RE.match(line)
        if seg_match and current_speaker:
            skip_headers = False
            time_str = seg_match.group(1)
            text_content = seg_match.group(2).strip()
            if text_content:
                segments.append(ParsedSegment(
                    speaker=current_speaker,
                    time=time_str,
                    text=text_content,
                    start_seconds=_time_to_seconds(time_str),
                ))
            continue

        # Plain text line under a speaker (no timestamp) — append to last segment or create new
        if current_speaker and line:
            skip_headers = False
            if segments and segments[-1].speaker == current_speaker:
                segments[-1].text += " " + line
            else:
                segments.append(ParsedSegment(
                    speaker=current_speaker,
                    time="00:00",
                    text=line,
                    start_seconds=0.0,
                ))

    if len(segments) >= 2:
        return ParsedTranscript(segments=segments, is_transcript=True)

    return ParsedTranscript(is_transcript=False)


def extract_text_from_file(file_path: Path) -> Optional[str]:
    """
    Extract text content from .txt or .docx file.

    Args:
        file_path: Path to the file

    Returns:
        Extracted text or None if extraction fails
    """
    suffix = file_path.suffix.lower()

    try:
        if suffix == '.txt':
            return _extract_from_txt(file_path)
        elif suffix == '.docx':
            return _extract_from_docx(file_path)
        else:
            logger.warning(f"Unsupported file type for text extraction: {suffix}")
            return None
    except Exception as e:
        logger.error(f"Text extraction failed for {file_path}: {e}")
        return None


def _extract_from_txt(file_path: Path) -> str:
    """Extract text from .txt file."""
    # Try different encodings
    for encoding in ['utf-8', 'utf-8-sig', 'cp1251', 'latin-1']:
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                return f.read()
        except UnicodeDecodeError:
            continue

    # Fallback: read as binary and decode with errors='replace'
    with open(file_path, 'rb') as f:
        return f.read().decode('utf-8', errors='replace')


def _extract_from_docx(file_path: Path) -> str:
    """Extract text from .docx file."""
    try:
        from docx import Document
    except ImportError:
        logger.error("python-docx not installed. Run: pip install python-docx")
        raise ImportError("python-docx required for .docx extraction")

    doc = Document(file_path)

    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                paragraphs.append(" | ".join(row_text))

    return "\n\n".join(paragraphs)


def is_text_file(filename: str) -> bool:
    """Check if file is a text file (.txt or .docx)."""
    suffix = Path(filename).suffix.lower()
    return suffix in ['.txt', '.docx']
