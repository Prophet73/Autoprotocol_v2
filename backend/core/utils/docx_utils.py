"""Shared DOCX utilities for domain report generators."""
from typing import Any, Optional

from docx import Document


def add_field(
    doc: Document,
    label: str,
    value: Any,
    is_list: bool = False,
    bold_label: bool = True,
    in_paragraph: Optional[Any] = None,
):
    """Add a field to DOCX document.

    Args:
        doc: python-docx Document instance
        label: Field label text (e.g. "Тема", "Ответственный")
        value: Field value — scalar, list, or bool/int/float
        is_list: If True, render list items as bullet points
        bold_label: If True, make the label bold
        in_paragraph: If provided, append to this paragraph instead of creating new one
    """
    if not value and not isinstance(value, (int, float, bool)):
        return

    p = in_paragraph if in_paragraph else doc.add_paragraph()

    if label:
        if in_paragraph and p.text:
            p.add_run(" ")
        run = p.add_run(f"{label}: ")
        if bold_label:
            run.bold = True

    if is_list and isinstance(value, list):
        if not value:
            if label and not in_paragraph:
                p.add_run("нет данных")
            elif label and in_paragraph:
                p.add_run(" нет данных")
            else:
                doc.add_paragraph("Нет данных", style='List Bullet')
        else:
            for item_idx, item_val in enumerate(value):
                if in_paragraph and item_idx == 0 and label:
                    p.add_run(str(item_val))
                    if item_idx < len(value) - 1:
                        p.add_run("; ")
                else:
                    doc.add_paragraph(str(item_val), style='List Bullet')
    elif isinstance(value, list) and not is_list:
        p.add_run("; ".join(map(str, value)))
    else:
        p.add_run(str(value))
