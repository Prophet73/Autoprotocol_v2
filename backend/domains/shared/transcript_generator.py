"""
Shared transcript generator — creates transcript.docx from TranscriptionResult.

Parameterized to work across all domains (construction, business, dct).
No LLM required — direct conversion of pipeline output.
"""

import logging
from typing import Optional, Any
from pathlib import Path
from datetime import datetime, timezone

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from backend.core.transcription.config import config

log = logging.getLogger(__name__)

# Get emotion mappings from config
EMOTION_EMOJI = config.emotions.emoji
EMOTION_LABELS = config.emotions.labels_ru
LANGUAGE_FLAGS = config.languages.flags


def generate_transcript_docx(
    result: Any,
    output_dir: Path,
    filename: Optional[str] = None,
    meeting_type: Optional[str] = None,
    meeting_date: Optional[str] = None,
    title_prefix: str = "Стенограмма",
    type_names: Optional[dict[str, str]] = None,
    default_type_label: str = "Совещание",
) -> Path:
    """
    Generate transcript.docx from transcription result.

    Args:
        result: TranscriptionResult from pipeline
        output_dir: Directory to save the file
        filename: Optional custom filename
        meeting_type: Type of meeting (key into type_names)
        meeting_date: Date of meeting (YYYY-MM-DD format)
        title_prefix: Prefix for the document title
        type_names: Dict mapping meeting_type -> display name
        default_type_label: Fallback label when meeting_type not in type_names

    Returns:
        Path to generated file
    """
    doc = Document()

    # Resolve type display name
    type_name = (type_names or {}).get(meeting_type, default_type_label) if meeting_type else None
    doc_title = f"{title_prefix}: {type_name}" if type_name else title_prefix

    # Title
    title = doc.add_heading(doc_title, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata section
    doc.add_heading("Информация о записи", level=1)

    # Format meeting date for display
    meeting_date_formatted = None
    if meeting_date:
        try:
            parsed_date = datetime.strptime(meeting_date, "%Y-%m-%d")
            meeting_date_formatted = parsed_date.strftime("%d.%m.%Y")
        except ValueError:
            meeting_date_formatted = meeting_date

    rows_count = 5 if meeting_date_formatted else 4
    meta_table = doc.add_table(rows=rows_count, cols=2)
    meta_table.style = "Table Grid"

    meta_data = [
        ("Файл", result.metadata.source_file if hasattr(result, 'metadata') else "—"),
        ("Длительность", result.metadata.duration_formatted if hasattr(result, 'metadata') else "—"),
    ]

    if type_name:
        meta_data.append(("Тип встречи", type_name))
    else:
        if meeting_date_formatted:
            meta_data.append(("Дата встречи", meeting_date_formatted))
        else:
            meta_data.append(("Дата обработки", datetime.now(timezone.utc).strftime("%d.%m.%Y %H:%M")))

    meta_data.append(
        ("Участников", str(result.speaker_count) if hasattr(result, 'speaker_count') else "—")
    )

    if meeting_date_formatted and type_name:
        meta_data.insert(2, ("Дата встречи", meeting_date_formatted))

    for i, (label, value) in enumerate(meta_data):
        meta_table.rows[i].cells[0].text = label
        meta_table.rows[i].cells[1].text = value

    doc.add_paragraph()  # spacing

    # Participants section
    speakers_list = getattr(result, 'speakers_list', None)
    if speakers_list:
        doc.add_heading("Участники", level=1)

        sorted_speakers = sorted(
            speakers_list,
            key=lambda s: s.total_time,
            reverse=True
        )

        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"

        headers = table.rows[0].cells
        headers[0].text = "Спикер"
        headers[1].text = "Время"
        headers[2].text = "Эмоция"
        headers[3].text = "Интерпретация"

        for cell in headers:
            for run in cell.paragraphs[0].runs:
                run.bold = True

        for speaker in sorted_speakers:
            row = table.add_row().cells
            row[0].text = speaker.speaker_id
            row[1].text = speaker.total_time_formatted

            if hasattr(speaker, "dominant_emotion") and speaker.dominant_emotion:
                emotion = speaker.dominant_emotion
                row[2].text = f"{emotion.emoji} {emotion.label_ru}"
            elif hasattr(speaker, 'emotion_counts') and speaker.emotion_counts:
                dominant = max(speaker.emotion_counts.items(), key=lambda x: x[1])[0]
                emoji = EMOTION_EMOJI.get(dominant, "")
                label = EMOTION_LABELS.get(dominant, dominant)
                row[2].text = f"{emoji} {label}"
            else:
                row[2].text = "Нейтрально"

            row[3].text = getattr(speaker, "interpretation", "Деловой тон")

        doc.add_paragraph()

    # Transcript section
    doc.add_heading("Транскрипция", level=1)

    segments = getattr(result, 'segments', None)
    if segments:
        current_speaker = None
        for segment in segments:
            if segment.speaker != current_speaker:
                current_speaker = segment.speaker
                speaker_para = doc.add_paragraph()
                speaker_para.add_run(f"\n{segment.speaker}").bold = True

            p = doc.add_paragraph()

            # Timestamp
            time_run = p.add_run(f"[{segment.start_formatted}] ")
            time_run.font.size = Pt(9)

            # Language flag
            lang = getattr(segment, "language", "ru")
            lang_flag = LANGUAGE_FLAGS.get(lang, "")
            if lang_flag:
                p.add_run(f"{lang_flag} ")

            # Emotion emoji
            if hasattr(segment, "emotion") and segment.emotion:
                emotion = segment.emotion
                emoji = EMOTION_EMOJI.get(emotion, "")
                if emoji:
                    p.add_run(f"{emoji} ")

            # Text with possible translation
            original_text = getattr(segment, "original_text", None)
            if original_text:
                p.add_run(original_text)
                trans_p = doc.add_paragraph()
                trans_run = trans_p.add_run(f"  → {segment.text}")
                trans_run.font.italic = True
            else:
                p.add_run(segment.text)

    # Save document
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    if filename is None:
        timestamp = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
        filename = f"transcript_{timestamp}.docx"

    output_path = output_dir / filename
    doc.save(str(output_path))
    log.info(f"Transcript saved: {output_path}")

    return output_path
