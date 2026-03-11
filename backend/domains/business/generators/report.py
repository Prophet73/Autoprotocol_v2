"""
Business Domain Report Generator.

Generates DOCX reports for Business meeting types:
- Negotiation
- Client Meeting
- Strategic Planning
- Presentation
- Work Meeting
- Brainstorm
- Lecture
"""
import logging
from typing import Optional, Any
from pathlib import Path

from docx import Document

from backend.core.utils.docx_utils import add_field
from backend.domains.business.schemas import (
    BusinessMeetingType,
    NegotiationResult,
    ClientMeetingResult,
    StrategicPlanningResult,
    PresentationResult,
    WorkMeetingResult,
    BrainstormResult,
    LectureResult,
)

log = logging.getLogger(__name__)


# =============================================================================
# 1. Negotiation
# =============================================================================

def _generate_negotiation_docx(
    result: NegotiationResult,
    output_path: Path,
    meeting_date: Optional[str] = None
) -> Path:
    doc = Document()

    title = "Протокол переговоров"
    if meeting_date:
        title += f" от {meeting_date}"
    doc.add_heading(title, level=1)

    add_field(doc, "Цель встречи", result.meeting_goal)

    if result.parties:
        doc.add_heading("Стороны переговоров", level=2)
        for party in result.parties:
            p = doc.add_paragraph()
            p.add_run(party.party_name).bold = True
            if party.representatives:
                add_field(doc, "Представители", party.representatives, is_list=True)

    if result.key_topics:
        doc.add_heading("Ключевые темы обсуждения", level=2)
        for i, topic in enumerate(result.key_topics, 1):
            doc.add_heading(f"{i}. {topic.topic}", level=3)
            add_field(doc, "Позиции сторон", topic.positions)
            add_field(doc, "Итог", topic.result)

    if result.agreements:
        doc.add_heading("Достигнутые договорённости", level=2)
        for item in result.agreements:
            doc.add_paragraph(item, style='List Bullet')

    if result.open_questions:
        doc.add_heading("Нерешённые вопросы", level=2)
        for item in result.open_questions:
            doc.add_paragraph(item, style='List Bullet')

    if result.action_items_for_us:
        doc.add_heading("Задачи для нашей стороны", level=2)
        for item in result.action_items_for_us:
            doc.add_paragraph(item, style='List Bullet')

    if result.action_items_for_counterpart:
        doc.add_heading("Задачи для контрагента", level=2)
        for item in result.action_items_for_counterpart:
            doc.add_paragraph(item, style='List Bullet')

    if result.internal_strategic_analysis:
        doc.add_heading("Внутренний стратегический анализ", level=2)
        if result.risk_level:
            add_field(doc, "Уровень рисков", result.risk_level)
        doc.add_paragraph(result.internal_strategic_analysis)

    doc.save(str(output_path))
    log.info(f"Business negotiation DOCX saved: {output_path}")
    return output_path


# =============================================================================
# 2. Client Meeting
# =============================================================================

def _generate_client_meeting_docx(
    result: ClientMeetingResult,
    output_path: Path,
    meeting_date: Optional[str] = None
) -> Path:
    doc = Document()

    title = "Протокол встречи с клиентом"
    if meeting_date:
        title += f" от {meeting_date}"
    doc.add_heading(title, level=1)

    add_field(doc, "Цель встречи", result.meeting_goal)
    add_field(doc, "Итог встречи", result.meeting_outcome)
    add_field(doc, "Заинтересованность клиента", result.interest_level)

    if result.client_info:
        doc.add_heading("Информация о клиенте", level=2)
        add_field(doc, "Компания", result.client_info.company)
        if result.client_info.representatives:
            add_field(doc, "Представители", result.client_info.representatives, is_list=True)

    if result.client_needs:
        doc.add_heading("Потребности клиента", level=2)
        for item in result.client_needs:
            doc.add_paragraph(item, style='List Bullet')

    if result.proposed_solutions:
        doc.add_heading("Предложенные решения", level=2)
        for item in result.proposed_solutions:
            doc.add_paragraph(item, style='List Bullet')

    if result.client_feedback:
        doc.add_heading("Обратная связь клиента", level=2)
        for item in result.client_feedback:
            doc.add_paragraph(item, style='List Bullet')

    if result.agreements:
        doc.add_heading("Договорённости", level=2)
        for item in result.agreements:
            doc.add_paragraph(item, style='List Bullet')

    if result.next_steps:
        doc.add_heading("Следующие шаги", level=2)
        for i, step in enumerate(result.next_steps, 1):
            p = doc.add_paragraph(f"{i}. {step.action}", style='List Number')
            if step.responsible:
                add_field(doc, "Ответственный", step.responsible, in_paragraph=p, bold_label=False)
            if step.deadline:
                add_field(doc, "Срок", step.deadline, in_paragraph=p, bold_label=False)

    doc.save(str(output_path))
    log.info(f"Business client meeting DOCX saved: {output_path}")
    return output_path


# =============================================================================
# 3. Strategic Planning
# =============================================================================

def _generate_strategic_planning_docx(
    result: StrategicPlanningResult,
    output_path: Path,
    meeting_date: Optional[str] = None
) -> Path:
    doc = Document()

    title = "Протокол стратегической сессии"
    if meeting_date:
        title += f" от {meeting_date}"
    doc.add_heading(title, level=1)

    add_field(doc, "Тема сессии", result.session_topic)

    if result.current_situation:
        doc.add_heading("Текущая ситуация", level=2)
        doc.add_paragraph(result.current_situation)

    if result.strategic_goals:
        doc.add_heading("Стратегические цели", level=2)
        for item in result.strategic_goals:
            doc.add_paragraph(item, style='List Bullet')

    if result.initiatives:
        doc.add_heading("Инициативы и проекты", level=2)
        for i, init in enumerate(result.initiatives, 1):
            p = doc.add_paragraph(f"{i}. {init.name}", style='List Number')
            add_field(doc, "Приоритет", init.priority, in_paragraph=p, bold_label=False)
            if init.responsible:
                add_field(doc, "Ответственный", init.responsible, in_paragraph=p, bold_label=False)
            if init.timeline:
                add_field(doc, "Сроки", init.timeline, in_paragraph=p, bold_label=False)

    if result.risks:
        doc.add_heading("Риски и ограничения", level=2)
        for item in result.risks:
            doc.add_paragraph(item, style='List Bullet')

    if result.kpis:
        doc.add_heading("Метрики / KPI", level=2)
        for kpi in result.kpis:
            p = doc.add_paragraph()
            p.add_run(kpi.metric).bold = True
            if kpi.target:
                p.add_run(f" — целевое значение: {kpi.target}")

    if result.next_steps:
        doc.add_heading("Следующие шаги", level=2)
        for i, step in enumerate(result.next_steps, 1):
            p = doc.add_paragraph(f"{i}. {step.action}", style='List Number')
            if step.responsible:
                add_field(doc, "Ответственный", step.responsible, in_paragraph=p, bold_label=False)
            if step.deadline:
                add_field(doc, "Срок", step.deadline, in_paragraph=p, bold_label=False)

    doc.save(str(output_path))
    log.info(f"Business strategic planning DOCX saved: {output_path}")
    return output_path


# =============================================================================
# 4. Presentation
# =============================================================================

def _generate_presentation_docx(
    result: PresentationResult,
    output_path: Path,
    meeting_date: Optional[str] = None
) -> Path:
    doc = Document()

    title = "Итоги презентации"
    if meeting_date:
        title += f" от {meeting_date}"
    doc.add_heading(title, level=1)

    add_field(doc, "Название", result.title)
    add_field(doc, "Докладчик", result.presenter)

    if result.key_messages:
        doc.add_heading("Ключевые тезисы", level=2)
        for item in result.key_messages:
            doc.add_paragraph(item, style='List Bullet')

    if result.conclusions:
        doc.add_heading("Выводы", level=2)
        for item in result.conclusions:
            doc.add_paragraph(item, style='List Bullet')

    if result.audience_questions:
        doc.add_heading("Вопросы аудитории", level=2)
        for qa in result.audience_questions:
            doc.add_heading(f"В: {qa.question}", level=3)
            if qa.answer:
                doc.add_paragraph(f"О: {qa.answer}")

    if result.decisions:
        doc.add_heading("Принятые решения", level=2)
        for item in result.decisions:
            doc.add_paragraph(item, style='List Bullet')

    if result.next_steps:
        doc.add_heading("Следующие шаги", level=2)
        for i, step in enumerate(result.next_steps, 1):
            p = doc.add_paragraph(f"{i}. {step.action}", style='List Number')
            if step.responsible:
                add_field(doc, "Ответственный", step.responsible, in_paragraph=p, bold_label=False)
            if step.deadline:
                add_field(doc, "Срок", step.deadline, in_paragraph=p, bold_label=False)

    doc.save(str(output_path))
    log.info(f"Business presentation DOCX saved: {output_path}")
    return output_path


# =============================================================================
# 5. Work Meeting
# =============================================================================

def _generate_work_meeting_docx(
    result: WorkMeetingResult,
    output_path: Path,
    meeting_date: Optional[str] = None
) -> Path:
    doc = Document()

    title = "Протокол рабочего совещания"
    if meeting_date:
        title += f" от {meeting_date}"
    doc.add_heading(title, level=1)

    add_field(doc, "Тема", result.meeting_topic)

    if result.summary:
        doc.add_heading("Краткое саммари", level=2)
        doc.add_paragraph(result.summary)

    if result.task_statuses:
        doc.add_heading("Статусы задач", level=2)
        for i, ts in enumerate(result.task_statuses, 1):
            p = doc.add_paragraph(f"{i}. {ts.task}", style='List Number')
            add_field(doc, "Статус", ts.status, in_paragraph=p, bold_label=False)
            if ts.responsible:
                add_field(doc, "Ответственный", ts.responsible, in_paragraph=p, bold_label=False)
            if ts.comment:
                add_field(doc, "Комментарий", ts.comment, in_paragraph=p, bold_label=False)

    if result.blockers:
        doc.add_heading("Блокеры и проблемы", level=2)
        for item in result.blockers:
            doc.add_paragraph(item, style='List Bullet')

    if result.decisions:
        doc.add_heading("Принятые решения", level=2)
        for item in result.decisions:
            doc.add_paragraph(item, style='List Bullet')

    if result.action_items:
        doc.add_heading("Поручения", level=2)
        for i, step in enumerate(result.action_items, 1):
            p = doc.add_paragraph(f"{i}. {step.action}", style='List Number')
            if step.responsible:
                add_field(doc, "Ответственный", step.responsible, in_paragraph=p, bold_label=False)
            if step.deadline:
                add_field(doc, "Срок", step.deadline, in_paragraph=p, bold_label=False)

    doc.save(str(output_path))
    log.info(f"Business work meeting DOCX saved: {output_path}")
    return output_path


# =============================================================================
# 6. Brainstorm
# =============================================================================

def _generate_brainstorm_docx(
    result: BrainstormResult,
    output_path: Path,
    meeting_date: Optional[str] = None
) -> Path:
    doc = Document()

    title = "Итоги сессии мозгового штурма"
    if meeting_date:
        title += f" от {meeting_date}"
    doc.add_heading(title, level=1)

    add_field(doc, "Тема сессии", result.session_topic)
    add_field(doc, "Обсуждаемая проблема", result.main_problem)

    if result.idea_clusters:
        doc.add_heading("Сгенерированные идеи", level=2)
        for cluster in result.idea_clusters:
            doc.add_heading(f"Кластер: {cluster.cluster_name}", level=3)
            if cluster.ideas:
                for idea in cluster.ideas:
                    doc.add_paragraph(idea, style='List Bullet')

    if result.top_ideas:
        doc.add_heading("Наиболее перспективные идеи", level=2)
        for i, idea in enumerate(result.top_ideas, 1):
            p = doc.add_paragraph(f"{i}. {idea.idea_description}", style='List Number')
            if idea.potential_impact:
                add_field(doc, "Влияние", idea.potential_impact, in_paragraph=p, bold_label=False)
            if idea.implementation_complexity:
                add_field(doc, "Сложность", idea.implementation_complexity, in_paragraph=p, bold_label=False)

    if result.parked_ideas:
        doc.add_heading("Отложенные идеи (Парковка)", level=2)
        for idea in result.parked_ideas:
            doc.add_paragraph(idea, style='List Bullet')

    if result.next_steps:
        doc.add_heading("Следующие шаги", level=2)
        for i, step in enumerate(result.next_steps, 1):
            p = doc.add_paragraph(f"{i}. {step.action_item}", style='List Number')
            if step.responsible:
                add_field(doc, "Ответственный", step.responsible, in_paragraph=p, bold_label=False)
            if step.deadline:
                add_field(doc, "Срок", step.deadline, in_paragraph=p, bold_label=False)

    doc.save(str(output_path))
    log.info(f"Business brainstorm DOCX saved: {output_path}")
    return output_path


# =============================================================================
# 7. Lecture
# =============================================================================

def _generate_lecture_docx(
    result: LectureResult,
    output_path: Path,
    meeting_date: Optional[str] = None
) -> Path:
    doc = Document()

    title = "Конспект лекции/вебинара"
    if meeting_date:
        title += f" от {meeting_date}"
    doc.add_heading(title, level=1)

    add_field(doc, "Название", result.webinar_title)

    if result.presentation_part:
        doc.add_heading("Основная часть", level=2)
        for block in result.presentation_part:
            doc.add_heading(block.block_title, level=3)
            if block.time_code:
                add_field(doc, "Тайм-код", block.time_code, bold_label=False)
            if block.key_idea:
                add_field(doc, "Ключевая мысль", block.key_idea, bold_label=False)
            if block.theses:
                doc.add_paragraph("Тезисы:")
                for thesis in block.theses:
                    doc.add_paragraph(thesis, style='List Bullet')

    if result.qa_part:
        doc.add_heading("Сессия Q&A", level=2)
        for qa_item in result.qa_part:
            doc.add_heading(f"Вопрос: {qa_item.question_title}", level=3)
            if qa_item.time_code:
                add_field(doc, "Тайм-код", qa_item.time_code, bold_label=False)
            if qa_item.key_answer_idea:
                add_field(doc, "Ответ", qa_item.key_answer_idea, bold_label=False)
            if qa_item.answer_theses:
                doc.add_paragraph("Тезисы ответа:")
                for thesis in qa_item.answer_theses:
                    doc.add_paragraph(thesis, style='List Bullet')

    if result.final_summary:
        doc.add_heading("Итоговые выводы", level=2)
        for item in result.final_summary:
            doc.add_paragraph(item, style='List Bullet')

    doc.save(str(output_path))
    log.info(f"Business lecture DOCX saved: {output_path}")
    return output_path


# =============================================================================
# Dispatcher
# =============================================================================

def generate_business_report(
    meeting_type: BusinessMeetingType,
    result: Any,
    output_path: Path,
    meeting_date: Optional[str] = None
) -> Path:
    """
    Generate DOCX report based on meeting type.

    Args:
        meeting_type: Type of Business meeting
        result: Parsed result object
        output_path: Path to save the DOCX file
        meeting_date: Optional meeting date string

    Returns:
        Path to the generated DOCX file
    """
    generators = {
        BusinessMeetingType.NEGOTIATION: _generate_negotiation_docx,
        BusinessMeetingType.CLIENT_MEETING: _generate_client_meeting_docx,
        BusinessMeetingType.STRATEGIC_PLANNING: _generate_strategic_planning_docx,
        BusinessMeetingType.PRESENTATION: _generate_presentation_docx,
        BusinessMeetingType.WORK_MEETING: _generate_work_meeting_docx,
        BusinessMeetingType.BRAINSTORM: _generate_brainstorm_docx,
        BusinessMeetingType.LECTURE: _generate_lecture_docx,
    }

    generator = generators.get(meeting_type)
    if not generator:
        raise ValueError(f"Unknown Business meeting type: {meeting_type}")

    return generator(result, output_path, meeting_date)
