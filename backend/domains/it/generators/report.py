"""
IT Report Generator.

Generates AI-analyzed IT meeting reports using Gemini.
"""
import os
import logging
from pathlib import Path
from datetime import datetime
from typing import TYPE_CHECKING, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

if TYPE_CHECKING:
    from backend.core.transcription.result import TranscriptionResult

from backend.config import get_prompt
from backend.domains.it.schemas import ITReport, ITMeetingType

logger = logging.getLogger(__name__)


def generate_report(
    result: "TranscriptionResult",
    output_dir: Path,
    meeting_type: str = "standup",
    system_prompt: Optional[str] = None,
    user_prompt: Optional[str] = None,
) -> Path:
    """
    Generate AI-analyzed IT report.

    Args:
        result: Transcription result
        output_dir: Directory to save report
        meeting_type: Type of IT meeting
        system_prompt: Optional custom system prompt
        user_prompt: Optional custom user prompt

    Returns:
        Path to generated report
    """
    # Check for Gemini API key
    api_key = os.getenv("GOOGLE_API_KEY") or os.getenv("GEMINI_API_KEY")
    if not api_key:
        logger.warning("No Gemini API key found, generating basic report")
        return _generate_basic_report(result, output_dir, meeting_type)

    try:
        import google.generativeai as genai
        genai.configure(api_key=api_key)

        # Get prompts
        if not system_prompt:
            system_prompt = get_prompt(f"domains.it.{meeting_type}.system")
        if not user_prompt:
            transcript_text = result.to_plain_text()
            user_prompt = get_prompt(
                f"domains.it.{meeting_type}.user",
                transcript=transcript_text
            )

        # Call Gemini
        model = genai.GenerativeModel(
            os.getenv("GEMINI_MODEL_NAME", "gemini-2.0-flash-exp")
        )

        response = model.generate_content(
            [system_prompt, user_prompt],
            generation_config={
                "temperature": 0.3,
                "max_output_tokens": 4096,
            }
        )

        # Parse response and generate document
        return _generate_report_document(
            result, output_dir, meeting_type, response.text
        )

    except Exception as e:
        logger.error(f"Error generating IT report with Gemini: {e}")
        return _generate_basic_report(result, output_dir, meeting_type)


def _generate_basic_report(
    result: "TranscriptionResult",
    output_dir: Path,
    meeting_type: str
) -> Path:
    """Generate basic report without AI analysis."""
    doc = Document()

    meeting_type_names = {
        "standup": "Daily Standup Report",
        "planning": "Sprint Planning Report",
        "retrospective": "Retrospective Report",
        "incident_review": "Incident Review Report",
        "architecture": "Architecture Discussion Report",
        "demo": "Sprint Demo Report",
    }

    # Title
    title = doc.add_heading(meeting_type_names.get(meeting_type, "IT Meeting Report"), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Info
    doc.add_paragraph(f"Date: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    doc.add_paragraph(f"Source: {result.source_file}")
    doc.add_paragraph()

    # Note about no AI
    doc.add_paragraph(
        "Note: AI analysis unavailable. "
        "Below is basic meeting information."
    )
    doc.add_paragraph()

    # Participants
    if result.speakers:
        doc.add_heading("Participants", level=1)
        for speaker_id, profile in result.speakers.items():
            emotion = getattr(profile, 'dominant_emotion', {})
            doc.add_paragraph(
                f"- {speaker_id}: "
                f"{int(getattr(profile, 'total_time', 0))} sec, "
                f"{getattr(profile, 'segment_count', 0)} segments, "
                f"mood: {emotion.get('label_ru', 'N/A')}"
            )

    # Key moments
    doc.add_heading("Key Moments", level=1)
    doc.add_paragraph("(First segments from meeting)")

    for i, segment in enumerate(result.segments[:10]):
        text = getattr(segment, 'text', '')[:200]
        if len(getattr(segment, 'text', '')) > 200:
            text += "..."
        doc.add_paragraph(f"{i+1}. {text}")

    # Save
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_path = output_dir / f"it_report_{timestamp}.docx"
    doc.save(str(output_path))

    logger.info(f"Basic IT report saved to {output_path}")
    return output_path


def _generate_report_document(
    result: "TranscriptionResult",
    output_dir: Path,
    meeting_type: str,
    ai_analysis: str
) -> Path:
    """Generate report document with AI analysis."""
    doc = Document()

    meeting_type_names = {
        "standup": "Daily Standup Report",
        "planning": "Sprint Planning Report",
        "retrospective": "Retrospective Report",
        "incident_review": "Incident Review Report",
        "architecture": "Architecture Discussion Report",
        "demo": "Sprint Demo Report",
    }

    # Title
    title = doc.add_heading(meeting_type_names.get(meeting_type, "IT Meeting Report"), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Info
    doc.add_paragraph(f"Date: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    doc.add_paragraph(f"Source: {result.source_file}")
    doc.add_paragraph()

    # AI Analysis
    doc.add_heading("AI Analysis", level=1)
    doc.add_paragraph(ai_analysis)

    # Participants
    if result.speakers:
        doc.add_heading("Participants", level=1)
        for speaker_id, profile in result.speakers.items():
            emotion = getattr(profile, 'dominant_emotion', {})
            doc.add_paragraph(
                f"- {speaker_id}: "
                f"{int(getattr(profile, 'total_time', 0))} sec, "
                f"{getattr(profile, 'segment_count', 0)} segments, "
                f"mood: {emotion.get('label_ru', 'N/A')}"
            )

    # Save
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_path = output_dir / f"it_report_{timestamp}.docx"
    doc.save(str(output_path))

    logger.info(f"IT report with AI analysis saved to {output_path}")
    return output_path
