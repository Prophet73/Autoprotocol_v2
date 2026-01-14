"""
IT Transcript Generator.

Generates formatted transcript document for IT meetings.
"""
import logging
from pathlib import Path
from datetime import datetime
from typing import TYPE_CHECKING

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

if TYPE_CHECKING:
    from backend.core.transcription.result import TranscriptionResult

logger = logging.getLogger(__name__)


def generate_transcript(
    result: "TranscriptionResult",
    output_dir: Path,
    meeting_type: str = "standup",
) -> Path:
    """
    Generate IT-formatted transcript document.

    Args:
        result: Transcription result with segments and speakers
        output_dir: Directory to save the document
        meeting_type: Type of IT meeting

    Returns:
        Path to generated document
    """
    doc = Document()

    # Title
    title = doc.add_heading("IT Meeting Transcript", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Meeting info
    meeting_type_names = {
        "standup": "Daily Standup",
        "planning": "Sprint Planning",
        "retrospective": "Retrospective",
        "incident_review": "Incident Review",
        "architecture": "Architecture Discussion",
        "demo": "Sprint Demo",
    }

    doc.add_paragraph(f"Meeting Type: {meeting_type_names.get(meeting_type, meeting_type)}")
    doc.add_paragraph(f"Date: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    doc.add_paragraph(f"Source: {result.source_file}")
    doc.add_paragraph()

    # Participants table
    if result.speakers:
        doc.add_heading("Participants", level=1)

        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Participant"
        hdr_cells[1].text = "Time (sec)"
        hdr_cells[2].text = "Segments"
        hdr_cells[3].text = "Mood"

        for cell in hdr_cells:
            cell.paragraphs[0].runs[0].bold = True

        # Data rows
        for speaker_id, profile in result.speakers.items():
            row = table.add_row().cells
            row[0].text = speaker_id
            row[1].text = str(int(getattr(profile, 'total_time', 0)))
            row[2].text = str(getattr(profile, 'segment_count', 0))

            emotion = getattr(profile, 'dominant_emotion', {})
            row[3].text = emotion.get('label_ru', 'N/A')

        doc.add_paragraph()

    # Transcript
    doc.add_heading("Transcript", level=1)

    for segment in result.segments:
        speaker = getattr(segment, 'speaker', 'Unknown')
        start = getattr(segment, 'start', 0)
        end = getattr(segment, 'end', 0)
        text = getattr(segment, 'text', '')

        # Format time as MM:SS
        start_fmt = f"{int(start // 60):02d}:{int(start % 60):02d}"
        end_fmt = f"{int(end // 60):02d}:{int(end % 60):02d}"

        # Add paragraph with speaker info
        p = doc.add_paragraph()
        run = p.add_run(f"[{start_fmt}-{end_fmt}] {speaker}: ")
        run.bold = True

        # Add text
        p.add_run(text)

    # Save
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_path = output_dir / f"it_transcript_{timestamp}.docx"
    doc.save(str(output_path))

    logger.info(f"IT transcript saved to {output_path}")
    return output_path
