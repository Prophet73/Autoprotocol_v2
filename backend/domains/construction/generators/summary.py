"""
Summary (Конспект) generator — topic-based meeting synopsis.

Unlike BasicReport (task-focused), this groups discussion by TOPICS,
providing full context, decisions, unresolved items, and disagreements
for each topic.

Output: summary.docx
"""

import os
import logging
from pathlib import Path
from datetime import datetime, timezone

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from backend.core.transcription.models import TranscriptionResult
from backend.domains.construction.schemas import SummaryReport
from backend.core.llm.llm_utils import run_llm_call
from backend.core.llm.client import get_llm_client
from backend.domains.construction.prompts import format_participants_for_prompt, SUMMARY_SYSTEM, SUMMARY_USER


from backend.shared.config import REPORT_MODEL as _DEFAULT_REPORT_MODEL
logger = logging.getLogger(__name__)


def get_summary_report(
    result: TranscriptionResult,
    meeting_date: str = None,
    participants: list = None,
) -> SummaryReport:
    """
    Generate SummaryReport from transcription via LLM.

    Args:
        result: TranscriptionResult from pipeline
        meeting_date: Optional meeting date (YYYY-MM-DD format)
        participants: Optional list of participants grouped by organization

    Returns:
        SummaryReport with topic-based meeting synopsis
    """
    from backend.core.llm.llm_utils import sanitize_transcript_for_llm
    transcript_text = sanitize_transcript_for_llm(result.to_plain_text())

    if not os.getenv("GOOGLE_API_KEY"):
        logger.warning("GOOGLE_API_KEY not set, returning empty SummaryReport")
        return SummaryReport(
            meeting_summary="Транскрипция обработана без LLM анализа",
            topics=[],
            key_takeaways=[],
        )

    if meeting_date:
        try:
            parsed_date = datetime.strptime(meeting_date, "%Y-%m-%d")
            meeting_date_formatted = parsed_date.strftime("%d.%m.%Y")
        except ValueError:
            meeting_date_formatted = meeting_date
    else:
        meeting_date_formatted = datetime.now(timezone.utc).strftime("%d.%m.%Y")

    return _extract_summary_via_llm(transcript_text, meeting_date_formatted, participants)


def _extract_summary_via_llm(transcript_text: str, meeting_date: str, participants: list = None) -> SummaryReport:
    """Extract SummaryReport from transcript using LLM."""

    # Format participants for prompt
    participants_text = format_participants_for_prompt(participants)
    participants_block = (
        f"\nУчастники совещания:\n{participants_text}\n"
        if participants_text else ""
    )

    full_prompt = SUMMARY_USER.format(
        transcript=transcript_text,
        meeting_date=meeting_date,
        participants_info=participants_block,
    )

    def _make_call(model: str):
        return lambda: get_llm_client().generate_content(
            model=model,
            contents=full_prompt,
            system_instruction=SUMMARY_SYSTEM,
            response_mime_type="application/json",
            response_schema=SummaryReport,
        )

    from backend.admin.settings.service import get_setting_value
    report_model = get_setting_value("gemini_report_model", _DEFAULT_REPORT_MODEL)

    try:
        response = run_llm_call(
            _make_call(report_model),
            model_name=report_model,
            make_call=_make_call,
        )

        report = SummaryReport.model_validate_json(response.text)
        logger.info(f"SummaryReport generated: {len(report.topics)} topics")
        return report

    except Exception as e:
        logger.warning("LLM SummaryReport extraction failed: %s", e)
        raise


def generate_summary(
    result: TranscriptionResult,
    output_dir: Path,
    meeting_date: str = None,
    participants: list = None,
) -> Path:
    """
    Generate summary.docx — topic-based meeting synopsis.

    Args:
        result: TranscriptionResult from pipeline
        output_dir: Directory to save the file
        meeting_date: Optional meeting date (YYYY-MM-DD)
        participants: Optional list of participants

    Returns:
        Path to generated summary.docx
    """
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    # Get structured data from LLM
    summary = get_summary_report(result, meeting_date, participants)

    # Build Word document
    doc = Document()

    # Title
    title = doc.add_heading("Конспект совещания", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Date
    if meeting_date:
        try:
            parsed = datetime.strptime(meeting_date, "%Y-%m-%d")
            date_str = parsed.strftime("%d.%m.%Y")
        except ValueError:
            date_str = meeting_date
        date_para = doc.add_paragraph(f"Дата: {date_str}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Meeting summary
    doc.add_heading("Общее резюме", level=2)
    doc.add_paragraph(summary.meeting_summary)

    # Topics
    for i, topic in enumerate(summary.topics, 1):
        doc.add_heading(f"{i}. {topic.title}", level=2)

        if topic.time_codes:
            tc_para = doc.add_paragraph()
            tc_run = tc_para.add_run(f"Тайм-коды: {', '.join(topic.time_codes)}")
            tc_run.italic = True
            tc_run.font.size = 95000  # ~7.5pt in EMU

        # Context
        if topic.context:
            doc.add_heading("Контекст", level=3)
            doc.add_paragraph(topic.context)

        # Discussion
        if topic.discussion:
            doc.add_heading("Обсуждение", level=3)
            doc.add_paragraph(topic.discussion)

        # Decisions
        if topic.decisions:
            doc.add_heading("Решения", level=3)
            for decision in topic.decisions:
                doc.add_paragraph(decision, style="List Bullet")

        # Unresolved
        if topic.unresolved:
            doc.add_heading("Нерешённые вопросы", level=3)
            for item in topic.unresolved:
                doc.add_paragraph(item, style="List Bullet")

        # Disagreements
        if topic.disagreements:
            doc.add_heading("Разногласия", level=3)
            for dis in topic.disagreements:
                para = doc.add_paragraph()
                parties_run = para.add_run(f"{dis.parties}: ")
                parties_run.bold = True
                para.add_run(dis.essence)
                if dis.outcome:
                    para.add_run(f" → {dis.outcome}")

        # Responsible
        if topic.responsible:
            resp_para = doc.add_paragraph()
            resp_run = resp_para.add_run("Ответственные: ")
            resp_run.bold = True
            resp_para.add_run(", ".join(topic.responsible))

    # Key takeaways
    if summary.key_takeaways:
        doc.add_heading("Главные выводы", level=2)
        for takeaway in summary.key_takeaways:
            doc.add_paragraph(takeaway, style="List Bullet")

    # Save
    timestamp = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
    filename = f"summary_{timestamp}.docx"
    output_path = output_dir / filename
    doc.save(str(output_path))

    logger.info(f"Summary generated: {output_path} ({len(summary.topics)} topics)")
    return output_path
