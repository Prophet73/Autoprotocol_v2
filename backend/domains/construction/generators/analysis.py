"""
Analysis generator - creates analysis.docx from TranscriptionResult via LLM.
Deep AI analysis for managers: status, indicators, challenges, recommendations.
"""

import os
import json
import logging
from pathlib import Path
from datetime import datetime
from typing import Optional

from google import genai

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from backend.core.transcription.models import TranscriptionResult
from backend.domains.construction.schemas import AIAnalysis, OverallStatus, Atmosphere
from backend.domains.construction.prompts import CONSTRUCTION_PROMPTS
from backend.domains.construction.generators.llm_utils import run_llm_call


# Model for reports (pro for quality)
REPORT_MODEL = os.getenv("GEMINI_REPORT_MODEL", "gemini-2.5-pro")
logger = logging.getLogger(__name__)

# Status colors and emojis
STATUS_CONFIG = {
    "stable": {"emoji": "🟢", "color": "00AA00", "label": "Стабильный"},
    "attention": {"emoji": "🟡", "color": "FFAA00", "label": "Требует внимания"},
    "critical": {"emoji": "🔴", "color": "FF0000", "label": "Критический"},
}

INDICATOR_EMOJI = {
    "ok": "✅",
    "risk": "⚠️",
    "critical": "🔴",
}


def generate_analysis(
    result: TranscriptionResult,
    output_dir: Path,
    filename: str = None,
) -> Path:
    """
    Generate analysis.docx from transcription via LLM.

    Args:
        result: TranscriptionResult from pipeline
        output_dir: Directory to save the file
        filename: Optional custom filename

    Returns:
        Path to generated file
    """
    # Get transcript text
    transcript_text = result.to_plain_text()

    # Call LLM for deep analysis
    if os.getenv("GOOGLE_API_KEY"):
        ai_analysis = _get_ai_analysis(transcript_text)
    else:
        ai_analysis = AIAnalysis(
            overall_status=OverallStatus.ATTENTION,
            executive_summary="GOOGLE_API_KEY не настроен. Анализ недоступен.",
            indicators=[],
            challenges=[],
            achievements=[],
            atmosphere=Atmosphere.WORKING,
            atmosphere_comment="",
        )

    # Create Word document
    doc = Document()

    # Title
    title = doc.add_heading("ИИ Анализ совещания", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Overall status
    status_key = ai_analysis.overall_status.value if hasattr(ai_analysis.overall_status, "value") else str(ai_analysis.overall_status)
    status_cfg = STATUS_CONFIG.get(status_key, STATUS_CONFIG["attention"])

    status_para = doc.add_paragraph()
    status_para.add_run("Общий статус: ").bold = True
    status_para.add_run(f"{status_cfg['emoji']} {status_cfg['label']}")

    doc.add_paragraph()

    # Metadata
    doc.add_heading("Информация о записи", level=1)
    meta_para = doc.add_paragraph()
    meta_para.add_run(f"Файл: {result.metadata.source_file}\n")
    meta_para.add_run(f"Длительность: {result.metadata.duration_formatted}\n")
    meta_para.add_run(f"Дата анализа: {datetime.now().strftime('%d.%m.%Y %H:%M')}\n")
    meta_para.add_run(f"Участников: {result.speaker_count}")

    # Executive Summary
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(ai_analysis.executive_summary)

    # Key Indicators
    if ai_analysis.indicators:
        doc.add_heading("Ключевые показатели", level=1)

        ind_table = doc.add_table(rows=len(ai_analysis.indicators) + 1, cols=3)
        ind_table.style = "Table Grid"

        # Header
        headers = ["Показатель", "Статус", "Комментарий"]
        for col, header in enumerate(headers):
            cell = ind_table.rows[0].cells[col]
            cell.text = header
            cell.paragraphs[0].runs[0].bold = True

        # Data
        for row_idx, indicator in enumerate(ai_analysis.indicators, 1):
            ind_table.rows[row_idx].cells[0].text = indicator.name
            status_emoji = INDICATOR_EMOJI.get(indicator.status, "⚪")
            ind_table.rows[row_idx].cells[1].text = f"{status_emoji} {indicator.status}"
            ind_table.rows[row_idx].cells[2].text = indicator.comment

        doc.add_paragraph()

    # Challenges and Recommendations
    if ai_analysis.challenges:
        doc.add_heading("Проблемы и рекомендации", level=1)

        for i, challenge in enumerate(ai_analysis.challenges, 1):
            # Problem
            p_problem = doc.add_paragraph()
            p_problem.add_run(f"{i}. Проблема: ").bold = True
            p_problem.add_run(challenge.problem)

            # Recommendation
            p_rec = doc.add_paragraph()
            p_rec.add_run("   → Рекомендация: ").italic = True
            p_rec.add_run(challenge.recommendation)

            if challenge.responsible:
                p_resp = doc.add_paragraph()
                p_resp.add_run(f"   Ответственный: {challenge.responsible}")

            doc.add_paragraph()  # spacing

    # Achievements
    if ai_analysis.achievements:
        doc.add_heading("Достижения и позитивные моменты", level=1)

        for achievement in ai_analysis.achievements:
            doc.add_paragraph(f"✅ {achievement}", style="List Bullet")

        doc.add_paragraph()

    # Atmosphere
    doc.add_heading("Атмосфера совещания", level=1)

    atm_value = ai_analysis.atmosphere.value if hasattr(ai_analysis.atmosphere, "value") else str(ai_analysis.atmosphere)
    atm_labels = {
        "calm": "Спокойное",
        "working": "Рабочее",
        "tense": "Напряжённое",
        "conflict": "Конфликтное",
    }
    atm_label = atm_labels.get(atm_value, atm_value)

    atm_para = doc.add_paragraph()
    atm_para.add_run(f"Уровень: {atm_label}\n").bold = True
    if ai_analysis.atmosphere_comment:
        atm_para.add_run(ai_analysis.atmosphere_comment)

    # Save document
    output_dir.mkdir(parents=True, exist_ok=True)

    if filename is None:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"analysis_{timestamp}.docx"

    output_path = output_dir / filename
    doc.save(str(output_path))

    return output_path


def _get_ai_analysis(transcript_text: str) -> AIAnalysis:
    """Get AI analysis from LLM."""
    client = genai.Client()

    system_prompt = CONSTRUCTION_PROMPTS.get("system", "")
    user_prompt = """
Проанализируй стенограмму совещания для руководителя.

1. Определи общий статус:
   - stable: серьёзных отклонений нет
   - attention: есть риски, требуется контроль
   - critical: угроза срыва (только при реальных проблемах!)

2. Напиши executive summary (2-3 предложения для руководителя)

3. Оцени 3-5 ключевых показателей:
   - Сроки (ok/risk/critical)
   - Бюджет (ok/risk/critical)
   - Ресурсы (ok/risk/critical)
   - Качество (ok/risk/critical)
   - Безопасность (ok/risk/critical)

4. Выдели 2-4 главные проблемы с рекомендациями

5. Найди 1-3 достижения или позитивных момента

6. Оцени атмосферу совещания:
   - calm: спокойное
   - working: рабочее напряжение
   - tense: напряжённо, споры
   - conflict: конфликт

Стенограмма:
---
{transcript}
---

Ответь в формате JSON.
""".format(transcript=transcript_text[:15000])

    try:
        response = run_llm_call(
            lambda: client.models.generate_content(
                model=REPORT_MODEL,
                contents=[system_prompt, user_prompt] if system_prompt else user_prompt,
                config={
                    "response_mime_type": "application/json",
                    "response_schema": AIAnalysis.model_json_schema(),
                },
            )
        )

        analysis_data = json.loads(response.text)
        return AIAnalysis.model_validate(analysis_data)

    except Exception as e:
        logger.warning("LLM analysis generation failed: %s", e)
        return AIAnalysis(
            overall_status=OverallStatus.ATTENTION,
            executive_summary=f"Ошибка генерации анализа: {e}",
            indicators=[],
            challenges=[],
            achievements=[],
            atmosphere=Atmosphere.WORKING,
            atmosphere_comment="",
        )
