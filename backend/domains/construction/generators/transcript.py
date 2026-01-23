"""
Transcript generator - creates transcript.docx from TranscriptionResult.
No LLM required - direct conversion of pipeline output.
"""

from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from backend.core.transcription.models import TranscriptionResult
from backend.core.transcription.config import config


# Get emotion mappings from config
EMOTION_EMOJI = config.emotions.emoji
EMOTION_LABELS = config.emotions.labels_ru
LANGUAGE_FLAGS = config.languages.flags


def generate_transcript(
    result: TranscriptionResult,
    output_dir: Path,
    filename: str = None,
    meeting_type: str = None,
    meeting_date: str = None,
) -> Path:
    """
    Generate transcript.docx from transcription result.

    Args:
        result: TranscriptionResult from pipeline
        output_dir: Directory to save the file
        filename: Optional custom filename
        meeting_type: Type of meeting (for title)
        meeting_date: Date of meeting (YYYY-MM-DD format)

    Returns:
        Path to generated file
    """
    doc = Document()

    # Title
    title = doc.add_heading("Транскрибация совещания", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata section
    doc.add_heading("Информация о записи", level=1)

    # Format meeting date for display
    meeting_date_formatted = None
    if meeting_date:
        try:
            from datetime import datetime as dt
            parsed_date = dt.strptime(meeting_date, "%Y-%m-%d")
            meeting_date_formatted = parsed_date.strftime("%d.%m.%Y")
        except ValueError:
            meeting_date_formatted = meeting_date

    meta_table = doc.add_table(rows=5 if meeting_date_formatted else 4, cols=2)
    meta_table.style = "Table Grid"

    meta_data = [
        ("Файл", result.metadata.source_file),
        ("Длительность", result.metadata.duration_formatted),
        ("Дата встречи", meeting_date_formatted) if meeting_date_formatted else ("Дата обработки", datetime.now().strftime("%d.%m.%Y %H:%M")),
        ("Участников", str(result.speaker_count)),
    ]
    if meeting_date_formatted:
        meta_data.insert(3, ("Дата обработки", datetime.now().strftime("%d.%m.%Y %H:%M")))

    for i, (label, value) in enumerate(meta_data):
        meta_table.rows[i].cells[0].text = label
        meta_table.rows[i].cells[1].text = value

    doc.add_paragraph()  # spacing

    # Participants section with table
    if result.speakers_list:
        doc.add_heading("Участники", level=1)

        # Sort by total time descending
        sorted_speakers = sorted(
            result.speakers_list,
            key=lambda s: s.total_time,
            reverse=True
        )

        # Create table
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"

        # Headers
        headers = table.rows[0].cells
        headers[0].text = "Спикер"
        headers[1].text = "Время"
        headers[2].text = "Эмоция"
        headers[3].text = "Интерпретация"

        for cell in headers:
            for run in cell.paragraphs[0].runs:
                run.bold = True

        # Data rows
        for speaker in sorted_speakers:
            row = table.add_row().cells
            row[0].text = speaker.speaker_id

            row[1].text = speaker.total_time_formatted

            # Emotion with emoji
            if hasattr(speaker, "dominant_emotion") and speaker.dominant_emotion:
                emotion = speaker.dominant_emotion
                row[2].text = f"{emotion.emoji} {emotion.label_ru}"
            else:
                # Get dominant from emotion_counts
                if speaker.emotion_counts:
                    dominant = max(speaker.emotion_counts.items(), key=lambda x: x[1])[0]
                    emoji = EMOTION_EMOJI.get(dominant, "")
                    label = EMOTION_LABELS.get(dominant, dominant)
                    row[2].text = f"{emoji} {label}"
                else:
                    row[2].text = "😐 Нейтрально"

            row[3].text = getattr(speaker, "interpretation", "Деловой тон")

        doc.add_paragraph()  # spacing

    # Transcript section
    doc.add_heading("Транскрипция", level=1)

    current_speaker = None
    for segment in result.segments:
        # Speaker header if changed
        if segment.speaker != current_speaker:
            current_speaker = segment.speaker
            speaker_para = doc.add_paragraph()
            speaker_para.add_run(f"\n{segment.speaker}").bold = True

        # Segment text with timestamp, language flag, and emotion
        p = doc.add_paragraph()

        # Timestamp
        time_run = p.add_run(f"[{segment.start_formatted}] ")
        time_run.font.size = Pt(9)

        # Language flag
        lang = getattr(segment, "language", "ru")
        lang_flag = LANGUAGE_FLAGS.get(lang, "")
        if lang_flag:
            p.add_run(f"{lang_flag} ")

        # Emotion emoji
        if hasattr(segment, "emotion") and segment.emotion:
            emotion = segment.emotion
            emoji = EMOTION_EMOJI.get(emotion, "")
            if emoji:
                p.add_run(f"{emoji} ")

        # Check if there's original text (translation case)
        original_text = getattr(segment, "original_text", None)
        if original_text:
            # Show original first
            p.add_run(original_text)
            # Then translation on new line
            trans_p = doc.add_paragraph()
            trans_run = trans_p.add_run(f"  → {segment.text}")
            trans_run.font.italic = True
        else:
            # Just the text
            p.add_run(segment.text)

    # Save document
    output_dir.mkdir(parents=True, exist_ok=True)

    if filename is None:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"transcript_{timestamp}.docx"

    output_path = output_dir / filename
    doc.save(str(output_path))

    return output_path
