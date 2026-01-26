"""
Report generator - creates report.docx from BasicReport.
Receives pre-generated BasicReport (from shared LLM call) and formats as Word document.
"""

import logging
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from backend.core.transcription.models import TranscriptionResult
from backend.domains.construction.schemas import BasicReport, TaskCategory, TaskPriority


logger = logging.getLogger(__name__)


def generate_report(
    result: TranscriptionResult,
    output_dir: Path,
    basic_report: BasicReport,
    filename: str = None,
    meeting_type: str = None,
    meeting_date: str = None,
    participants: list = None,
    include_emotions: bool = True,
) -> Path:
    """
    Generate report.docx from BasicReport.

    Args:
        result: TranscriptionResult from pipeline (for metadata)
        output_dir: Directory to save the file
        basic_report: Pre-generated BasicReport from shared LLM call
        filename: Optional custom filename
        meeting_type: Type of meeting (overrides BasicReport)
        meeting_date: Date of meeting (YYYY-MM-DD format)
        participants: Optional list of participants
        include_emotions: Whether to include emotions section (default True)

    Returns:
        Path to generated file
    """
    # Override meeting_type if provided
    effective_meeting_type = meeting_type or basic_report.meeting_type

    # Format meeting date for display
    meeting_date_formatted = None
    if meeting_date:
        try:
            parsed_date = datetime.strptime(meeting_date, "%Y-%m-%d")
            meeting_date_formatted = parsed_date.strftime("%d.%m.%Y")
        except ValueError:
            meeting_date_formatted = meeting_date

    # Create Word document
    doc = Document()

    # Title
    title = doc.add_heading("Протокол совещания", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata
    doc.add_heading("Информация о записи", level=1)

    meta_rows = 5 if meeting_date_formatted else 4
    meta_table = doc.add_table(rows=meta_rows, cols=2)
    meta_table.style = "Table Grid"

    meta_data = [
        ("Файл", result.metadata.source_file),
        ("Длительность", result.metadata.duration_formatted),
    ]
    if meeting_date_formatted:
        meta_data.append(("Дата встречи", meeting_date_formatted))
    meta_data.append(("Дата обработки", datetime.now().strftime("%d.%m.%Y %H:%M")))
    meta_data.append(("Тип совещания", effective_meeting_type))

    for i, (label, value) in enumerate(meta_data):
        meta_table.rows[i].cells[0].text = label
        meta_table.rows[i].cells[1].text = value
        meta_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # Participants
    if participants:
        doc.add_heading("Участники совещания", level=1)
        _add_participants_table(doc, participants)
        doc.add_paragraph()

    # Summary
    doc.add_heading("Краткое содержание", level=1)
    doc.add_paragraph(basic_report.meeting_summary)

    # Participants with emotions (only if requested and speakers have emotion data)
    if include_emotions and result.speakers_list:
        # Check if any speaker has emotion data
        has_emotions = any(
            hasattr(s, "dominant_emotion") and s.dominant_emotion
            for s in result.speakers_list
        )
        if has_emotions:
            doc.add_heading("Участники и эмоции", level=1)

            speaker_table = doc.add_table(rows=len(result.speakers_list) + 1, cols=3)
            speaker_table.style = "Table Grid"

            # Header
            headers = ["Спикер", "Время", "Доминирующая эмоция"]
            for col, header in enumerate(headers):
                cell = speaker_table.rows[0].cells[col]
                cell.text = header
                cell.paragraphs[0].runs[0].bold = True

            # Data rows
            for row_idx, speaker in enumerate(result.speakers_list, 1):
                speaker_table.rows[row_idx].cells[0].text = speaker.speaker_id
                speaker_table.rows[row_idx].cells[1].text = speaker.total_time_formatted

                if hasattr(speaker, "dominant_emotion") and speaker.dominant_emotion:
                    emotion = speaker.dominant_emotion
                    emotion_text = f"{emotion.label_ru} {emotion.emoji}"
                else:
                    emotion_text = "—"
                speaker_table.rows[row_idx].cells[2].text = emotion_text

            doc.add_paragraph()

    # Tasks
    if basic_report.tasks:
        doc.add_heading("Задачи", level=1)

        task_table = doc.add_table(rows=len(basic_report.tasks) + 1, cols=5)
        task_table.style = "Table Grid"

        # Header
        task_headers = ["№", "Категория", "Задача", "Ответственный", "Срок"]
        for col, header in enumerate(task_headers):
            cell = task_table.rows[0].cells[col]
            cell.text = header
            cell.paragraphs[0].runs[0].bold = True

        # Task rows
        for row_idx, task in enumerate(basic_report.tasks, 1):
            task_table.rows[row_idx].cells[0].text = str(row_idx)
            category = task.category.value if isinstance(task.category, TaskCategory) else str(task.category)
            task_table.rows[row_idx].cells[1].text = category
            task_table.rows[row_idx].cells[2].text = task.description
            task_table.rows[row_idx].cells[3].text = task.responsible or "—"
            task_table.rows[row_idx].cells[4].text = task.deadline or "—"

        doc.add_paragraph()

    # Expert analysis
    doc.add_heading("Экспертный анализ", level=1)
    doc.add_paragraph(basic_report.expert_analysis)

    # Save document
    output_dir.mkdir(parents=True, exist_ok=True)

    if filename is None:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"report_{timestamp}.docx"

    output_path = output_dir / filename
    doc.save(str(output_path))

    return output_path


def _add_participants_table(doc: Document, participants: list) -> None:
    """Add participants table to Word document."""
    if not participants:
        return

    # Role labels mapping
    role_labels = {
        "GENERAL": "Генподрядчик",
        "CUSTOMER": "Заказчик",
        "ЗАКАЗЧИК": "Заказчик",
        "TECHNICAL_CUSTOMER": "Технический заказчик",
        "ТЕХНИЧЕСКИЙ ЗАКАЗЧИК": "Технический заказчик",
        "DESIGNER": "Проектировщик",
        "SUBCONTRACTOR": "Субподрядчик",
        "SUPPLIER": "Поставщик",
        "INVESTOR": "Инвестор",
    }

    # Count total people
    total_people = sum(len(p.get("people", [])) for p in participants)
    if total_people == 0:
        return

    # Create table: header + data rows
    table = doc.add_table(rows=total_people + 1, cols=4)
    table.style = "Table Grid"

    # Header
    headers = ["Организация", "Роль", "ФИО", "Должность"]
    for col, header in enumerate(headers):
        cell = table.rows[0].cells[col]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True

    # Data rows
    row_idx = 1
    for org in participants:
        org_name = org.get("organization", "")
        role_raw = org.get("role", "")
        role = role_labels.get(role_raw.upper(), role_raw)

        for person in org.get("people", []):
            table.rows[row_idx].cells[0].text = org_name
            table.rows[row_idx].cells[1].text = role
            # Handle both formats: string "Name (Position)" or dict {"name": ..., "position": ...}
            if isinstance(person, str):
                # Parse "Иванов Иван Петрови (Директор)" format
                if "(" in person and person.endswith(")"):
                    name_part = person[:person.rfind("(")].strip()
                    position_part = person[person.rfind("(")+1:-1].strip()
                else:
                    name_part = person
                    position_part = ""
                table.rows[row_idx].cells[2].text = name_part
                table.rows[row_idx].cells[3].text = position_part
            else:
                table.rows[row_idx].cells[2].text = person.get("name", "")
                table.rows[row_idx].cells[3].text = person.get("position", "")
            row_idx += 1
