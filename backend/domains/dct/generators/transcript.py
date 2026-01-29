"""
DCT Domain Transcript Generator.

Generates transcript documents for DCT meetings.
"""
import logging
from typing import Optional, Any
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from backend.core.transcription.config import config

log = logging.getLogger(__name__)

# Get emotion mappings from config
EMOTION_EMOJI = config.emotions.emoji
EMOTION_LABELS = config.emotions.labels_ru
LANGUAGE_FLAGS = config.languages.flags


def generate_dct_transcript(
    result: Any,
    output_dir: Path,
    filename: str = None,
    meeting_type: str = "brainstorm",
    meeting_date: Optional[str] = None
) -> Path:
    """
    Generate transcript DOCX for DCT meeting.

    Args:
        result: TranscriptionResult from pipeline
        output_dir: Directory to save the DOCX file
        filename: Optional custom filename
        meeting_type: Type of meeting (brainstorm, production, negotiation, lecture)
        meeting_date: Optional meeting date string (YYYY-MM-DD)

    Returns:
        Path to the generated DOCX file
    """
    doc = Document()

    # Meeting type display names
    type_names = {
        "brainstorm": "Мозговой штурм",
        "production": "Производственное совещание",
        "negotiation": "Переговоры с контрагентом",
        "lecture": "Лекция/Вебинар",
    }

    type_name = type_names.get(meeting_type, "Совещание ДЦТ")

    # Title
    title = doc.add_heading(f"Стенограмма: {type_name}", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata section
    doc.add_heading("Информация о записи", level=1)

    # Format meeting date for display
    meeting_date_formatted = None
    if meeting_date:
        try:
            parsed_date = datetime.strptime(meeting_date, "%Y-%m-%d")
            meeting_date_formatted = parsed_date.strftime("%d.%m.%Y")
        except ValueError:
            meeting_date_formatted = meeting_date

    rows_count = 5 if meeting_date_formatted else 4
    meta_table = doc.add_table(rows=rows_count, cols=2)
    meta_table.style = "Table Grid"

    meta_data = [
        ("Файл", result.metadata.source_file if hasattr(result, 'metadata') else "—"),
        ("Длительность", result.metadata.duration_formatted if hasattr(result, 'metadata') else "—"),
        ("Тип встречи", type_name),
        ("Участников", str(result.speaker_count) if hasattr(result, 'speaker_count') else "—"),
    ]

    if meeting_date_formatted:
        meta_data.insert(2, ("Дата встречи", meeting_date_formatted))

    for i, (label, value) in enumerate(meta_data):
        meta_table.rows[i].cells[0].text = label
        meta_table.rows[i].cells[1].text = value

    doc.add_paragraph()  # spacing

    # Participants section
    if hasattr(result, 'speakers_list') and result.speakers_list:
        doc.add_heading("Участники", level=1)

        sorted_speakers = sorted(
            result.speakers_list,
            key=lambda s: s.total_time,
            reverse=True
        )

        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"

        headers = table.rows[0].cells
        headers[0].text = "Спикер"
        headers[1].text = "Время"
        headers[2].text = "Эмоция"
        headers[3].text = "Интерпретация"

        for cell in headers:
            for run in cell.paragraphs[0].runs:
                run.bold = True

        for speaker in sorted_speakers:
            row = table.add_row().cells
            row[0].text = speaker.speaker_id
            row[1].text = speaker.total_time_formatted

            if hasattr(speaker, "dominant_emotion") and speaker.dominant_emotion:
                emotion = speaker.dominant_emotion
                row[2].text = f"{emotion.emoji} {emotion.label_ru}"
            elif hasattr(speaker, 'emotion_counts') and speaker.emotion_counts:
                dominant = max(speaker.emotion_counts.items(), key=lambda x: x[1])[0]
                emoji = EMOTION_EMOJI.get(dominant, "")
                label = EMOTION_LABELS.get(dominant, dominant)
                row[2].text = f"{emoji} {label}"
            else:
                row[2].text = "😐 Нейтрально"

            row[3].text = getattr(speaker, "interpretation", "Деловой тон")

        doc.add_paragraph()

    # Transcript section
    doc.add_heading("Транскрипция", level=1)

    if hasattr(result, 'segments'):
        current_speaker = None
        for segment in result.segments:
            if segment.speaker != current_speaker:
                current_speaker = segment.speaker
                speaker_para = doc.add_paragraph()
                speaker_para.add_run(f"\n{segment.speaker}").bold = True

            p = doc.add_paragraph()

            # Timestamp
            time_run = p.add_run(f"[{segment.start_formatted}] ")
            time_run.font.size = Pt(9)

            # Language flag
            lang = getattr(segment, "language", "ru")
            lang_flag = LANGUAGE_FLAGS.get(lang, "")
            if lang_flag:
                p.add_run(f"{lang_flag} ")

            # Emotion emoji
            if hasattr(segment, "emotion") and segment.emotion:
                emotion = segment.emotion
                emoji = EMOTION_EMOJI.get(emotion, "")
                if emoji:
                    p.add_run(f"{emoji} ")

            # Text with possible translation
            original_text = getattr(segment, "original_text", None)
            if original_text:
                p.add_run(original_text)
                trans_p = doc.add_paragraph()
                trans_run = trans_p.add_run(f"  → {segment.text}")
                trans_run.font.italic = True
            else:
                p.add_run(segment.text)

    # Save document
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    if filename is None:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"transcript_{timestamp}.docx"

    output_path = output_dir / filename
    doc.save(str(output_path))
    log.info(f"DCT transcript saved: {output_path}")

    return output_path


# Alias for pipeline compatibility
generate_transcript = generate_dct_transcript
