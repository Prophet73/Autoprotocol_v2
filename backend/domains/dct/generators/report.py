"""
DCT Domain Report Generator.

Generates DOCX reports for DCT meeting types:
- Brainstorm sessions
- Production meetings
- Negotiations
- Lectures/Webinars
"""
import logging
from typing import Optional, Any
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from backend.domains.dct.schemas import (
    DCTMeetingType,
    BrainstormResult,
    ProductionMeetingResult,
    NegotiationResult,
    LectureResult,
)

log = logging.getLogger(__name__)


def _add_field_to_docx(
    doc: Document,
    label: str,
    value: Any,
    is_list: bool = False,
    bold_label: bool = True,
    in_paragraph: Optional[Any] = None
):
    """Add a field to DOCX document."""
    if value or isinstance(value, (int, float, bool)):
        p = in_paragraph if in_paragraph else doc.add_paragraph()

        if label:
            if in_paragraph and p.text:
                p.add_run(" ")
            run = p.add_run(f"{label}: ")
            if bold_label:
                run.bold = True

        if is_list and isinstance(value, list):
            if not value:
                if label and not in_paragraph:
                    p.add_run("нет данных")
                elif label and in_paragraph:
                    p.add_run(" нет данных")
                else:
                    doc.add_paragraph("Нет данных", style='List Bullet')
            else:
                for item_idx, item_val in enumerate(value):
                    if in_paragraph and item_idx == 0 and label:
                        p.add_run(str(item_val))
                        if item_idx < len(value) - 1:
                            p.add_run("; ")
                    else:
                        doc.add_paragraph(str(item_val), style='List Bullet')
        elif isinstance(value, list) and not is_list:
            p.add_run("; ".join(map(str, value)))
        else:
            p.add_run(str(value))


def generate_brainstorm_docx(
    result: BrainstormResult,
    output_path: Path,
    meeting_date: Optional[str] = None
) -> Path:
    """Generate DOCX for brainstorm session."""
    doc = Document()

    title = "Итоги сессии мозгового штурма"
    if meeting_date:
        title += f" от {meeting_date}"
    doc.add_heading(title, level=1)

    _add_field_to_docx(doc, "Тема сессии", result.session_topic)
    _add_field_to_docx(doc, "Обсуждаемая проблема", result.main_problem)

    if result.idea_clusters:
        doc.add_heading("Сгенерированные идеи", level=2)
        for cluster in result.idea_clusters:
            doc.add_heading(f"Кластер: {cluster.cluster_name}", level=3)
            if cluster.ideas:
                for idea in cluster.ideas:
                    doc.add_paragraph(idea, style='List Bullet')

    if result.top_ideas:
        doc.add_heading("Наиболее перспективные идеи", level=2)
        for i, idea in enumerate(result.top_ideas, 1):
            p = doc.add_paragraph(f"{i}. {idea.idea_description}", style='List Number')
            if idea.potential_impact:
                _add_field_to_docx(doc, "Влияние", idea.potential_impact, in_paragraph=p, bold_label=False)
            if idea.implementation_complexity:
                _add_field_to_docx(doc, "Сложность", idea.implementation_complexity, in_paragraph=p, bold_label=False)

    if result.parked_ideas:
        doc.add_heading("Отложенные идеи (Парковка)", level=2)
        for idea in result.parked_ideas:
            doc.add_paragraph(idea, style='List Bullet')

    if result.next_steps:
        doc.add_heading("Следующие шаги", level=2)
        for i, step in enumerate(result.next_steps, 1):
            p = doc.add_paragraph(f"{i}. {step.action_item}", style='List Number')
            if step.responsible:
                _add_field_to_docx(doc, "Ответственный", step.responsible, in_paragraph=p, bold_label=False)

    doc.save(str(output_path))
    log.info(f"Brainstorm DOCX saved: {output_path}")
    return output_path


def generate_production_docx(
    result: ProductionMeetingResult,
    output_path: Path,
    meeting_date: Optional[str] = None
) -> Path:
    """Generate DOCX for production meeting."""
    doc = Document()

    title = "Протокол производственного совещания"
    if meeting_date:
        title += f" от {meeting_date}"
    doc.add_heading(title, level=1)

    _add_field_to_docx(doc, "Объект", result.object_name)

    if result.attendees:
        _add_field_to_docx(doc, "Присутствовали", result.attendees, is_list=True)

    if result.past_tasks_control:
        doc.add_heading("Контроль исполнения ранее поставленных задач", level=2)
        for i, task in enumerate(result.past_tasks_control, 1):
            p = doc.add_paragraph(f"{i}. {task.task_description}", style='List Number')
            _add_field_to_docx(doc, "Статус", task.status, in_paragraph=p, bold_label=False)
            if task.comment:
                _add_field_to_docx(doc, "Комментарий", task.comment, in_paragraph=p, bold_label=False)

    if result.work_progress_analysis:
        doc.add_heading("Анализ хода выполнения работ", level=2)
        for item in result.work_progress_analysis:
            p = doc.add_paragraph()
            p.add_run(f"{item.work_block_name}: ").bold = True
            p.add_run(item.status_summary)

    if result.resources_and_supply:
        doc.add_heading("Обеспеченность ресурсами и поставками", level=2)
        res = result.resources_and_supply
        _add_field_to_docx(doc, "Людские ресурсы", res.manpower)
        _add_field_to_docx(doc, "Техника", res.machinery)
        _add_field_to_docx(doc, "Материалы", res.materials)

    if result.safety_and_labor_protection:
        doc.add_heading("Вопросы охраны труда и ТБ", level=2)
        for item in result.safety_and_labor_protection:
            doc.add_paragraph(item, style='List Bullet')

    if result.new_tasks:
        doc.add_heading("Новые задачи", level=2)
        for i, task in enumerate(result.new_tasks, 1):
            p = doc.add_paragraph(f"{i}. {task.task_description}", style='List Number')
            _add_field_to_docx(doc, "Ответственный", task.responsible, in_paragraph=p, bold_label=False)
            if task.deadline:
                _add_field_to_docx(doc, "Срок", task.deadline, in_paragraph=p, bold_label=False)

    doc.save(str(output_path))
    log.info(f"Production meeting DOCX saved: {output_path}")
    return output_path


def generate_negotiation_docx(
    result: NegotiationResult,
    output_path: Path,
    meeting_date: Optional[str] = None
) -> Path:
    """Generate DOCX for negotiation protocol."""
    doc = Document()

    title = "Протокол переговоров"
    if meeting_date:
        title += f" от {meeting_date}"
    doc.add_heading(title, level=1)

    _add_field_to_docx(doc, "Цель встречи", result.meeting_goal)
    _add_field_to_docx(doc, "Компания-контрагент", result.counterpart_company)

    if result.topics_discussed:
        doc.add_heading("Обсуждаемые вопросы", level=2)
        for i, topic in enumerate(result.topics_discussed, 1):
            doc.add_heading(f"Тема {i}: {topic.topic_title}", level=3)
            _add_field_to_docx(doc, "Предложение", topic.proposal_summary, bold_label=False)
            if topic.value_for_company:
                _add_field_to_docx(doc, "Ценность для нас", topic.value_for_company, is_list=True, bold_label=False)
            if topic.risks_and_objections:
                _add_field_to_docx(doc, "Риски и возражения", topic.risks_and_objections, is_list=True, bold_label=False)
            if topic.terms_and_cost:
                _add_field_to_docx(doc, "Условия", topic.terms_and_cost, is_list=True, bold_label=False)

    if result.action_items:
        doc.add_heading("Принятые решения", level=2)
        if result.action_items.for_us:
            _add_field_to_docx(doc, "Задачи для нашей стороны", result.action_items.for_us, is_list=True)
        if result.action_items.for_counterpart:
            _add_field_to_docx(doc, "Задачи для контрагента", result.action_items.for_counterpart, is_list=True)

    if result.internal_strategic_analysis:
        doc.add_heading("Внутренний стратегический анализ", level=2)
        doc.add_paragraph(result.internal_strategic_analysis)

    doc.save(str(output_path))
    log.info(f"Negotiation DOCX saved: {output_path}")
    return output_path


def generate_lecture_docx(
    result: LectureResult,
    output_path: Path,
    meeting_date: Optional[str] = None
) -> Path:
    """Generate DOCX for lecture/webinar summary."""
    doc = Document()

    title = "Конспект лекции/вебинара"
    if meeting_date:
        title += f" от {meeting_date}"
    doc.add_heading(title, level=1)

    _add_field_to_docx(doc, "Название", result.webinar_title)

    if result.presentation_part:
        doc.add_heading("Основная часть", level=2)
        for block in result.presentation_part:
            doc.add_heading(block.block_title, level=3)
            if block.time_code:
                _add_field_to_docx(doc, "Тайм-код", block.time_code, bold_label=False)
            if block.key_idea:
                _add_field_to_docx(doc, "Ключевая мысль", block.key_idea, bold_label=False)
            if block.theses:
                doc.add_paragraph("Тезисы:")
                for thesis in block.theses:
                    doc.add_paragraph(thesis, style='List Bullet')

    if result.qa_part:
        doc.add_heading("Сессия Q&A", level=2)
        for qa_item in result.qa_part:
            doc.add_heading(f"Вопрос: {qa_item.question_title}", level=3)
            if qa_item.time_code:
                _add_field_to_docx(doc, "Тайм-код", qa_item.time_code, bold_label=False)
            if qa_item.key_answer_idea:
                _add_field_to_docx(doc, "Ответ", qa_item.key_answer_idea, bold_label=False)
            if qa_item.answer_theses:
                doc.add_paragraph("Тезисы ответа:")
                for thesis in qa_item.answer_theses:
                    doc.add_paragraph(thesis, style='List Bullet')

    if result.final_summary:
        doc.add_heading("Итоговые выводы", level=2)
        for item in result.final_summary:
            doc.add_paragraph(item, style='List Bullet')

    doc.save(str(output_path))
    log.info(f"Lecture DOCX saved: {output_path}")
    return output_path


def generate_dct_report(
    meeting_type: DCTMeetingType,
    result: Any,
    output_path: Path,
    meeting_date: Optional[str] = None
) -> Path:
    """
    Generate DOCX report based on meeting type.

    Args:
        meeting_type: Type of DCT meeting
        result: Parsed result object (BrainstormResult, ProductionMeetingResult, etc.)
        output_path: Path to save the DOCX file
        meeting_date: Optional meeting date string

    Returns:
        Path to the generated DOCX file
    """
    if meeting_type == DCTMeetingType.BRAINSTORM:
        return generate_brainstorm_docx(result, output_path, meeting_date)
    elif meeting_type == DCTMeetingType.PRODUCTION:
        return generate_production_docx(result, output_path, meeting_date)
    elif meeting_type == DCTMeetingType.NEGOTIATION:
        return generate_negotiation_docx(result, output_path, meeting_date)
    elif meeting_type == DCTMeetingType.LECTURE:
        return generate_lecture_docx(result, output_path, meeting_date)
    else:
        raise ValueError(f"Unknown DCT meeting type: {meeting_type}")
