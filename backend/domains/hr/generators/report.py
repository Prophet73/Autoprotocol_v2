"""
HR Report Generator.

Generates AI-analyzed HR meeting reports using Gemini.
"""
import os
import json
import logging
from pathlib import Path
from datetime import datetime
from typing import TYPE_CHECKING, Optional

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

if TYPE_CHECKING:
    from backend.core.transcription.result import TranscriptionResult

from backend.config import get_prompt
from backend.domains.hr.schemas import HRReport, HRMeetingType

logger = logging.getLogger(__name__)


def generate_report(
    result: "TranscriptionResult",
    output_dir: Path,
    meeting_type: str = "one_on_one",
    system_prompt: Optional[str] = None,
    user_prompt: Optional[str] = None,
) -> Path:
    """
    Generate AI-analyzed HR report.

    Args:
        result: Transcription result
        output_dir: Directory to save report
        meeting_type: Type of HR meeting
        system_prompt: Optional custom system prompt
        user_prompt: Optional custom user prompt

    Returns:
        Path to generated report
    """
    # Check for Gemini API key
    api_key = os.getenv("GOOGLE_API_KEY") or os.getenv("GEMINI_API_KEY")
    if not api_key:
        logger.warning("No Gemini API key found, generating basic report")
        return _generate_basic_report(result, output_dir, meeting_type)

    try:
        import google.generativeai as genai
        genai.configure(api_key=api_key)

        # Get prompts
        if not system_prompt:
            system_prompt = get_prompt(f"domains.hr.{meeting_type}.system")
        if not user_prompt:
            transcript_text = result.to_plain_text()
            user_prompt = get_prompt(
                f"domains.hr.{meeting_type}.user",
                transcript=transcript_text
            )

        # Call Gemini
        model = genai.GenerativeModel(
            os.getenv("GEMINI_MODEL_NAME", "gemini-2.0-flash-exp")
        )

        response = model.generate_content(
            [system_prompt, user_prompt],
            generation_config={
                "temperature": 0.3,
                "max_output_tokens": 4096,
            }
        )

        # Parse response and generate document
        return _generate_report_document(
            result, output_dir, meeting_type, response.text
        )

    except Exception as e:
        logger.error(f"Error generating HR report with Gemini: {e}")
        return _generate_basic_report(result, output_dir, meeting_type)


def _generate_basic_report(
    result: "TranscriptionResult",
    output_dir: Path,
    meeting_type: str
) -> Path:
    """Generate basic report without AI analysis."""
    doc = Document()

    # Title
    meeting_type_names = {
        "recruitment": "Отчёт о собеседовании",
        "one_on_one": "Отчёт о встрече 1-на-1",
        "performance_review": "Отчёт о Performance Review",
        "team_meeting": "Отчёт о командном совещании",
        "onboarding": "Отчёт об Onboarding",
    }

    title = doc.add_heading(meeting_type_names.get(meeting_type, "HR Отчёт"), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Info
    doc.add_paragraph(f"Дата: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    doc.add_paragraph(f"Файл: {result.source_file}")
    doc.add_paragraph()

    # Note about no AI
    doc.add_paragraph(
        "Примечание: ИИ-анализ недоступен. "
        "Ниже представлена базовая информация о встрече."
    )
    doc.add_paragraph()

    # Participants
    if result.speakers:
        doc.add_heading("Участники", level=1)
        for speaker_id, profile in result.speakers.items():
            emotion = getattr(profile, 'dominant_emotion', {})
            doc.add_paragraph(
                f"- {speaker_id}: "
                f"{int(getattr(profile, 'total_time', 0))} сек, "
                f"{getattr(profile, 'segment_count', 0)} реплик, "
                f"настроение: {emotion.get('label_ru', 'Н/Д')}"
            )

    # Key moments (first segments)
    doc.add_heading("Ключевые моменты", level=1)
    doc.add_paragraph("(Первые реплики встречи)")

    for i, segment in enumerate(result.segments[:10]):
        text = getattr(segment, 'text', '')[:200]
        if len(getattr(segment, 'text', '')) > 200:
            text += "..."
        doc.add_paragraph(f"{i+1}. {text}")

    # Save
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_path = output_dir / f"hr_report_{timestamp}.docx"
    doc.save(str(output_path))

    logger.info(f"Basic HR report saved to {output_path}")
    return output_path


def _generate_report_document(
    result: "TranscriptionResult",
    output_dir: Path,
    meeting_type: str,
    ai_analysis: str
) -> Path:
    """Generate report document with AI analysis."""
    doc = Document()

    meeting_type_names = {
        "recruitment": "Отчёт о собеседовании",
        "one_on_one": "Отчёт о встрече 1-на-1",
        "performance_review": "Отчёт о Performance Review",
        "team_meeting": "Отчёт о командном совещании",
        "onboarding": "Отчёт об Onboarding",
    }

    # Title
    title = doc.add_heading(meeting_type_names.get(meeting_type, "HR Отчёт"), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Info
    doc.add_paragraph(f"Дата: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    doc.add_paragraph(f"Файл: {result.source_file}")
    doc.add_paragraph()

    # AI Analysis
    doc.add_heading("ИИ-Анализ", level=1)
    doc.add_paragraph(ai_analysis)

    # Participants
    if result.speakers:
        doc.add_heading("Участники", level=1)
        for speaker_id, profile in result.speakers.items():
            emotion = getattr(profile, 'dominant_emotion', {})
            doc.add_paragraph(
                f"- {speaker_id}: "
                f"{int(getattr(profile, 'total_time', 0))} сек, "
                f"{getattr(profile, 'segment_count', 0)} реплик, "
                f"настроение: {emotion.get('label_ru', 'Н/Д')}"
            )

    # Save
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_path = output_dir / f"hr_report_{timestamp}.docx"
    doc.save(str(output_path))

    logger.info(f"HR report with AI analysis saved to {output_path}")
    return output_path
