"""
HR Transcript Generator.

Generates formatted transcript document for HR meetings.
"""
import logging
from pathlib import Path
from datetime import datetime
from typing import TYPE_CHECKING

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

if TYPE_CHECKING:
    from backend.core.transcription.result import TranscriptionResult

logger = logging.getLogger(__name__)


def generate_transcript(
    result: "TranscriptionResult",
    output_dir: Path,
    meeting_type: str = "one_on_one",
) -> Path:
    """
    Generate HR-formatted transcript document.

    Args:
        result: Transcription result with segments and speakers
        output_dir: Directory to save the document
        meeting_type: Type of HR meeting

    Returns:
        Path to generated document
    """
    doc = Document()

    # Title
    title = doc.add_heading("HR Meeting Transcript", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Meeting info
    meeting_type_names = {
        "recruitment": "Собеседование",
        "one_on_one": "Встреча 1-на-1",
        "performance_review": "Performance Review",
        "team_meeting": "Командное совещание",
        "onboarding": "Onboarding",
    }

    doc.add_paragraph(f"Тип встречи: {meeting_type_names.get(meeting_type, meeting_type)}")
    doc.add_paragraph(f"Дата: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    doc.add_paragraph(f"Файл: {result.source_file}")
    doc.add_paragraph()

    # Participants table
    if result.speakers:
        doc.add_heading("Участники", level=1)

        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Участник"
        hdr_cells[1].text = "Время (сек)"
        hdr_cells[2].text = "Реплик"
        hdr_cells[3].text = "Настроение"

        for cell in hdr_cells:
            cell.paragraphs[0].runs[0].bold = True

        # Data rows
        for speaker_id, profile in result.speakers.items():
            row = table.add_row().cells
            row[0].text = speaker_id
            row[1].text = str(int(getattr(profile, 'total_time', 0)))
            row[2].text = str(getattr(profile, 'segment_count', 0))

            emotion = getattr(profile, 'dominant_emotion', {})
            row[3].text = emotion.get('label_ru', 'Н/Д')

        doc.add_paragraph()

    # Transcript
    doc.add_heading("Транскрипция", level=1)

    for segment in result.segments:
        # Speaker and time
        speaker = getattr(segment, 'speaker', 'Unknown')
        start = getattr(segment, 'start', 0)
        end = getattr(segment, 'end', 0)
        text = getattr(segment, 'text', '')
        emotion = getattr(segment, 'emotion', None)

        # Format time as MM:SS
        start_fmt = f"{int(start // 60):02d}:{int(start % 60):02d}"
        end_fmt = f"{int(end // 60):02d}:{int(end % 60):02d}"

        # Add paragraph with speaker info
        p = doc.add_paragraph()
        run = p.add_run(f"[{start_fmt}-{end_fmt}] {speaker}")
        run.bold = True
        if emotion:
            run = p.add_run(f" ({emotion})")
            run.italic = True

        # Add text
        doc.add_paragraph(text)

    # Save
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_path = output_dir / f"hr_transcript_{timestamp}.docx"
    doc.save(str(output_path))

    logger.info(f"HR transcript saved to {output_path}")
    return output_path
