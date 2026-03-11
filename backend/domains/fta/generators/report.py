"""
FTA Domain Report Generator.

Generates DOCX reports for FTA audit meetings.
"""
import logging
from typing import Optional, Any
from pathlib import Path

from docx import Document

from backend.core.utils.docx_utils import add_field
from backend.domains.fta.schemas import FTAMeetingType, AuditResult

log = logging.getLogger(__name__)


def _generate_audit_docx(
    result: AuditResult,
    output_path: Path,
    meeting_date: Optional[str] = None
) -> Path:
    """Generate DOCX report for audit."""
    doc = Document()

    title = "Отчёт по результатам аудита"
    if meeting_date:
        title += f" от {meeting_date}"
    doc.add_heading(title, level=1)

    add_field(doc, "Предмет аудита", result.audit_subject)
    if result.audit_scope:
        add_field(doc, "Охват проверки", result.audit_scope)
    if result.overall_rating:
        add_field(doc, "Общая оценка", result.overall_rating)
    if result.participants:
        add_field(doc, "Участники", result.participants, is_list=True)

    if result.findings:
        doc.add_heading("Выявленные замечания", level=2)
        for i, finding in enumerate(result.findings, 1):
            p = doc.add_paragraph(f"{i}. {finding.finding}", style='List Number')
            add_field(
                doc, "Критичность", finding.severity,
                in_paragraph=p, bold_label=False,
            )
            if finding.area:
                add_field(
                    doc, "Область", finding.area,
                    in_paragraph=p, bold_label=False,
                )
            if finding.recommendation:
                add_field(
                    doc, "Рекомендация", finding.recommendation,
                    in_paragraph=p, bold_label=False,
                )

    if result.positive_observations:
        doc.add_heading("Положительные наблюдения", level=2)
        for obs in result.positive_observations:
            doc.add_paragraph(obs, style='List Bullet')

    if result.risks:
        doc.add_heading("Выявленные риски", level=2)
        for risk in result.risks:
            doc.add_paragraph(risk, style='List Bullet')

    if result.corrective_actions:
        doc.add_heading("Корректирующие мероприятия", level=2)
        for i, action in enumerate(result.corrective_actions, 1):
            p = doc.add_paragraph(f"{i}. {action.action}", style='List Number')
            if action.responsible:
                add_field(
                    doc, "Ответственный", action.responsible,
                    in_paragraph=p, bold_label=False,
                )
            if action.deadline:
                add_field(
                    doc, "Срок", action.deadline,
                    in_paragraph=p, bold_label=False,
                )

    if result.conclusions:
        doc.add_heading("Выводы и заключение", level=2)
        for conclusion in result.conclusions:
            doc.add_paragraph(conclusion, style='List Bullet')

    doc.save(str(output_path))
    log.info(f"FTA audit DOCX saved: {output_path}")
    return output_path


# =============================================================================
# Dispatcher
# =============================================================================

def generate_fta_report(
    meeting_type: FTAMeetingType,
    result: Any,
    output_path: Path,
    meeting_date: Optional[str] = None
) -> Path:
    """
    Generate DOCX report based on meeting type.

    Args:
        meeting_type: Type of FTA meeting
        result: Parsed result object
        output_path: Path to save the DOCX file
        meeting_date: Optional meeting date string

    Returns:
        Path to the generated DOCX file
    """
    generators = {
        FTAMeetingType.AUDIT: _generate_audit_docx,
    }

    generator = generators.get(meeting_type)
    if not generator:
        raise ValueError(f"Unknown FTA meeting type: {meeting_type}")

    return generator(result, output_path, meeting_date)
