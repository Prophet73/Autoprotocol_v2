"""
CEO Domain Report Generator.

Generates DOCX reports for CEO NOTECH meetings.
"""
import logging
from typing import Optional, Any
from pathlib import Path

from docx import Document

from backend.core.utils.docx_utils import add_field
from backend.domains.ceo.schemas import CEOMeetingType, NotechResult

log = logging.getLogger(__name__)


def _generate_notech_docx(
    result: NotechResult,
    output_path: Path,
    meeting_date: Optional[str] = None
) -> Path:
    """Generate DOCX report for NOTECH meeting."""
    doc = Document()

    title = "Протокол совещания НОТЕХ"
    if meeting_date:
        title += f" от {meeting_date}"
    doc.add_heading(title, level=1)

    add_field(doc, "Тема", result.meeting_topic)

    if result.summary:
        doc.add_heading("Краткое саммари", level=2)
        doc.add_paragraph(result.summary)

    if result.attendees:
        add_field(doc, "Участники", result.attendees, is_list=True)

    # Вопросы повестки — основное содержание протокола
    for i, q in enumerate(result.questions, 1):
        doc.add_heading(f"Вопрос №{i}: {q.title}", level=2)
        doc.add_paragraph(q.description)

        if q.value_points:
            p = doc.add_paragraph()
            run = p.add_run("Ценность:")
            run.bold = True
            for point in q.value_points:
                doc.add_paragraph(point, style='List Bullet')

        if q.decision:
            p = doc.add_paragraph()
            run = p.add_run("Решение: ")
            run.bold = True
            p.add_run(q.decision)

        if q.discussion_details:
            p = doc.add_paragraph()
            run = p.add_run("Детали обсуждения:")
            run.bold = True
            for detail in q.discussion_details:
                doc.add_paragraph(detail, style='List Bullet')

        if q.risks:
            p = doc.add_paragraph()
            run = p.add_run("Риски/Возражения:")
            run.bold = True
            for risk in q.risks:
                doc.add_paragraph(risk, style='List Bullet')

    # Принятые решения / задачи
    if result.action_items:
        doc.add_heading("Принятые решения", level=2)
        for item in result.action_items:
            doc.add_paragraph(item, style='List Bullet')

    doc.save(str(output_path))
    log.info(f"CEO NOTECH DOCX saved: {output_path}")
    return output_path


# =============================================================================
# Dispatcher
# =============================================================================

def generate_ceo_report(
    meeting_type: CEOMeetingType,
    result: Any,
    output_path: Path,
    meeting_date: Optional[str] = None
) -> Path:
    """
    Generate DOCX report based on meeting type.

    Args:
        meeting_type: Type of CEO meeting
        result: Parsed result object
        output_path: Path to save the DOCX file
        meeting_date: Optional meeting date string

    Returns:
        Path to the generated DOCX file
    """
    generators = {
        CEOMeetingType.NOTECH: _generate_notech_docx,
    }

    generator = generators.get(meeting_type)
    if not generator:
        raise ValueError(f"Unknown CEO meeting type: {meeting_type}")

    return generator(result, output_path, meeting_date)
